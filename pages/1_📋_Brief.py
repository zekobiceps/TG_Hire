import streamlit as st
import requests
import json
import time
from datetime import datetime
import os
import pickle
import re
import io
from urllib.parse import quote

# -------------------- Imports optionnels pour l'export --------------------
try:
    from reportlab.lib.pagesizes import A4
    from reportlab.platypus import SimpleDocTemplate, Table, TableStyle, Paragraph, Spacer
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib.units import inch
    from reportlab.lib import colors
    PDF_AVAILABLE = True
except ImportError:
    PDF_AVAILABLE = False

try:
    from docx import Document
    WORD_AVAILABLE = True
except ImportError:
    WORD_AVAILABLE = False

# -------------------- Checklist simplifiée --------------------
SIMPLIFIED_CHECKLIST = {
    "Contexte du Poste et Environnement": [
        "Pourquoi ce poste est-il ouvert?",
        "Fourchette budgétaire (entre X et Y)",
        "Date de prise de poste souhaitée",
        "Équipe (taille, composition)",
        "Manager (poste, expertise, style)",
        "Collaborations internes/externes",
        "Lieux de travail et déplacements"
    ],
    "Missions et Responsabilités": [
        "Mission principale du poste",
        "Objectifs à atteindre (3-5 maximum)",
        "Sur quoi la performance sera évaluée?",
        "3-5 Principales tâches quotidiennes",
        "2 Tâches les plus importantes/critiques",
        "Outils informatiques à maitriser"
    ],
    "Compétences - Modèle KSA": [],
    "Profil et Formation": [
        "Expérience minimum requise",
        "Formation/diplôme nécessaire"
    ],
    "Stratégie de Recrutement": [
        "Pourquoi recruter maintenant?",
        "Difficultés anticipées",
        "Mot-clés cruciaux (CV screening)",
        "Canaux de sourcing prioritaires",
        "Plans B : Autres postes, Revoir certains critères...",
        "Exemple d'un profil cible sur LinkedIn",
        "Processus de sélection étape par étape"
    ]
}

# -------------------- Modèle KSA --------------------
KSA_MODEL = {
    "Knowledge (Connaissances)": [],
    "Skills (Savoir-faire)": [],
    "Abilities (Aptitudes)": []
}

# -------------------- Templates --------------------
BRIEF_TEMPLATES = {
    "Template Vide": {category: {item: {"valeur": "", "importance": 3} for item in items}
                      for category, items in SIMPLIFIED_CHECKLIST.items()},
    "Chargé de Recrutement": {
        "Contexte du Poste et Environnement": {
            "Pourquoi ce poste est-il ouvert?": {"valeur": "Départ d'un collaborateur", "importance": 3},
            "Fourchette budgétaire (entre X et Y)": {"valeur": "10 000 - 15 000 DH", "importance": 3},
            "Date de prise de poste souhaitée": {"valeur": "ASAP", "importance": 3},
            "Équipe (taille, composition)": {"valeur": "4 personnes", "importance": 3},
            "Manager (poste, expertise, style)": {"valeur": "Responsable RH", "importance": 3},
            "Collaborations internes/externes": {"valeur": "Managers opérationnels, prestataires", "importance": 3},
            "Lieux de travail et déplacements": {"valeur": "Casablanca avec déplacements ponctuels", "importance": 3}
        },
        "Missions et Responsabilités": {
            "Mission principale du poste": {"valeur": "Recrutement de profils techniques", "importance": 3},
            "Objectifs à atteindre (3-5 maximum)": {"valeur": "5 recrutements/mois", "importance": 3},
            "Sur quoi la performance sera évaluée?": {"valeur": "Taux de succès des recrutements", "importance": 3},
            "3-5 Principales tâches quotidiennes": {"valeur": "Sourcing, entretiens, reporting", "importance": 3},
            "2 Tâches les plus importantes/critiques": {"valeur": "Closing des candidats", "importance": 3},
            "Outils informatiques à maitriser": {"valeur": "LinkedIn Recruiter, Excel", "importance": 3}
        },
        "Compétences - Modèle KSA": {},
        "Profil et Formation": {
            "Expérience minimum requise": {"valeur": "2 ans en recrutement", "importance": 3},
            "Formation/diplôme nécessaire": {"valeur": "Bac+3 RH", "importance": 3}
        },
        "Stratégie de Recrutement": {
            "Pourquoi recruter maintenant?": {"valeur": "Croissance activité", "importance": 3},
            "Difficultés anticipées": {"valeur": "Marché pénurique sur profils techniques", "importance": 3},
            "Mot-clés cruciaux (CV screening)": {"valeur": "Recrutement, RH, Sourcing", "importance": 3},
            "Canaux de sourcing prioritaires": {"valeur": "LinkedIn, cooptation", "importance": 3},
            "Plans B : Autres postes, Revoir certains critères...": {"valeur": "", "importance": 3},
            "Exemple d'un profil cible sur LinkedIn": {"valeur": "", "importance": 3},
            "Processus de sélection étape par étape": {"valeur": "Entretien RH + Technique", "importance": 3}
        }
    }
}

# -------------------- Persistence --------------------
def save_briefs():
    with open("saved_briefs.pkl", "wb") as f:
        pickle.dump(st.session_state.saved_briefs, f)

def load_briefs():
    if os.path.exists("saved_briefs.pkl"):
        with open("saved_briefs.pkl", "rb") as f:
            return pickle.load(f)
    return {}

# -------------------- Utils --------------------
def generate_automatic_brief_name():
    now = datetime.now()
    return f"{now.strftime('%d%m%y')}-{st.session_state.get('poste_intitule','poste')}-{st.session_state.get('manager_nom','manager')}"

def filter_briefs(saved_briefs, month=None, recruteur=None, poste=None, manager=None):
    results = {}
    for name, data in saved_briefs.items():
        if month and month not in name:
            continue
        if recruteur and data.get("recruteur") != recruteur:
            continue
        if poste and poste.lower() not in data.get("poste_intitule","").lower():
            continue
        if manager and manager.lower() not in data.get("manager_nom","").lower():
            continue
        results[name] = data
    return results

# -------------------- Init session --------------------
def init_session_state():
    defaults = {
        "poste_intitule": "",
        "manager_nom": "",
        "recruteur": "Zakaria",
        "affectation_type": "Chantier",
        "affectation_nom": "",
        "current_brief_name": "",
        "brief_data": {category: {item: {"valeur": "", "importance": 3} for item in items}
                       for category, items in SIMPLIFIED_CHECKLIST.items()},
        "ksa_data": {},
        "comment_libre": "",
        "saved_briefs": load_briefs(),
        "filtered_briefs": {},
        "show_filtered_results": False,
    }
    for k, v in defaults.items():
        if k not in st.session_state:
            st.session_state[k] = v

init_session_state()

# -------------------- Page config --------------------
st.set_page_config(page_title="TG-Hire IA - Brief", page_icon="📋", layout="wide")

st.title("📋 Brief Recrutement")

# -------------------- Sélecteur de phase --------------------
brief_phase = st.radio("Phase du Brief:", ["📁 Gestion", "🔄 Avant-brief", "✅ Réunion de brief"], horizontal=True)
st.session_state.brief_phase = brief_phase

# -------------------- Phase Gestion --------------------
if brief_phase == "📁 Gestion":
    col1, col2 = st.columns(2)
    with col1:
        st.subheader("Informations de base")
        st.session_state.poste_intitule = st.text_input("Intitulé du poste", st.session_state.poste_intitule)
        st.session_state.manager_nom = st.text_input("Nom du manager", st.session_state.manager_nom)
        st.session_state.recruteur = st.selectbox("Recruteur", ["Zakaria", "Sara", "Jalal", "Bouchra"], index=0)

        st.session_state.affectation_type = st.selectbox("Affectation", ["Chantier", "Direction"], index=0)
        st.session_state.affectation_nom = st.text_input("Nom de l'affectation", st.session_state.affectation_nom)

        if st.session_state.poste_intitule and st.session_state.manager_nom:
            suggested = generate_automatic_brief_name()
            st.session_state.current_brief_name = st.text_input("Nom du brief", value=suggested)

    with col2:
        st.subheader("Chargement & Templates")
        filter_month = st.selectbox("Mois", [""]+[f"{i:02d}" for i in range(1,13)])
        filter_recruteur = st.selectbox("Recruteur", ["","Zakaria","Sara","Jalal","Bouchra"])
        filter_poste = st.text_input("Poste")
        filter_manager = st.text_input("Manager")

        if st.button("🔍 Rechercher briefs"):
            st.session_state.filtered_briefs = filter_briefs(st.session_state.saved_briefs, filter_month, filter_recruteur, filter_poste, filter_manager)
            st.session_state.show_filtered_results = True

        if st.session_state.show_filtered_results:
            if st.session_state.filtered_briefs:
                selected = st.selectbox("Choisir un brief", [""]+list(st.session_state.filtered_briefs.keys()))
                if selected and st.button("📂 Charger"):
                    loaded = st.session_state.filtered_briefs[selected]
                    for k in ["poste_intitule","manager_nom","recruteur","affectation_type","affectation_nom","brief_data","ksa_data","comment_libre"]:
                        st.session_state[k] = loaded.get(k, st.session_state[k])
                    st.success("Brief chargé")
                    st.session_state.current_brief_name = selected
                if selected and st.button("🗑️ Supprimer"):
                    del st.session_state.saved_briefs[selected]
                    save_briefs()
                    st.rerun()
            else:
                st.info("Aucun brief trouvé")

        template_choice = st.selectbox("Choisir un template", list(BRIEF_TEMPLATES.keys()))
        if st.button("Appliquer template"):
            st.session_state.brief_data = BRIEF_TEMPLATES[template_choice].copy()
            st.success("Template appliqué")

    if st.button("💾 Sauvegarder le brief"):
        if not st.session_state.current_brief_name:
            st.session_state.current_brief_name = generate_automatic_brief_name()
        st.session_state.saved_briefs[st.session_state.current_brief_name] = {
            "poste_intitule": st.session_state.poste_intitule,
            "manager_nom": st.session_state.manager_nom,
            "recruteur": st.session_state.recruteur,
            "affectation_type": st.session_state.affectation_type,
            "affectation_nom": st.session_state.affectation_nom,
            "brief_data": st.session_state.brief_data,
            "ksa_data": st.session_state.ksa_data,
            "comment_libre": st.session_state.comment_libre
        }
        save_briefs()
        st.success("Brief sauvegardé")

# -------------------- Avant-brief & Réunion --------------------
elif brief_phase in ["🔄 Avant-brief","✅ Réunion de brief"]:
    st.subheader(f"{brief_phase}")
    categories = list(SIMPLIFIED_CHECKLIST.keys())
    section = st.selectbox("Section", categories)

    if section == "Compétences - Modèle KSA":
        st.subheader("📊 Modèle KSA")
        for cat in KSA_MODEL.keys():
            st.markdown(f"**{cat}**")
            new = st.text_input(f"Ajouter compétence {cat}", key=f"new_{cat}")
            if st.button(f"Ajouter {cat}", key=f"btn_{cat}") and new:
                if cat not in st.session_state.ksa_data:
                    st.session_state.ksa_data[cat] = {}
                st.session_state.ksa_data[cat][new] = {"niveau":"Intermédiaire","priorite":"Indispensable","evaluateur":"Manager"}
                st.rerun()

            if cat in st.session_state.ksa_data:
                for comp,details in st.session_state.ksa_data[cat].items():
                    col1,col2,col3,col4,col5 = st.columns([2,1,1,1,1])
                    with col1: st.text(comp)
                    with col2: details["niveau"] = st.selectbox("Niveau",["Débutant","Intermédiaire","Expert"], key=f"niv_{cat}_{comp}")
                    with col3: details["priorite"] = st.selectbox("Priorité",["Indispensable","Souhaitable"], key=f"pri_{cat}_{comp}")
                    with col4: details["evaluateur"] = st.selectbox("Évaluateur",["Manager","Recruteur","Les deux"], key=f"eval_{cat}_{comp}")
                    with col5:
                        if st.button("🗑️", key=f"del_{cat}_{comp}"):
                            del st.session_state.ksa_data[cat][comp]
                            st.rerun()
    else:
        st.subheader(section)
        for item in SIMPLIFIED_CHECKLIST[section]:
            current = st.session_state.brief_data[section].get(item, {"valeur":"","importance":3})
            st.session_state.brief_data[section][item]["valeur"] = st.text_area(item, current["valeur"], height=80)

    st.subheader("💬 Commentaires libres")
    st.session_state.comment_libre = st.text_area("Vos notes", st.session_state.comment_libre, height=120)

    if st.button("💾 Sauvegarder"):
        if not st.session_state.current_brief_name:
            st.session_state.current_brief_name = generate_automatic_brief_name()
        st.session_state.saved_briefs[st.session_state.current_brief_name] = {
            "poste_intitule": st.session_state.poste_intitule,
            "manager_nom": st.session_state.manager_nom,
            "recruteur": st.session_state.recruteur,
            "affectation_type": st.session_state.affectation_type,
            "affectation_nom": st.session_state.affectation_nom,
            "brief_data": st.session_state.brief_data,
            "ksa_data": st.session_state.ksa_data,
            "comment_libre": st.session_state.comment_libre
        }
        save_briefs()
        st.success("Brief sauvegardé")

# -------------------- Export --------------------
st.markdown("---")
st.subheader("📄 Export du Brief")
col1,col2 = st.columns(2)
with col1:
    if PDF_AVAILABLE and st.button("📄 Exporter PDF"):
        buffer = io.BytesIO()
        doc = SimpleDocTemplate(buffer, pagesize=A4)
        styles = getSampleStyleSheet()
        story = [Paragraph(f"Brief - {st.session_state.poste_intitule}", styles["Title"])]
        for category,items in st.session_state.brief_data.items():
            story.append(Paragraph(category, styles["Heading2"]))
            for item,data in items.items():
                if data["valeur"]:
                    story.append(Paragraph(f"{item}: {data['valeur']}", styles["Normal"]))
        doc.build(story)
        buffer.seek(0)
        st.download_button("⬇️ Télécharger PDF", buffer, "brief.pdf", "application/pdf")
with col2:
    if WORD_AVAILABLE and st.button("📄 Exporter Word"):
        doc = Document()
        doc.add_heading(f"Brief - {st.session_state.poste_intitule}", 0)
        for category,items in st.session_state.brief_data.items():
            doc.add_heading(category, level=1)
            for item,data in items.items():
                if data["valeur"]:
                    doc.add_heading(item, level=2)
                    doc.add_paragraph(data["valeur"])
        buffer = io.BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        st.download_button("⬇️ Télécharger Word", buffer, "brief.docx", "application/vnd.openxmlformats-officedocument.wordprocessingml.document")
