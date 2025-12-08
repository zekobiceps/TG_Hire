import streamlit as st
import sys, os
from datetime import datetime, timedelta
import json
import pandas as pd
from datetime import date
import random
from io import BytesIO
from textwrap import wrap
try:
    from docx import Document
    WORD_LIB_CUSTOM = True
except Exception:
    WORD_LIB_CUSTOM = False
try:
    from reportlab.lib.pagesizes import A4
    from reportlab.pdfgen import canvas
    from reportlab.lib.units import mm
    PDF_LIB_OK = True
except Exception:
    PDF_LIB_OK = False

# ✅ permet d'accéder à utils.py à la racine
sys.path.append(os.path.abspath(os.path.join(os.path.dirname(__file__), "..")))

from utils import (
    init_session_state,
    PDF_AVAILABLE,
    WORD_AVAILABLE,
    save_briefs,
    load_briefs,
    get_example_for_field,
    export_brief_pdf,
    export_brief_word,
    generate_checklist_advice,
    filter_briefs,
    generate_automatic_brief_name,
    save_library,
    generate_ai_question,
    test_deepseek_connection,
    save_brief_to_gsheet,
    export_brief_pdf_pretty,
    refresh_saved_briefs,
    load_all_local_briefs,
    save_ksa_matrix_to_current_brief   # <-- AJOUT
)

# --- CSS pour augmenter la taille du texte des onglets ---
st.markdown("""
    <style>
    .stTabs [data-baseweb="tab"]{height:auto!important;padding:10px 16px!important;}
    .stButton > button{border-radius:4px;padding:0.5rem 1rem;font-weight:500;white-space:nowrap;}
    .streamlit-expanderHeader{font-weight:600;}
    .stDataFrame{width:100%;}
    .stTextArea textarea{min-height:100px;resize:vertical;white-space:pre-wrap!important;}
    /* cible plus spécifique pour les boutons dans les forms (st.form) */
    .ai-red-btn button, .ai-red-btn input[type="submit"]{background:#C40000!important;color:#fff!important;border:1px solid #960000!important;font-weight:600!important;}
    .ai-red-btn button:hover, .ai-red-btn input[type="submit"]:hover{background:#E00000!important;}
    form .ai-red-btn button, .stForm .ai-red-btn button, form .ai-red-btn input[type="submit"], .stForm .ai-red-btn input[type="submit"] { background:#C40000!important; color:#fff!important; border:1px solid #960000!important; }
    .ai-suggestion-box{background:linear-gradient(135deg,#e6ffed,#f4fff7);border-left:5px solid #23C552;padding:0.6rem 0.8rem;margin:0.3rem 0 0.8rem;border-radius:6px;font-size:.9rem;}
    .score-cible{font-size:24px!important;font-weight:700;color:#0057B7;margin-top:.5rem;}
    .brief-name{padding-top:0.05rem;font-size:1.05rem;font-weight:650;display:flex;align-items:center;}
    .filtered-brief-line{display:flex;align-items:center;gap:0.4rem;margin-bottom:0.15rem;}
    .filtered-brief-line .actions{display:flex;gap:0.3rem;}
    .filtered-brief-line button{padding:0.25rem 0.6rem !important;font-size:0.70rem !important;}
    .synthese-two-cols{display:flex;gap:2rem;}
    .synthese-col{flex:1;min-width:200px;}
    .synthese-col h4{margin-bottom:0.4rem;}
    </style>
""", unsafe_allow_html=True)

def generate_custom_pdf(brief_name:str, data:dict, ksa:pd.DataFrame|None)->BytesIO|None:
    if not PDF_LIB_OK:
        return None
    try:
        from reportlab.lib.pagesizes import A4
        from reportlab.pdfgen import canvas
        buffer = BytesIO()
        c = canvas.Canvas(buffer, pagesize=A4)
        w,h = A4
        y = h-40
        # Recherche robuste du logo et centrage
        possible_paths = [
            os.path.join(os.path.dirname(__file__), '..', 'tgcc.png'),
            os.path.join(os.path.dirname(__file__), '..', 'static', 'tgcc.png'),
            os.path.join(os.getcwd(), 'tgcc.png')
        ]
        logo_path = next((p for p in possible_paths if os.path.exists(p)), None)
        if logo_path:
            try:
                logo_width = 140
                x_center = (w - logo_width) / 2
                c.drawImage(logo_path, x_center, y-30, width=logo_width, preserveAspectRatio=True, mask='auto')
            except Exception:
                pass
        c.setFont('Helvetica-Bold',18); c.drawString(40,y-10,f"Brief - {brief_name}")
        meta = [
            f"Poste : {data.get('POSTE_INTITULE') or data.get('poste_intitule','')}",
            f"Manager : {data.get('MANAGER_NOM') or data.get('manager_nom','')}",
            f"Recruteur : {data.get('RECRUTEUR') or data.get('recruteur','')}",
            f"Affectation : {data.get('AFFECTATION_NOM') or data.get('affectation_nom','')} ({data.get('AFFECTATION_TYPE') or data.get('affectation_type','')})",
            f"Date : {data.get('DATE_BRIEF') or data.get('date_brief','')}"
        ]
        y-=55; c.setFont('Helvetica',10)
        for line in meta:
            c.drawString(40,y,line); y-=14
        sections_pdf=[
            ("Contexte",["RAISON_OUVERTURE","IMPACT_STRATEGIQUE","TACHES_PRINCIPALES"]),
            ("Must-have",["MUST_HAVE_EXP","MUST_HAVE_DIP","MUST_HAVE_COMPETENCES","MUST_HAVE_SOFTSKILLS"]),
            ("Nice-to-have",["NICE_TO_HAVE_EXP","NICE_TO_HAVE_DIP","NICE_TO_HAVE_COMPETENCES"]),
            ("Conditions",["RATTACHEMENT","BUDGET"]),
            ("Sourcing",["ENTREPRISES_PROFIL","SYNONYMES_POSTE","CANAUX_PROFIL"]),
            ("Notes",["COMMENTAIRES","NOTES_LIBRES"])
        ]
        for title,keys in sections_pdf:
            c.setFont('Helvetica-Bold',12); c.drawString(40,y-4,title); y-=18; c.setFont('Helvetica',9)
            for k in keys:
                val=data.get(k,'');
                if not val: continue
                wrapped=wrap(f"- {k.title().replace('_',' ')} : {val}",110)
                for wline in wrapped:
                    if y<60: c.showPage(); y=h-60; c.setFont('Helvetica',9)
                    c.drawString(50,y,wline); y-=12
            y-=4
        if ksa is not None and isinstance(ksa,pd.DataFrame) and not ksa.empty:
            if y<140: c.showPage(); y=h-60
            c.setFont('Helvetica-Bold',12); c.drawString(40,y,'Matrice KSA'); y-=20
            cols=["Rubrique","Critère","Type de question","Évaluation (1-5)"]
            col_widths=[70,200,130,80]
            x0=40
            header_y=y
            c.setFont('Helvetica-Bold',9)
            for i,col in enumerate(cols):
                c.drawString(x0+sum(col_widths[:i]), header_y, col)
            y=header_y-14
            c.setFont('Helvetica',8)
            for _,row in ksa.iterrows():
                crit=str(row.get('Critère',''))
                crit_lines=wrap(crit,45) or ['']
                rub=str(row.get('Rubrique',''))
                tq=str(row.get('Type de question',''))
                evalv=str(row.get('Évaluation (1-5)',''))
                for li,txt in enumerate(crit_lines):
                    if y<60:
                        c.showPage(); y=h-60; header_y=y; c.setFont('Helvetica-Bold',9)
                        for i,col in enumerate(cols):
                            c.drawString(x0+sum(col_widths[:i]), header_y, col)
                        y=header_y-14; c.setFont('Helvetica',8)
                    if li==0:
                        c.drawString(x0, y, rub)
                        c.drawString(x0+col_widths[0], y, txt)
                        c.drawString(x0+col_widths[0]+col_widths[1], y, tq)
                        c.drawString(x0+col_widths[0]+col_widths[1]+col_widths[2], y, evalv)
                    else:
                        c.drawString(x0+col_widths[0], y, txt)
                    y-=10
                y-=2
        c.showPage(); c.save(); buffer.seek(0); return buffer
    except Exception:
        return None

def generate_custom_word(brief_name:str,data:dict,ksa:pd.DataFrame|None)->BytesIO|None:
    if not WORD_LIB_CUSTOM:
        return None
    try:
        doc=Document()
        logo_path = os.path.join(os.path.dirname(__file__), '..', 'tgcc.png')
        if os.path.exists(logo_path):
            try:
                from docx.shared import Inches
                doc.add_picture(logo_path, width=Inches(2.2))
            except Exception:
                pass
        doc.add_heading(f"Brief - {brief_name}",0)
        for label,key in [("Poste","POSTE_INTITULE"),("Manager","MANAGER_NOM"),("Recruteur","RECRUTEUR"),("Affectation","AFFECTATION_NOM"),("Date","DATE_BRIEF")]:
            p=doc.add_paragraph(); p.add_run(f"{label} : ").bold=True; p.add_run(str(data.get(key) or data.get(key.lower(),'') or ''))
        mapping=[
            ("Contexte",["RAISON_OUVERTURE","IMPACT_STRATEGIQUE","TACHES_PRINCIPALES"]),
            ("Must-have",["MUST_HAVE_EXP","MUST_HAVE_DIP","MUST_HAVE_COMPETENCES","MUST_HAVE_SOFTSKILLS"]),
            ("Nice-to-have",["NICE_TO_HAVE_EXP","NICE_TO_HAVE_DIP","NICE_TO_HAVE_COMPETENCES"]),
            ("Conditions",["RATTACHEMENT","BUDGET"]),
            ("Sourcing",["ENTREPRISES_PROFIL","SYNONYMES_POSTE","CANAUX_PROFIL"]),
            ("Notes",["COMMENTAIRES","NOTES_LIBRES"])
        ]
        for title,keys in mapping:
            doc.add_heading(title,2)
            for k in keys:
                val=data.get(k,'')
                if val:
                    doc.add_paragraph(f"{k.title().replace('_',' ')} : {val}",style='List Bullet')
        if ksa is not None and isinstance(ksa,pd.DataFrame) and not ksa.empty:
            doc.add_heading('Matrice KSA',2)
            table=doc.add_table(rows=1,cols=4)
            hdrs=["Rubrique","Critère","Type","Éval"]
            for i,h in enumerate(hdrs): table.rows[0].cells[i].text=h
            for _,row in ksa.iterrows():
                cells=table.add_row().cells
                cells[0].text=str(row.get('Rubrique',''))
                cells[1].text=str(row.get('Critère',''))
                cells[2].text=str(row.get('Type de question',''))
                cells[3].text=str(row.get('Évaluation (1-5)',''))
        bio=BytesIO(); doc.save(bio); bio.seek(0); return bio
    except Exception:
        return None

def render_ksa_matrix():
    """Affiche la matrice KSA sous forme de tableau et permet l'ajout de critères."""
    try:
        with st.expander("ℹ️ Explications de la méthode KSA", expanded=False):
            st.markdown("""### Méthode KSA (Knowledge, Skills, Abilities)
La matrice KSA permet d'évaluer un candidat sur trois dimensions complémentaires :

**Knowledge (Connaissances)** – Savoirs théoriques / réglementaires.  
**Skills (Compétences)** – Savoirs-faire techniques & opérationnels.  
**Abilities (Aptitudes)** – Capacités comportementales / cognitives dans l'action.

| Type de question | Objectif | Cible K/S/A |
|------------------|----------|-------------|
| Comportementale (STAR) | Situation réelle passée | Abilities (principal) + Skills |
| Technique | Vérifier la maîtrise concrète | Knowledge & Skills |
| Situationnelle | Réaction hypothétique | Abilities + Skills |
| Générale | Motivation / structuration | Mix selon formulation |

Échelle (1–5) : 1=Insuffisant | 3=Autonome attendu | 5=Référent.  
Limiter à 4–7 critères forts. Une question = un critère avec cible claire.
""")
        if "ksa_matrix" not in st.session_state:
            st.session_state.ksa_matrix = pd.DataFrame(columns=[
                "Rubrique", "Critère", "Type de question", "Cible / Standard attendu",
                "Échelle d'évaluation (1-5)", "Évaluateur"
            ])

        placeholder_dict = {
            "Comportementale": "Ex: Décrivez une situation où vous avez géré une équipe sous pression (méthode STAR).",
            "Situationnelle": "Ex: Que feriez-vous si un délai de chantier était menacé par un retard de livraison ?",
            "Technique": "Ex: Expliquez comment vous utilisez AutoCAD pour la modélisation de structures BTP.",
            "Générale": "Ex: Parlez-moi de votre expérience globale dans le secteur BTP."
        }

        with st.expander("➕ Ajouter un critère", expanded=True):
            # (bloc inchangé sauf que toutes les occurrences unsafe_allow_html sont confirmées)
            with st.form(key="add_ksa_criterion_form"):
                st.markdown("<div style='padding: 5px;'>", unsafe_allow_html=True)
                cible = st.text_area("Cible / Standard attendu",
                                     placeholder="Définissez la cible ou le standard attendu pour ce critère.",
                                     key="new_cible", height=100,
                                     value=st.session_state.get("new_cible", ""))
                st.markdown("</div>", unsafe_allow_html=True)

                col1, col2, col3 = st.columns([1, 1, 1.5])
                with col1:
                    st.markdown("<div style='padding: 5px;'>", unsafe_allow_html=True)
                    rubrique = st.selectbox(
                        "Rubrique", ["Knowledge", "Skills", "Abilities"],
                        key="new_rubrique",
                        index=["Knowledge", "Skills", "Abilities"].index(
                            st.session_state.get("new_rubrique", "Knowledge")
                        ) if st.session_state.get("new_rubrique") in ["Knowledge", "Skills", "Abilities"] else 0
                    )
                    st.markdown("</div>", unsafe_allow_html=True)
                with col2:
                    st.markdown("<div style='padding: 5px;'>", unsafe_allow_html=True)
                    critere = st.text_input("Critère", key="new_critere",
                                            value=st.session_state.get("new_critere", ""))
                    type_question = st.selectbox(
                        "Type de question",
                        ["Comportementale", "Situationnelle", "Technique", "Générale"],
                        key="new_type_question",
                        index=["Comportementale", "Situationnelle", "Technique", "Générale"].index(
                            st.session_state.get("new_type_question", "Comportementale")
                        ) if st.session_state.get("new_type_question") in
                           ["Comportementale", "Situationnelle", "Technique", "Générale"] else 0
                    )
                    st.markdown("</div>", unsafe_allow_html=True)
                with col3:
                    st.markdown("<div style='padding: 5px;'>", unsafe_allow_html=True)
                    evaluation = st.slider("Échelle d'évaluation (1-5)", 1, 5,
                                           value=st.session_state.get("new_evaluation", 3),
                                           step=1, key="new_evaluation")
                    evaluateur = st.selectbox(
                        "Évaluateur", ["Manager", "Recruteur", "Les deux"],
                        key="new_evaluateur",
                        index=["Manager", "Recruteur", "Les deux"].index(
                            st.session_state.get("new_evaluateur", "Manager")
                        ) if st.session_state.get("new_evaluateur") in
                           ["Manager", "Recruteur", "Les deux"] else 0
                    )
                    st.markdown("</div>", unsafe_allow_html=True)

                st.markdown("---")
                st.markdown("**Demander une question à l'IA**")
                ai_prompt = st.text_input(
                    "Décrivez ce que l'IA doit générer :",
                    placeholder="Ex: une question générale sur l'expérience en gestion de projet",
                    key="ai_prompt", value=st.session_state.get("ai_prompt", "")
                )
                st.checkbox("⚡ Mode rapide (réponse concise)", key="concise_checkbox")

                col_buttons = st.columns([1,1])
                with col_buttons[0]:
                    st.markdown("<div class='ai-red-btn'>", unsafe_allow_html=True)
                    gen_click = st.form_submit_button("💡 Générer question IA", width="stretch")
                    st.markdown("</div>", unsafe_allow_html=True)
                with col_buttons[1]:
                    add_click = st.form_submit_button("➕ Ajouter le critère", width="stretch")
                if gen_click:
                    if ai_prompt:
                        try:
                            ai_response = generate_ai_question(ai_prompt, concise=st.session_state.concise_checkbox)
                            st.session_state.ai_response = ai_response
                        except Exception as e:
                            st.error(f"Erreur génération IA : {e}")
                    else:
                        st.error("Veuillez entrer un prompt pour l'IA")
                if add_click:
                        if not critere or not cible:
                            st.error("Veuillez remplir au moins le critère et la cible.")
                        else:
                            new_row = pd.DataFrame([{
                                "Rubrique": rubrique,
                                "Critère": critere,
                                "Type de question": type_question,
                                "Cible / Standard attendu": cible,
                                "Échelle d'évaluation (1-5)": evaluation,
                                "Évaluateur": evaluateur
                            }])
                            st.session_state.ksa_matrix = pd.concat(
                                [st.session_state.ksa_matrix, new_row],
                                ignore_index=True
                            )
                            st.success("✅ Critère ajouté.")
                            st.rerun()

            if "ai_response" in st.session_state and st.session_state.ai_response:
                st.success(f"**Question :** {st.session_state.ai_response}")
                st.session_state.ai_response = ""

        if not st.session_state.ksa_matrix.empty:
            st.session_state.ksa_matrix = st.data_editor(
                st.session_state.ksa_matrix,
                hide_index=True,
                column_config={
                    "Rubrique": st.column_config.SelectboxColumn(
                        "Rubrique", options=["Knowledge", "Skills", "Abilities"], required=True),
                    "Critère": st.column_config.TextColumn("Critère", required=True),
                    "Type de question": st.column_config.SelectboxColumn(
                        "Type de question",
                        options=["Comportementale", "Situationnelle", "Technique", "Générale"],
                        required=True),
                    "Cible / Standard attendu": st.column_config.TextColumn(
                        "Cible / Standard attendu", required=True),
                    "Échelle d'évaluation (1-5)": st.column_config.NumberColumn(
                        "Échelle d'évaluation (1-5)", min_value=1, max_value=5, step=1, format="%d"),
                    "Évaluateur": st.column_config.SelectboxColumn(
                        "Évaluateur", options=["Manager", "Recruteur", "Les deux"], required=True),
                },
                num_rows="dynamic",
                width="stretch",
            )
    except Exception as e:
        st.error(f"❌ Erreur dans render_ksa_matrix: {e}")

def delete_current_brief():
    """Supprime le brief actuel et retourne à l'onglet Gestion"""
    if "current_brief_name" in st.session_state and st.session_state.current_brief_name:
        brief_name = st.session_state.current_brief_name
        file_path = os.path.join("briefs", f"{brief_name}.json")
        if os.path.exists(file_path):
            os.remove(file_path)
            st.session_state.saved_briefs.pop(brief_name, None)
            save_briefs()
            
            st.session_state.current_brief_name = ""
            st.session_state.avant_brief_completed = True
            st.session_state.reunion_completed = True
            st.session_state.reunion_step = 1
            
            keys_to_reset = [
                "manager_nom", "niveau_hierarchique", "affectation_type", 
                "recruteur", "affectation_nom", "date_brief", "raison_ouverture",
                "impact_strategique", "rattachement", "taches_principales",
                "must_have_experience", "must_have_diplomes", "must_have_competences",
                "must_have_softskills", "nice_to_have_experience", "nice_to_have_diplomes",
                "nice_to_have_competences", "entreprises_profil", "synonymes_poste",
                "canaux_profil", "budget", "commentaires", "notes_libres"
            ]
            
            for key in keys_to_reset:
                if key in st.session_state:
                    del st.session_state[key]
            
            st.success(f"✅ Brief '{brief_name}' supprimé avec succès")
            st.session_state.brief_phase = "📁 Gestion"
            # Met à jour la liste des briefs et les statistiques
            try:
                refresh_saved_briefs()
            except Exception:
                pass
            st.rerun()

# ---------------- INIT ----------------
init_session_state()
st.set_page_config(
    page_title="Brief de Calibration",
    page_icon=None,
    layout="wide",
    initial_sidebar_state="expanded"
)

if "brief_phase" not in st.session_state:
    st.session_state.brief_phase = "📁 Gestion"

if "reunion_step" not in st.session_state:
    st.session_state.reunion_step = 1

if "filtered_briefs" not in st.session_state:
    st.session_state.filtered_briefs = {}
if "show_filtered_results" not in st.session_state:
    # Contrôle d'affichage des briefs filtrés (par défaut caché tant qu'on n'a pas cliqué sur Filtrer)
    st.session_state.show_filtered_results = False

if "avant_brief_completed" not in st.session_state:
    st.session_state.avant_brief_completed = True

if "reunion_completed" not in st.session_state:
    st.session_state.reunion_completed = True

if "current_tab" not in st.session_state:
    st.session_state.current_tab = "Gestion"

if "save_message" not in st.session_state:
    st.session_state.save_message = None
if "save_message_tab" not in st.session_state:
    st.session_state.save_message_tab = None

if "import_brief_flag" not in st.session_state:
    st.session_state.import_brief_flag = False
if "brief_to_import" not in st.session_state:
    st.session_state.brief_to_import = None

key_mapping = {
    "POSTE_INTITULE": "poste_intitule",
    "MANAGER_NOM": "manager_nom",
    "RECRUTEUR": "recruteur",
    "AFFECTATION_TYPE": "affectation_type",
    "AFFECTATION_NOM": "affectation_nom",
    "DATE_BRIEF": "date_brief",
    "RAISON_OUVERTURE": "raison_ouverture",
    "IMPACT_STRATEGIQUE": "impact_strategique",
    "TACHES_PRINCIPALES": "taches_principales",
    "MUST_HAVE_EXP": "must_have_experience",
    "MUST_HAVE_DIP": "must_have_diplomes",
    "MUST_HAVE_COMPETENCES": "must_have_competences",
    "MUST_HAVE_SOFTSKILLS": "must_have_softskills",
    "NICE_TO_HAVE_EXP": "nice_to_have_experience",
    "NICE_TO_HAVE_DIP": "nice_to_have_diplomes",
    "NICE_TO_HAVE_COMPETENCES": "nice_to_have_competences",
    "RATTACHEMENT": "rattachement",
    "BUDGET": "budget",
    "ENTREPRISES_PROFIL": "entreprises_profil",
    "SYNONYMES_POSTE": "synonymes_poste",
    "CANAUX_PROFIL": "canaux_profil",
    "LIEN_PROFIL_1": "profil_link_1",
    "LIEN_PROFIL_2": "profil_link_2",
    "LIEN_PROFIL_3": "profil_link_3",
    "COMMENTAIRES": "commentaires",
    "NOTES_LIBRES": "notes_libres"
}

sections = [
    {"title": "Contexte du poste", "fields": [
        ("Raison de l'ouverture", "raison_ouverture", ""),
        ("Impact stratégique", "impact_strategique", ""),
        ("Tâches principales", "taches_principales", "")]},
    {"title": "Must-have (Indispensables)", "fields": [
        ("Expérience", "must_have_experience", ""),
        ("Connaissances / Diplômes / Certifications", "must_have_diplomes", ""),
        ("Compétences / Outils", "must_have_competences", ""),
        ("Soft skills / aptitudes comportementales", "must_have_softskills", "")]},
    {"title": "Nice-to-have (Atouts)", "fields": [
        ("Expérience additionnelle", "nice_to_have_experience", ""),
        ("Diplômes / Certifications valorisantes", "nice_to_have_diplomes", ""),
        ("Compétences complémentaires", "nice_to_have_competences", "")]},
    {"title": "Conditions et contraintes", "fields": [
        ("Localisation", "rattachement", ""),
        ("Budget recrutement", "budget", "")]},
    {"title": "Sourcing et marché", "fields": [
        ("Entreprises où trouver ce profil", "entreprises_profil", ""),
        ("Synonymes / intitulés proches", "synonymes_poste", ""),
        ("Canaux à utiliser", "canaux_profil", "")]},
    {"title": "Profils pertinents", "fields": [
        ("Lien profil 1", "profil_link_1", ""),
        ("Lien profil 2", "profil_link_2", ""),
        ("Lien profil 3", "profil_link_3", "")]},
    {"title": "Notes libres", "fields": [
        ("Points à discuter ou à clarifier avec le manager", "commentaires", ""),
        ("Case libre", "notes_libres", "")]},
]

# --- Bloc d'import automatique ---
if st.session_state.import_brief_flag and st.session_state.brief_to_import:
    brief = load_briefs().get(st.session_state.brief_to_import, {})
    for sheet_key, session_key in key_mapping.items():
        st.session_state[session_key] = brief.get(sheet_key, "")
    for section in sections:
        for title, field_key, _ in section["fields"]:
            unique_key = f"{section['title'].replace(' ', '_')}_{field_key}"
            st.session_state[unique_key] = brief.get(key_mapping.get(field_key.upper(), field_key), "")
    st.session_state.current_brief_name = st.session_state.brief_to_import
    st.session_state.avant_brief_completed = True
    st.session_state.reunion_completed = True
    st.session_state.reunion_step = 1
    st.session_state.import_brief_flag = False
    st.session_state.brief_to_import = None
    st.rerun()
# -------------------- Sidebar --------------------
with st.sidebar:
    st.title("📊 Statistiques Brief")

    # Rafraîchit (fusion) avant calcul
    merged = refresh_saved_briefs()
    total_briefs = len(merged)
    st.metric("📋 Briefs créés", total_briefs)

    st.divider()
    if st.button("Tester Connexion IA", key="test_deepseek"):
        test_deepseek_connection()

# ---------------- NAVIGATION PRINCIPALE ----------------
st.title("Brief de Calibration")

# Define tabs before using them
tabs = st.tabs([
    "📁 Gestion",
    "🔄 Avant-brief",
    "✅ Réunion de brief",
    "📝 Synthèse"
])

st.markdown("""
<style>
div[data-testid="stTabs"] button p {
    font-size: 18px;
}
.stTextArea textarea {
    white-space: pre-wrap !important;
}
div[data-baseweb="input"] > div,
div[data-baseweb="textarea"] > div,
div[data-baseweb="select"] > div {
    background-color: var(--background-color) !important;
    border-color: var(--border-color) !important;
}
.stTextInput > div > div > input,
.stTextArea > div > div > textarea,
.stSelectbox > div > div > select,
.stDateInput > div > div > input {
    background-color: var(--background-color) !important;
    color: var(--text-color) !important;
    border: 1px solid var(--border-color) !important;
}
:root {
    --background-color: #f0f2f6;
    --text-color: #31333F;
    --border-color: #d0d0d0;
}
@media (prefers-color-scheme: dark) {
    :root {
        --background-color: #0e1117;
        --text-color: #fafafa;
        --border-color: #555555;
    }
}
</style>
""", unsafe_allow_html=True)

if "current_brief_name" not in st.session_state:
    st.session_state.current_brief_name = ""

all_field_keys = [
    "raison_ouverture", "impact_strategique", "taches_principales",
    "must_have_experience", "must_have_diplomes", "must_have_competences", "must_have_softskills",
    "nice_to_have_experience", "nice_to_have_diplomes", "nice_to_have_competences",
    "rattachement", "budget",
    "entreprises_profil", "synonymes_poste", "canaux_profil",
    "profil_link_1", "profil_link_2", "profil_link_3",
    "commentaires", "notes_libres",
    "canaux_prioritaires", "criteres_exclusion", "processus_evaluation", "manager_notes"
]

# ---------------- ONGLET GESTION ----------------
with tabs[0]:
    brief_data = st.session_state.saved_briefs.get(st.session_state.current_brief_name, {}) if st.session_state.current_brief_name else {}

    col_left, col_right = st.columns([2, 2])

    # Bloc "Créer un brief"
    with col_left:
        st.markdown('<h3 style="margin-bottom: 0.3rem;">📋 Informations de base</h3>', unsafe_allow_html=True)
        col1, col2, col3 = st.columns(3)
        with col1:
            st.text_input("Poste à recruter", key="poste_intitule", value=brief_data.get("poste_intitule", ""))
        with col2:
            st.text_input("Manager", key="manager_nom", value=brief_data.get("manager_nom", ""))
        with col3:
            st.selectbox("Recruteur", ["Zakaria", "Jalal", "Sara", "Ghita", "Bouchra"], key="recruteur",
                index=["Zakaria", "Jalal", "Sara", "Ghita", "Bouchra"].index(brief_data.get("recruteur", "Zakaria")) if brief_data.get("recruteur", "Zakaria") in ["Zakaria", "Jalal", "Sara", "Ghita", "Bouchra"] else 0)
        col4, col5, col6 = st.columns(3)
        with col4:
            st.selectbox("Type d'affectation", ["Chantier", "Siège", "Dépôt"], key="affectation_type",
                index=["Chantier", "Siège", "Dépôt"].index(brief_data.get("affectation_type", "Chantier")) if brief_data.get("affectation_type", "Chantier") in ["Chantier", "Siège", "Dépôt"] else 0)
        with col5:
            st.text_input("Nom affectation", key="affectation_nom", value=brief_data.get("affectation_nom", ""))
        with col6:
            # --- Normalisation robuste de la date avant le widget ---
            def _parse_date_any(v):
                if isinstance(v, date):
                    return v
                if isinstance(v, datetime):
                    return v.date()
                if isinstance(v, str):
                    for fmt in ("%Y-%m-%d", "%d/%m/%Y", "%Y/%m/%d"):
                        try:
                            return datetime.strptime(v, fmt).date()
                        except:
                            continue
                return date.today()

            # Valeur prioritaire : session_state si déjà existante
            if "date_brief" in st.session_state:
                # Sécurise le type avant d'appeler le widget (sinon warning)
                if not isinstance(st.session_state.date_brief, (date, datetime)):
                    st.session_state.date_brief = _parse_date_any(st.session_state.date_brief)
                elif isinstance(st.session_state.date_brief, datetime):
                    st.session_state.date_brief = st.session_state.date_brief.date()
                # IMPORTANT : ne pas passer 'value=' si key déjà présent (évite le warning jaune)
                chosen_date = st.date_input("Date du brief", key="date_brief")
            else:
                raw = brief_data.get("date_brief", brief_data.get("DATE_BRIEF", date.today()))
                date_brief_value = _parse_date_any(raw)
                chosen_date = st.date_input("Date du brief", value=date_brief_value, key="date_brief")

        # Si un brief est actuellement sélectionné, proposer de sauvegarder les modifications
        if st.session_state.get("current_brief_name"):
            if st.button("💾 Sauvegarder les modifications", type="primary", width="stretch", key="save_existing_brief"):
                current = st.session_state.current_brief_name
                bd = st.session_state.saved_briefs.get(current, {})
                # Mettre à jour les champs de base
                for low in ["poste_intitule","manager_nom","recruteur","affectation_type","affectation_nom","date_brief"]:
                    if low in st.session_state:
                        v = st.session_state.get(low, "")
                        bd[low] = v
                        bd[low.upper()] = v

                # Champs étendus
                for k in all_field_keys:
                    v = st.session_state.get(k, "")
                    bd[k] = v
                    if k.startswith("profil_link_"):
                        suffix = k.split("_")[-1]
                        bd[f"LIEN_PROFIL_{suffix}"] = v
                    else:
                        bd[k.upper()] = v

                # KSA
                if "ksa_matrix" in st.session_state and isinstance(st.session_state.ksa_matrix, pd.DataFrame):
                    bd["KSA_MATRIX_JSON"] = st.session_state.ksa_matrix.to_json(orient="records", force_ascii=False)

                st.session_state.saved_briefs[current] = bd
                save_briefs()
                # Google Sheets save is async, doesn't block the UI
                try:
                    save_brief_to_gsheet(current, bd)
                except Exception:
                    pass  # Local save succeeded, Sheets is optional
                refresh_saved_briefs()
                st.success("✅ Modifications sauvegardées.")
                st.rerun()
        else:
            # Mode création d'un nouveau brief
            if st.button("💾 Créer brief", type="primary", width="stretch", key="create_brief"):
                required_fields = ["poste_intitule", "manager_nom", "affectation_nom", "date_brief"]
                missing_fields = [field for field in required_fields if not st.session_state.get(field)]
                if missing_fields:
                    st.error(f"Veuillez remplir les champs suivants : {', '.join(missing_fields)}")
                else:
                    date_arg = st.session_state.date_brief
                    if isinstance(date_arg, str):
                        for fmt in ("%Y-%m-%d", "%d/%m/%Y"):
                            try:
                                date_arg = datetime.strptime(date_arg, fmt).date()
                                break
                            except:
                                continue
                    if isinstance(date_arg, datetime):
                        date_arg = date_arg.date()

                    new_brief_name = generate_automatic_brief_name(
                        st.session_state.poste_intitule,
                        st.session_state.manager_nom,
                        date_arg
                    )
                    
                    date_str = ""
                    if isinstance(date_arg, (date, datetime)):
                        date_str = date_arg.strftime("%Y-%m-%d")
                    else:
                        date_str = str(date_arg)

                    new_brief_data = {}
                    base_pairs = {
                        "poste_intitule": "POSTE_INTITULE",
                        "manager_nom": "MANAGER_NOM",
                        "recruteur": "RECRUTEUR",
                        "affectation_type": "AFFECTATION_TYPE",
                        "affectation_nom": "AFFECTATION_NOM",
                        "date_brief": "DATE_BRIEF"
                    }
                    for low, up in base_pairs.items():
                        v = date_str if low == "date_brief" else st.session_state.get(low, "")
                        new_brief_data[low] = v
                        new_brief_data[up] = v

                    for k in all_field_keys:
                        v = st.session_state.get(k, "")
                        new_brief_data[k] = v
                        if k.startswith("profil_link_"):
                            suffix = k.split("_")[-1]
                            new_brief_data[f"LIEN_PROFIL_{suffix}"] = v
                        else:
                            new_brief_data[k.upper()] = v

                    new_brief_data["BRIEF_NAME"] = new_brief_name

                    # Mise à jour session + sauvegarde
                    st.session_state.saved_briefs[new_brief_name] = new_brief_data
                    st.session_state.current_brief_name = new_brief_name
                    st.session_state.reunion_step = 1
                    st.session_state.reunion_completed = False
                    save_briefs()
                    # Google Sheets save is async, doesn't block the UI
                    try:
                        save_brief_to_gsheet(new_brief_name, new_brief_data)
                    except Exception:
                        pass
                    # Recalcule les stats immédiatement
                    refresh_saved_briefs()
                    st.success(f"✅ Brief '{new_brief_name}' créé avec succès !")
                    st.rerun()

    # Bloc "Filtrer les briefs"
    with col_right:
        st.markdown('<h3 style="margin-bottom: 0.3rem;">🔍 Filtrer les briefs</h3>', unsafe_allow_html=True)
        col_filter1, col_filter2, col_filter3 = st.columns(3)
        with col_filter1:
            st.date_input("Date", key="filter_date", value=None)
        with col_filter2:
            st.text_input("Recruteur", key="filter_recruteur", value=st.session_state.get("filter_recruteur", ""))
        with col_filter3:
            st.text_input("Manager", key="filter_manager", value=st.session_state.get("filter_manager", ""))

        col_filter4, col_filter5, col_filter6 = st.columns(3)
        with col_filter4:
            st.selectbox("Affectation", ["", "Chantier", "Siège", "Dépôt"], key="filter_affectation",
                        index=["", "Chantier", "Siège", "Dépôt"].index(st.session_state.get("filter_affectation", ""))
                        if st.session_state.get("filter_affectation") in ["", "Chantier", "Siège", "Dépôt"] else 0)
        with col_filter5:
            st.text_input("Nom affectation", key="filter_nom_affectation", value=st.session_state.get("filter_nom_affectation", ""))
        with col_filter6:
            st.selectbox("Type de brief", ["", "Standard", "Urgent", "Stratégique"], key="filter_brief_type",
                        index=["", "Standard", "Urgent", "Stratégique"].index(st.session_state.get("filter_brief_type", ""))
                        if st.session_state.get("filter_brief_type") in ["", "Standard", "Urgent", "Stratégique"] else 0)

        if st.button("🔎 Filtrer", width="stretch", key="apply_filter"):
            filter_month = st.session_state.filter_date.strftime("%m") if st.session_state.filter_date else ""
            st.session_state.filtered_briefs = filter_briefs(
                st.session_state.saved_briefs,
                filter_month,
                st.session_state.filter_recruteur,
                st.session_state.filter_brief_type,
                st.session_state.filter_manager,
                st.session_state.filter_affectation,
                st.session_state.filter_nom_affectation
            )
            st.session_state.show_filtered_results = True
            st.rerun()

        # Résultats filtrés (uniquement après clic sur Filtrer)
        if st.session_state.get("show_filtered_results", False):
            briefs_to_show = st.session_state.filtered_briefs
            st.markdown('<h3 style="margin-top: 1rem; margin-bottom: 0.3rem;">📋 Briefs sauvegardés</h3>', unsafe_allow_html=True)
            if briefs_to_show:
                for name, brief in briefs_to_show.items():
                    coln, cola = st.columns([5,4])
                    with coln:
                        st.markdown(f"<div class='brief-name' style='font-size:1rem;font-weight:700;'>{name}</div>", unsafe_allow_html=True)
                    with cola:
                        c_edit, c_del = st.columns([1,1])
                        with c_edit:
                            if st.button("📝 Éditer", key=f"edit_{name}", width="stretch"):
                                st.session_state.import_brief_flag=True; st.session_state.brief_to_import=name; st.rerun()
                        with c_del:
                            if st.button("🗑️ Supprimer", key=f"del_{name}", width="stretch"):
                                # Suppression locale
                                st.session_state.saved_briefs.pop(name, None)
                                save_briefs()
                                st.session_state.filtered_briefs.pop(name, None)
                                # Tentative suppression dans Google Sheets
                                try:
                                    from utils import delete_brief_from_gsheet, refresh_saved_briefs
                                    delete_brief_from_gsheet(name)
                                    refresh_saved_briefs()
                                except Exception:
                                    try:
                                        refresh_saved_briefs()
                                    except Exception:
                                        pass
                                st.success("Supprimé")
                                st.rerun()
            else:
                st.info("Aucun brief sauvegardé ou correspondant aux filtres.")
            # CSS pour réduire la taille des boutons éditer/supprimer
            st.markdown("""
            <style>
            [data-testid='stButton'] button {
                font-size: 13px !important;
                padding: 2px 10px !important;
                min-width: 50px !important;
                height: 28px !important;
            }
            .brief-name { font-size: 1rem !important; font-weight: 700 !important; }
            </style>
            """, unsafe_allow_html=True)

# ---------------- AVANT-BRIEF ----------------
# Dans l'onglet Avant-brief (tabs[1])
with tabs[1]:
    brief_data = load_briefs().get(st.session_state.current_brief_name, {}) if st.session_state.current_brief_name else {}

    if not st.session_state.current_brief_name:
        st.info("Crée ou sélectionne un brief dans l’onglet Gestion.")
    else:
        st.markdown("### 💡 Assistance IA (conseils ciblés)")
        c_ai1, c_ai2 = st.columns([3,1])
        with c_ai1:
            # Exclure la section 'Profils pertinents' du menu déroulant
            ai_field_options = [f"{s['title']} ➜ {title}"
                                for s in sections if s['title'] != 'Profils pertinents'
                                for title,_,_ in s["fields"]]
            ai_selected = st.selectbox("Champ à enrichir", ai_field_options, key="ab_ai_field")
        with c_ai2:
            st.markdown("<div class='ai-red-btn'>", unsafe_allow_html=True)
            if st.button("💡 Générer suggestion IA", key="btn_ai_suggestion"):
                sec, fld = ai_selected.split(" ➜ ",1)
                for s in sections:
                    if s["title"] == sec:
                        for title, key, _ in s["fields"]:
                            if title == fld:
                                adv = generate_checklist_advice(s["title"], title)
                                ex = get_example_for_field(s["title"], title)
                                st.session_state[f"advice_{key}"] = adv + ("\nExemple:\n" + ex if ex else "")
                                break
            st.markdown("</div>", unsafe_allow_html=True)

        with st.form(key="avant_brief_form"):
            for section in sections:
                with st.expander(f"📋 {section['title']}", expanded=False):
                    for title, key, placeholder in section["fields"]:
                        sheet_key = key.upper()
                        
                        # Sanitize session_state to prevent TypeError in text_area
                        if key in st.session_state:
                            if st.session_state[key] is None:
                                st.session_state[key] = ""
                            else:
                                # Force string conversion for any non-string type in session_state
                                st.session_state[key] = str(st.session_state[key])

                        # Ensure value is a string, handle None
                        # Fallback strategy: 
                        # 1. session_state[key]
                        # 2. brief_data[key.upper()] (Standard)
                        # 3. brief_data[key] (Lowercase fallback)
                        # 4. brief_data[Legacy] (Old keys fallback)
                        
                        legacy_map = {
                            "MUST_HAVE_EXPERIENCE": "MUST_HAVE_EXP",
                            "MUST_HAVE_DIPLOMES": "MUST_HAVE_DIP",
                            "NICE_TO_HAVE_EXPERIENCE": "NICE_TO_HAVE_EXP",
                            "NICE_TO_HAVE_DIPLOMES": "NICE_TO_HAVE_DIP"
                        }
                        
                        val_from_data = brief_data.get(sheet_key)
                        if val_from_data is None or val_from_data == "":
                            val_from_data = brief_data.get(key)
                        if (val_from_data is None or val_from_data == "") and sheet_key in legacy_map:
                            val_from_data = brief_data.get(legacy_map[sheet_key])
                            
                        raw_value = st.session_state.get(key, val_from_data if val_from_data is not None else "")
                        value = str(raw_value) if raw_value is not None else ""
                        
                        st.text_area(
                            title,
                            value=value,
                            key=key,
                            placeholder=placeholder,
                            height=140
                        )
                        if st.session_state.get(f"advice_{key}", ""):
                            suggestion_html = st.session_state[f"advice_{key}"].replace('\n', '<br>')
                            st.markdown(f"<div class='ai-suggestion-box'><strong>Suggestion IA :</strong><br>{suggestion_html}</div>", unsafe_allow_html=True)
            if st.form_submit_button("💾 Enregistrer modifications", type="primary"):
                if st.session_state.current_brief_name:
                    current = st.session_state.current_brief_name
                    bd = st.session_state.saved_briefs.get(current, {})
                    for s in sections:
                        for _, key, _ in s["fields"]:
                            val = st.session_state.get(key,"")
                            bd[key] = val
                            if key.startswith("profil_link_"):
                                suf = key.split("_")[-1]
                                bd[f"LIEN_PROFIL_{suf}"] = val
                            else:
                                bd[key.upper()] = val
                    # Champs de base si modifiés
                    for low in ["poste_intitule","manager_nom","recruteur","affectation_type","affectation_nom","date_brief"]:
                        if low in st.session_state:
                            v = st.session_state.get(low,"")
                            bd[low] = v
                            bd[low.upper()] = v
                    st.session_state.saved_briefs[current] = bd
                    save_briefs()
                    
                    # Google Sheets save
                    try:
                        save_brief_to_gsheet(current, bd)
                    except Exception as e:
                        st.error(f"Erreur lors de la sauvegarde Google Sheets: {e}")
                        
                    st.success("Avant-brief sauvegardé avec succès.")


# ================== REUNION DE BRIEF (MODIFS) ==================
with tabs[2]:
    st.subheader("✅ Réunion de brief")
    if not st.session_state.current_brief_name:
        st.info("Sélectionnez ou créez un brief.")
    else:
        total_steps = 4
        if "reunion_step" not in st.session_state or not isinstance(st.session_state.reunion_step, int):
            st.session_state.reunion_step = 1
        step = st.session_state.reunion_step
        st.progress(int(((step-1)/(total_steps-1))*100), text=f"Étape {step}/{total_steps}")

        brief_data = st.session_state.saved_briefs.get(st.session_state.current_brief_name, {})
        try:
            manager_comments = json.loads(brief_data.get("MANAGER_COMMENTS_JSON","{}"))
        except Exception:
            manager_comments = {}

        # ----- Étape 1 -----
        if step == 1:
            st.markdown("### 📝 Étape 1 : Vue consolidée & commentaires manager")
            rows = []
            for s in sections:
                if s["title"] == "Profils pertinents":
                    continue
                for title, key, _ in s["fields"]:
                    rows.append({
                        "Section": s["title"],
                        "Item": title,
                        "Infos": brief_data.get(key.upper()) or brief_data.get(key,""),
                        "Commentaire manager": manager_comments.get(key,""),
                        "_key": key
                    })
            if not rows:
                st.warning("Avant-brief vide.")
            else:
                base_df = pd.DataFrame(rows)
                display_df = base_df.drop(columns=["_key"])
                edited_df = st.data_editor(
                    display_df,
                    hide_index=True,
                    column_config={
                        "Section": st.column_config.TextColumn(disabled=True),
                        "Item": st.column_config.TextColumn(disabled=True),
                        "Infos": st.column_config.TextColumn(disabled=True),
                        "Commentaire manager": st.column_config.TextColumn()
                    },
                    key="etape1_editor",
                    width="stretch"
                )
                if st.button("💾 Enregistrer commentaires", key="save_mgr_step1"):
                    new_com = {}
                    for i, row in edited_df.iterrows():
                        val = row["Commentaire manager"]
                        orig = base_df.loc[i, "_key"]
                        if val:
                            new_com[orig] = val
                    brief_data["manager_comments"] = new_com
                    brief_data["MANAGER_COMMENTS_JSON"] = json.dumps(new_com, ensure_ascii=False)
                    st.session_state.saved_briefs[st.session_state.current_brief_name] = brief_data
                    save_briefs()
                    # Google Sheets save is async, doesn't block the UI
                    try:
                        save_brief_to_gsheet(st.session_state.current_brief_name, brief_data)
                    except Exception:
                        pass
                    st.success("Commentaires sauvegardés.")
                    st.rerun()

        # ----- Étape 2 -----
        elif step == 2:
            st.markdown("### 📊 Étape 2 : Matrice KSA")
            with st.expander("ℹ️ Matrice KSA : Les Critères d'Évaluation", expanded=False):
                st.markdown("""**Matrice KSA : Les Critères d'Évaluation 🛠️🧠🤝**
La Matrice KSA (Knowledge, Skills, Abilities) permet de structurer l'évaluation du candidat en trois piliers :

**Knowledge (K) 📚 :** Connaissances théoriques.
**Skills (S) 🔧 :** Compétences pratiques / opérationnelles.
**Abilities (A) 🧭 :** Aptitudes comportementales (méthode STAR pour évaluer une situation réelle passée).

**Choix du type de question :**
• Comportementale (STAR) → surtout Abilities (A).  
• Technique → Knowledge (K) & Skills (S).  
• Situationnelle → Abilities (A) + Skills (S) sur un scénario hypothétique.  
• Générale → Vision / motivation / structuration.  

**Échelle 1-5 – Cible Standard :** 1=Insuffisant, 3=Acceptable (autonome moyen terme), 5=Expert (peut former).  
**Bonnes pratiques :** 4–7 critères, une question = un critère, évaluateur défini, cible claire.
""")

            # Charger JSON existant si DataFrame vide
            if ("ksa_matrix" not in st.session_state or
                not isinstance(st.session_state.ksa_matrix, pd.DataFrame) or
                st.session_state.ksa_matrix.empty):
                js = brief_data.get("KSA_MATRIX_JSON","")
                if js:
                    try:
                        df_imp = pd.DataFrame(json.loads(js))
                        exp_cols = ["Rubrique","Critère","Type de question","Question pour l'entretien","Évaluation (1-5)","Évaluateur"]
                        for c in exp_cols:
                            if c not in df_imp.columns: df_imp[c] = ""
                        st.session_state.ksa_matrix = df_imp[exp_cols]
                    except Exception:
                        st.session_state.ksa_matrix = pd.DataFrame(columns=exp_cols)
                else:
                    st.session_state.ksa_matrix = pd.DataFrame(columns=[
                        "Rubrique","Critère","Type de question","Question pour l'entretien","Évaluation (1-5)","Évaluateur"
                    ])

            # Injection différée d'une question IA générée (évite conflit de state après widget)
            if 'ks_question_new' in st.session_state:
                st.session_state.ks_question = st.session_state.ks_question_new
                del st.session_state.ks_question_new
                st.rerun()

            with st.form("add_ksa_form_step2"):
                c1,c2,c3,c4 = st.columns(4)
                with c1:
                    rubrique = st.selectbox("Rubrique", ["Knowledge","Skills","Abilities"], key="ks_rubrique")
                with c2:
                    critere = st.text_input("Critère", key="ks_critere", value=st.session_state.get("ks_critere",""))
                with c3:
                    type_q = st.selectbox("Type de question", ["Comportementale","Situationnelle","Technique","Générale"], key="ks_type_q")
                with c4:
                    evaluateur = st.selectbox("Évaluateur", ["Recruteur","Manager","Les deux"], key="ks_evaluteur")

                qc, ec = st.columns([3,1])
                with qc:
                    question = st.text_area("Question pour l'entretien", key="ks_question",
                                            value=st.session_state.get("ks_question",""),
                                            height=70,
                                            placeholder="Ex: Si un retard critique survient sur deux lots en parallèle...")
                with ec:
                    evaluation = st.slider("Évaluation (1-5)", 1,5,3, key="ks_evaluation")

                ai_prompt = st.text_input("Prompt IA (génération question)", key="ks_ai_prompt",
                                          value=st.session_state.get("ks_ai_prompt", ""),
                                          placeholder="Ex: question comportementale sur gestion de conflit priorisation")
                # Par défaut, mode concis activé
                concise = st.checkbox("⚡ Mode rapide (réponse concise)", key="ks_concise", value=True)

                bcol1, bcol2 = st.columns([1,1])
                with bcol1:
                    st.markdown("<div class='ai-red-btn' style='margin-top:0.4rem;'>", unsafe_allow_html=True)
                    gen_btn = st.form_submit_button("💡 Générer question IA", width="stretch")
                    st.markdown("</div>", unsafe_allow_html=True)
                with bcol2:
                    st.markdown("<div class='add-crit-btn' style='margin-top:0.4rem;'>", unsafe_allow_html=True)
                    add_btn = st.form_submit_button("➕ Ajouter", width="stretch")
                    st.markdown("</div>", unsafe_allow_html=True)

                if gen_btn:
                    if not ai_prompt:
                        st.warning("Indique un prompt.")
                    else:
                        try:
                            with st.spinner("⏳ Génération en cours par l'IA..."):
                                resp = generate_ai_question(ai_prompt, concise=concise)
                            # Nettoyage : retire les '**' et le préfixe "Question:" ou "Réponse:"
                            resp = resp.replace('**', '')
                            if isinstance(resp, str) and resp.lower().startswith("question:"):
                                resp = resp.split(":",1)[1].strip()
                            elif isinstance(resp, str) and resp.lower().startswith("réponse:"):
                                resp = resp.split(":",1)[1].strip()
                            st.session_state.ks_generated_question = resp
                            st.success("Question générée.")
                            st.rerun()
                        except Exception as e:
                            st.error(f"Erreur IA: {e}")

                if add_btn:
                    if not critere or not question:
                        st.error("Critère + question requis.")
                    else:
                        new_row = {
                            "Rubrique": rubrique,
                            "Critère": critere,
                            "Type de question": type_q,
                            "Question pour l'entretien": question,
                            "Évaluation (1-5)": evaluation,
                            "Évaluateur": evaluateur
                        }
                        if st.session_state.ksa_matrix.empty:
                            st.session_state.ksa_matrix = pd.DataFrame([new_row])
                        else:
                            st.session_state.ksa_matrix = pd.concat(
                                [st.session_state.ksa_matrix, pd.DataFrame([new_row])],
                                ignore_index=True
                            )
                        save_ksa_matrix_to_current_brief()
                        st.success("Critère ajouté.")
                        # Nettoie l'affichage suggestion pour prochaine entrée
                        if 'ks_generated_question' in st.session_state:
                            del st.session_state.ks_generated_question
                        st.rerun()

                # Affichage de la question générée (si existe) sous les boutons
                if st.session_state.get('ks_generated_question'):
                    q_disp = st.session_state.ks_generated_question.replace('\n','<br>')
                    st.markdown(f"<div class='ai-suggestion-box' style='color:#228B22;font-size:16px;'>Question générée :<br>{q_disp}</div>", unsafe_allow_html=True)

            if not st.session_state.ksa_matrix.empty:
                cols_order = ["Rubrique","Critère","Type de question","Question pour l'entretien","Évaluation (1-5)","Évaluateur"]
                for c in cols_order:
                    if c not in st.session_state.ksa_matrix.columns:
                        st.session_state.ksa_matrix[c] = ""
                edited = st.data_editor(
                    st.session_state.ksa_matrix[cols_order],
                    hide_index=True,
                    width="stretch",
                    key="ksa_editor_step2",
                    column_config={
                        "Rubrique": st.column_config.SelectboxColumn("Rubrique", options=["Knowledge","Skills","Abilities"]),
                        "Type de question": st.column_config.SelectboxColumn("Type de question",
                            options=["Comportementale","Situationnelle","Technique","Générale"]),
                        "Évaluation (1-5)": st.column_config.NumberColumn("Évaluation (1-5)", min_value=1, max_value=5),
                        "Évaluateur": st.column_config.SelectboxColumn("Évaluateur", options=["Recruteur","Manager","Les deux"])
                    }
                )
                if not edited.equals(st.session_state.ksa_matrix[cols_order]):
                    st.session_state.ksa_matrix = edited
                    save_ksa_matrix_to_current_brief()
                # Score
                try:
                    avg = round(st.session_state.ksa_matrix["Évaluation (1-5)"].astype(float).mean(),2)
                    st.markdown(f"<div class='score-cible'>🎯 Score cible : {avg} / 5</div>", unsafe_allow_html=True)
                except Exception:
                    pass
            else:
                st.info("Aucun critère KSA pour l’instant.")

        # ----- Étape 3 -----
        elif step == 3:
            st.markdown("### 🛠️ Étape 3 : Stratégie & Processus")
            channels = ["LinkedIn","Jobboards","Jobzyn","Chasse de tête","Annonces","CVthèques"]
            if "canaux_prioritaires" not in st.session_state or not isinstance(st.session_state.canaux_prioritaires, list):
                st.session_state.canaux_prioritaires = []
            st.multiselect("Canaux prioritaires", channels,
                           key="canaux_prioritaires",
                           default=st.session_state.canaux_prioritaires)
            default_steps = ("1-Sourcing & Qualification (Recruteur):\nLe Recruteur trie, vérifie les critères éliminatoires (Budget, Mobilité) et qualifie les connaissances de base (Knowledge). Il envoie la shortlist qualifiée au Manager sous 5 jours.\n\n"
                             "2-Entretien Manager (Knowledge & Skills):\nLe Manager évalue la maîtrise technique et les compétences pratiques (Skills) du candidat, qui est convoqué à partir de la shortlist.\n\n"
                             "3-Entretien Recruteur (Aptitudes/Abilities):\nLe Recruteur se concentre sur l'évaluation des aptitudes comportementales (Abilities) et de l'adéquation à la culture d'entreprise (fit).\n\n"
                             "4-Clôture et Validation DRH:\nLe dossier complet du candidat retenu (avec l'avis KSA de chacun) est soumis à la DRH pour validation finale.")
            if not st.session_state.get("processus_evaluation") and not brief_data.get("PROCESSUS_EVALUATION"):
                st.session_state.processus_evaluation = default_steps
            c_excl, c_proc = st.columns([1,1.2])
            with c_excl:
                if "criteres_exclusion" in st.session_state and st.session_state["criteres_exclusion"] is None:
                    st.session_state["criteres_exclusion"] = ""
                val_excl = brief_data.get("CRITERES_EXCLUSION", st.session_state.get("criteres_exclusion",""))
                st.text_area("🚫 Critères d'exclusion", key="criteres_exclusion", height=260, value=str(val_excl) if val_excl is not None else "", placeholder="Ex: Moins de 3 ans d'expérience / Refus mobilité / Salaire hors budget")
            with c_proc:
                if "processus_evaluation" in st.session_state and st.session_state["processus_evaluation"] is None:
                    st.session_state["processus_evaluation"] = ""
                val_proc = st.session_state.get("processus_evaluation", brief_data.get("PROCESSUS_EVALUATION",""))
                st.text_area("✅ Etapes suivantes", key="processus_evaluation", height=420, value=str(val_proc) if val_proc is not None else "", placeholder="Ex: Étapes du processus d'évaluation pour ce poste")
            # Sauvegarde auto (pas de bouton)
            brief_data["canaux_prioritaires"] = st.session_state.get("canaux_prioritaires", [])
            brief_data["CANAUX_PRIORITAIRES"] = json.dumps(brief_data["canaux_prioritaires"], ensure_ascii=False)
            for low, up in {"criteres_exclusion":"CRITERES_EXCLUSION","processus_evaluation":"PROCESSUS_EVALUATION"}.items():
                val = st.session_state.get(low,"")
                brief_data[low] = val
                brief_data[up] = val
            st.session_state.saved_briefs[st.session_state.current_brief_name] = brief_data
            save_briefs()
            # Google Sheets save is async, doesn't block the UI
            try:
                save_brief_to_gsheet(st.session_state.current_brief_name, brief_data)
            except Exception:
                pass

        # ----- Étape 4 (NOUVELLE) : Validation finale -----
        elif step == 4:
            st.markdown("### ✅ Étape 4 : Validation finale")
            if "manager_notes" in st.session_state and st.session_state["manager_notes"] is None:
                st.session_state["manager_notes"] = ""
            val_notes = brief_data.get("MANAGER_NOTES", st.session_state.get("manager_notes",""))
            st.text_area("🗒️ Notes / commentaires finaux du manager",
                         key="manager_notes",
                         height=200,
                         value=str(val_notes) if val_notes is not None else "")
            if st.button("💾 Finaliser le brief", key="finalize_brief", type="primary"):
                brief_data["MANAGER_NOTES"] = st.session_state.get("manager_notes","")
                if "ksa_matrix" in st.session_state and not st.session_state.ksa_matrix.empty:
                    save_ksa_matrix_to_current_brief()
                st.session_state.saved_briefs[st.session_state.current_brief_name] = brief_data
                st.session_state.reunion_completed = True
                save_briefs()
                # Google Sheets save is mandatory here
                try:
                    if save_brief_to_gsheet(st.session_state.current_brief_name, brief_data):
                        st.success("Brief finalisé et sauvegardé sur Google Sheets.")
                        st.rerun()
                    else:
                        st.error("Erreur lors de la sauvegarde sur Google Sheets. Veuillez réessayer.")
                except Exception as e:
                    st.error(f"Erreur critique lors de la sauvegarde : {e}")

        # ----- Navigation -----
        nav_prev, nav_next = st.columns([1,1])
        with nav_prev:
            st.markdown("<div class='nav-small'>", unsafe_allow_html=True)
            if step > 1 and st.button("⬅️ Précédent", key=f"rb_prev_{step}", width="stretch"):
                st.session_state.reunion_step -= 1
                st.rerun()
            st.markdown("</div>", unsafe_allow_html=True)
        with nav_next:
            st.markdown("<div class='nav-small'>", unsafe_allow_html=True)
            if step < 4 and st.button("Suivant ➡️", key=f"rb_next_{step}", width="stretch"):
                # Auto-save léger (local uniquement pour les étapes intermédiaires)
                if step in (1,2,3):
                    brief_data["CRITERES_EXCLUSION"] = st.session_state.get("criteres_exclusion","")
                    brief_data["PROCESSUS_EVALUATION"] = st.session_state.get("processus_evaluation","")
                    brief_data["MANAGER_NOTES"] = st.session_state.get("manager_notes","")
                    if "ksa_matrix" in st.session_state:
                        save_ksa_matrix_to_current_brief()
                    st.session_state.saved_briefs[st.session_state.current_brief_name] = brief_data
                    save_briefs()
                st.session_state.reunion_step += 1
                st.rerun()
            st.markdown("</div>", unsafe_allow_html=True)

# ================== SYNTHÈSE (AJOUT Score moyen) ==================
with tabs[3]:
    if not st.session_state.current_brief_name:
        st.info("Sélectionnez un brief.")
    elif not st.session_state.reunion_completed:
        st.warning("Complétez la réunion avant la synthèse.")
    else:
        st.subheader(f"Synthèse - {st.session_state.current_brief_name}")
        bd = load_briefs().get(st.session_state.current_brief_name, {})
        st.write(f"- **Poste :** {bd.get('POSTE_INTITULE') or bd.get('poste_intitule','')}")
        st.write(f"- **Manager :** {bd.get('MANAGER_NOM') or bd.get('manager_nom','')}")
        st.write(f"- **Affectation :** {(bd.get('AFFECTATION_NOM') or '')} ({bd.get('AFFECTATION_TYPE') or ''})")
        st.write(f"- **Date :** {bd.get('DATE_BRIEF') or bd.get('date_brief','')}")

        st.markdown("### Détails")
        for s in sections:
            with st.expander(f"📋 {s['title']}", expanded=False):
                for title, key, _ in s["fields"]:
                    val = bd.get(key.upper()) or bd.get(key) or ""
                    if val:
                        st.write(f"- **{title} :** {val}")

        # Charger KSA si absent
        if ("ksa_matrix" not in st.session_state or
            not isinstance(st.session_state.ksa_matrix, pd.DataFrame) or
            st.session_state.ksa_matrix.empty):
            js = bd.get("KSA_MATRIX_JSON","")
            if js:
                try:
                    st.session_state.ksa_matrix = pd.DataFrame(json.loads(js))
                except Exception:
                    st.session_state.ksa_matrix = pd.DataFrame()

        if "ksa_matrix" in st.session_state and isinstance(st.session_state.ksa_matrix, pd.DataFrame) and not st.session_state.ksa_matrix.empty:
            st.subheader("📊 Matrice KSA")
            show_df = st.session_state.ksa_matrix.copy()
            st.dataframe(show_df, hide_index=True, width="stretch")
            try:
                if "Évaluation (1-5)" in show_df.columns:
                    avg2 = round(show_df["Évaluation (1-5)"].astype(float).mean(),2)
                    st.markdown(f"<div class='score-cible'>🎯 Score cible : {avg2} / 5</div>", unsafe_allow_html=True)
            except Exception:
                pass
        else:
            st.info("Pas de matrice KSA.")

    # Suppression du titre "Actions & Export"
        col_act, col_exp = st.columns(2)
        with col_act:
            st.markdown("#### Actions")
            a1, a2 = st.columns([1,1])
            with a1:
                if st.button("💾 Sauvegarder synthèse", type="primary", key="btn_save_synthese", width="stretch"):
                    # Construire un dict complet à partir de la session (tous les onglets)
                    try:
                        name = st.session_state.current_brief_name
                        brief_data = {}

                        # Champs mappés (clé sheet -> clé session)
                        for sheet_key, session_key in key_mapping.items():
                            brief_data[sheet_key] = st.session_state.get(session_key, "")

                        # Champs des sections (construits depuis les clés locales)
                        for section in sections:
                            for title, field_key, _ in section["fields"]:
                                header = field_key.upper()
                                unique_key = f"{section['title'].replace(' ', '_')}_{field_key}"
                                brief_data[header] = st.session_state.get(unique_key, st.session_state.get(field_key, ""))

                        # KSA (conserver DataFrame afin que save_brief_to_gsheet puisse le convertir)
                        ksa = st.session_state.get("ksa_matrix")
                        if ksa is not None:
                            brief_data["ksa_matrix"] = ksa

                        # Sauvegarde en session locale
                        st.session_state.saved_briefs[name] = brief_data
                        save_briefs()

                        # Sauvegarde sur Google Sheets (async)
                        try:
                            save_brief_to_gsheet(name, brief_data)
                        except Exception:
                            pass  # Local save succeeded, Sheets is optional

                        # Rafraîchir les briefs chargés
                        try:
                            refresh_saved_briefs()
                        except Exception:
                            pass

                        st.success("Synthèse sauvegardée.")
                    except Exception as e:
                        st.error(f"Erreur lors de la sauvegarde de la synthèse : {e}")
            with a2:
                if st.button("🗑️ Supprimer le brief", key="btn_delete_synthese", width="stretch"):
                    name = st.session_state.current_brief_name
                    st.session_state.saved_briefs.pop(name, None)
                    save_briefs()
                    # Tentative suppression sur Google Sheets + refresh
                    try:
                        from utils import delete_brief_from_gsheet, refresh_saved_briefs
                        delete_brief_from_gsheet(name)
                        refresh_saved_briefs()
                    except Exception:
                        try:
                            refresh_saved_briefs()
                        except Exception:
                            pass
                    st.session_state.current_brief_name = ""
                    st.success("Brief supprimé.")
                    st.rerun()
        with col_exp:
            st.markdown("#### Export")
            e1, e2 = st.columns([1,1])
            pdf_buf = generate_custom_pdf(st.session_state.current_brief_name, bd, st.session_state.ksa_matrix if "ksa_matrix" in st.session_state else None)
            with e1:
                if pdf_buf:
                    st.download_button("⬇️ PDF", data=pdf_buf, file_name=f"{st.session_state.current_brief_name}.pdf", mime="application/pdf", width="stretch")
                else:
                    st.info("PDF indisponible.")
            word_buf = generate_custom_word(st.session_state.current_brief_name, bd, st.session_state.ksa_matrix if "ksa_matrix" in st.session_state else None)
            with e2:
                if word_buf:
                    st.download_button("⬇️ Word", data=word_buf, file_name=f"{st.session_state.current_brief_name}.docx", mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document", width="stretch")
                else:
                    st.info("Word indisponible.")