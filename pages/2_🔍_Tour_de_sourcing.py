import streamlit as st
from utils import save_sourcing_entry_to_gsheet, load_sourcing_entries_from_gsheet

import streamlit as st
import requests
from bs4 import BeautifulSoup
import re
import time
from urllib.parse import quote
from datetime import datetime
import json
import os
import hashlib
import pandas as pd
from collections import Counter

# Fonctions pour gestion des tokens Google Sheets
@st.cache_resource
def get_tokens_gsheet_client():
    """Connexion sp√©cifique √† la feuille Tokens"""
    try:
        import gspread
        from utils import _build_service_account_info_from_st_secrets
        
        service_account_info = _build_service_account_info_from_st_secrets()
        gc = gspread.service_account_from_dict(service_account_info)
        
        # Utiliser la m√™me URL que pour Sourcing DB mais feuille "Tokens"
        SOURCING_SHEET_URL = "https://docs.google.com/spreadsheets/d/1Yw99SS4vU5v0DuD7S1AwaEispJCo-cwioxSsAYnzRkE/edit"
        spreadsheet = gc.open_by_url(SOURCING_SHEET_URL)
        
        # Essayer d'acc√©der √† la feuille Tokens, la cr√©er si elle n'existe pas
        try:
            worksheet = spreadsheet.worksheet("Tokens")
        except gspread.WorksheetNotFound:
            worksheet = spreadsheet.add_worksheet(title="Tokens", rows="1000", cols="10")
            # Ajouter les en-t√™tes
            headers = ["timestamp", "type", "function", "user", "tokens", "cumulative_total", "action"]
            worksheet.append_row(headers)
        
        return worksheet
    except Exception as e:
        st.warning(f"Connexion Tokens Google Sheets indisponible: {e}")
        return None

def save_tokens_to_gsheet(tokens, function_name="General", user="Unknown", reset=False):
    """Sauvegarde les tokens utilis√©s dans Google Sheets avec historique"""
    try:
        worksheet = get_tokens_gsheet_client()
        if worksheet is None:
            return False
        
        if reset:
            # Entr√©e de reset total
            row = [
                datetime.now().strftime("%Y-%m-%d %H:%M:%S"),
                "Token Reset",
                "Reset Total", 
                user,
                0,
                0,
                "RESET TOTAL"
            ]
        else:
            # R√©cup√©rer le total actuel
            current_total = st.session_state.get("api_usage", {}).get("used_tokens", 0)
            
            # Entr√©e d'usage normal
            row = [
                datetime.now().strftime("%Y-%m-%d %H:%M:%S"),
                "Token Usage",
                function_name,
                user,
                tokens,
                current_total,
                "USAGE"
            ]
        
        worksheet.append_row(row)
        return True
    except Exception as e:
        st.warning(f"Erreur sauvegarde tokens: {e}")
        return False

def load_total_tokens_from_gsheet():
    """Charge le total cumul√© des tokens depuis Google Sheets"""
    try:
        worksheet = get_tokens_gsheet_client()
        if worksheet is None:
            return 0
        
        # R√©cup√©rer toutes les entr√©es
        records = worksheet.get_all_records()
        
        if records:
            # Trouver la derni√®re entr√©e pour r√©cup√©rer le total cumul√©
            for record in reversed(records):
                if record.get("action") == "RESET TOTAL":
                    return 0  # Si le dernier √©tait un reset, commencer √† 0
                elif record.get("action") == "USAGE":
                    return int(record.get("cumulative_total", 0))
        
        return 0  # Aucune entr√©e trouv√©e, commencer √† 0
    except Exception as e:
        st.warning(f"Erreur chargement tokens: {e}")
        return 0

# Configuration pour l'appel √† l'IA DeepSeek
SYSTEM_PROMPT = """
Tu es 'TG-Hire Assistant', un expert IA sp√©cialis√© dans le recrutement pour le secteur du BTP (B√¢timent et Travaux Publics) au Maroc.
Ton r√¥le est d'aider un recruteur humain √† optimiser ses t√¢ches quotidiennes.
Tes r√©ponses doivent √™tre :
1.  **Contextualis√©es** : Toujours adapt√©es au march√© de l'emploi marocain et aux sp√©cificit√©s du secteur du BTP.
2.  **Professionnelles et Pr√©cises** : Fournis des informations concr√®tes et structur√©es.
3.  **Orient√©es Action** : Propose des listes, des questions, des mod√®les de texte, etc.
4.  **Adaptables** : Tu dois ajuster la longueur de ta r√©ponse (courte, normale, d√©taill√©e) selon la demande.
"""

def get_deepseek_response(prompt, history, length, function_name="General"):
    api_key = st.secrets.get("DEEPSEEK_API_KEY")
    if not api_key: return {"content": "Erreur: Cl√© API DeepSeek manquante.", "usage": 0}
    
    final_prompt = f"{prompt}\n\n(Instruction: Fournir une r√©ponse de longueur '{length}')"
    messages = [{"role": "system", "content": SYSTEM_PROMPT}] + history + [{"role": "user", "content": final_prompt}]
    
    try:
        response = requests.post(
            "https://api.deepseek.com/v1/chat/completions",
            headers={"Content-Type": "application/json", "Authorization": f"Bearer {api_key}"},
            json={"model": "deepseek-chat", "messages": messages, "max_tokens": 2048}
        )
        response.raise_for_status()
        data = response.json()
        content = data["choices"][0]["message"]["content"]
        usage = data.get("usage", {}).get("total_tokens", 0)
        
        # Mettre √† jour les statistiques de tokens
        if "api_usage" not in st.session_state:
            st.session_state["api_usage"] = {"current_session_tokens": 0, "used_tokens": 0}
        
        st.session_state["api_usage"]["current_session_tokens"] += usage
        st.session_state["api_usage"]["used_tokens"] += usage
        
        # Sauvegarder dans Google Sheets avec d√©tails de la fonction
        user = st.session_state.get("user", "Unknown")
        save_tokens_to_gsheet(usage, function_name, user)
        
        return {"content": content, "usage": usage}
    except Exception as e:
        return {"content": f"‚ùå Erreur API DeepSeek: {e}", "usage": 0}

# Fonction de debug pour l'analyse LinkedIn
def debug_linkedin_analysis(url, ia_result):
    """Debug d√©taill√© pour l'analyse LinkedIn"""
    debug_info = {
        "url_fournie": url,
        "url_valide": bool(url and ("linkedin.com" in url or "linkedin" in url.lower())),
        "reponse_api": ia_result,
        "contenu_present": bool(ia_result.get("content")),
        "longueur_contenu": len(str(ia_result.get("content", ""))) if ia_result.get("content") else 0
    }
    return debug_info

# Import optionnel pour l'extraction PDF
try:
    import PyPDF2
    PDF_AVAILABLE = True
except ImportError:
    try:
        import pdfplumber
        PDF_AVAILABLE = True
    except ImportError:
        PDF_AVAILABLE = False

# Fichier de persistance pour la biblioth√®que
LIB_FILE = "library_entries.json"

# Fonction pour g√©n√©rer des messages InMail
def generate_inmail(donnees_profil, poste_accroche, entreprise, ton_message, longueur_message, cta_option, genre_profil, context):
    """
    G√©n√®re un message InMail personnalis√© bas√© sur les param√®tres fournis.
    """
    # Adaptation du ton
    tone = "formel" if ton_message == "Formel" else "d√©contract√©"
    
    # Adaptation de la longueur
    if longueur_message == "Court":
        max_words = 50
    elif longueur_message == "Moyen":
        max_words = 100
    else:
        max_words = 150
    
    # Construction du message de base
    greeting = f"Bonjour {'Madame' if genre_profil == 'F√©minin' else 'Monsieur'},"
    
    # Corps du message adapt√© au contexte
    if context == "entreprise":
        body = f"""
Je suis impressionn√©(e) par votre profil et votre exp√©rience. Nous avons une opportunit√© {poste_accroche} chez {entreprise} qui pourrait vous int√©resser.

Vos comp√©tences correspondent parfaitement √† ce que nous recherchons. Cette position offre de belles perspectives d'√©volution dans un environnement stimulant.
        """.strip()
    else:
        body = f"Votre profil pour le poste de {poste_accroche} chez {entreprise} a retenu notre attention."
    
    # Call-to-action
    if cta_option == "Appel t√©l√©phonique":
        cta = "Seriez-vous disponible pour un bref appel t√©l√©phonique cette semaine ?"
    elif cta_option == "Rencontre":
        cta = "Seriez-vous int√©ress√©(e) par une rencontre pour discuter de cette opportunit√© ?"
    else:
        cta = "N'h√©sitez pas √† me contacter si cette opportunit√© vous int√©resse."
    
    # Assembly du message final
    message = f"{greeting}\n\n{body}\n\n{cta}\n\nCordialement"
    
    # Limitation de la longueur si n√©cessaire
    words = message.split()
    if len(words) > max_words:
        message = " ".join(words[:max_words]) + "..."
    
    return message

# --- CSS pour augmenter la taille du texte des onglets ---
st.markdown("""
<style>
div[data-testid="stTabs"] button p {
    font-size: 18px; 
}
</style>
<script>
function copyToClipboard(text){
    navigator.clipboard.writeText(text).then(()=>{
        const ev = new Event('clipboard-copied');
        document.dispatchEvent(ev);
    });
}
document.addEventListener('click', function(e){
    if(e.target && e.target.dataset && e.target.dataset.copy){
         copyToClipboard(e.target.dataset.copy);
    }
});
</script>
""", unsafe_allow_html=True)

# -------------------- Configuration initiale --------------------
def _load_library_entries():
    if os.path.exists(LIB_FILE):
        try:
            with open(LIB_FILE, 'r', encoding='utf-8') as f:
                data = json.load(f)
            if isinstance(data, list):
                return data
        except Exception:
            return []
    return []

def init_session_state():
    """Initialise les variables de session"""
    # Charger le total cumul√© depuis Google Sheets au premier chargement
    if "tokens_loaded" not in st.session_state:
        total_from_sheets = load_total_tokens_from_gsheet()
        st.session_state["tokens_loaded"] = True
        api_usage = {"current_session_tokens": 0, "used_tokens": total_from_sheets}
    else:
        api_usage = st.session_state.get("api_usage", {"current_session_tokens": 0, "used_tokens": 0})
    
    defaults = {
        "api_usage": api_usage,
        "library_entries": _load_library_entries(),
        "magicien_history": [],
        "boolean_query": "",
        "boolean_snapshot": {},
        "xray_query": "",
        "xray_snapshot": {},
        "cse_query": "",
        "dogpile_query": "",
        "scraper_result": "",
        "scraper_emails": set(),
        "inmail_message": "",
        "perm_result": [],
        "inmail_objet": "",
        "inmail_generated": False,
        "inmail_profil_data": {},
    }
    for k, v in defaults.items():
        if k not in st.session_state:
            st.session_state[k] = v

def save_library_entries():
    """Sauvegarde les entr√©es de la biblioth√®que en JSON"""
    try:
        with open(LIB_FILE, 'w', encoding='utf-8') as f:
            json.dump(st.session_state.library_entries, f, ensure_ascii=False, indent=2)
    except Exception as e:
        st.warning(f"‚ö†Ô∏è √âchec de sauvegarde biblioth√®que: {e}")

def _split_terms(raw: str) -> list:
    if not raw:
        return []
    # support virgule / point-virgule / saut de ligne
    separators = [',', ';', '\n']
    for sep in separators:
        raw = raw.replace(sep, '|')
    terms = [t.strip() for t in raw.split('|') if t.strip()]
    # d√©duplication en conservant l'ordre
    seen = set(); ordered = []
    for t in terms:
        low = t.lower()
        if low not in seen:
            seen.add(low); ordered.append(t)
    return ordered

def _or_group(terms: list[str]) -> str:
    if not terms:
        return ''
    if len(terms) == 1:
        return f'"{terms[0]}"'
    return '(' + ' OR '.join(f'"{t}"' for t in terms) + ')'

def _and_group(terms: list[str]) -> str:
    if not terms:
        return ''
    if len(terms) == 1:
        return f'"{terms[0]}"'
    return '(' + ' AND '.join(f'"{t}"' for t in terms) + ')'

def generate_boolean_query(poste: str, synonymes: str, competences_obligatoires: str,
                           competences_optionnelles: str, exclusions: str, localisation: str, secteur: str,
                           employeur: str | None = None) -> str:
    """G√©n√®re une requ√™te boolean normalis√©e.
    - Support multi-termes (virgule / ; / retour ligne)
    - D√©duplication
    - Groupes OR / AND corrects
    """
    parts: list[str] = []
    if poste or synonymes:
        # Combine poste and synonyms with OR
        poste_and_synonyms = _or_group([poste] + _split_terms(synonymes) if poste else _split_terms(synonymes))
        parts.append(poste_and_synonyms)
    comp_ob = _split_terms(competences_obligatoires)
    if comp_ob:
        parts.append(_and_group(comp_ob))
    comp_opt = _split_terms(competences_optionnelles)
    if comp_opt:
        parts.append(_or_group(comp_opt))
    if localisation:
        parts.append(f'"{localisation}"')
    if secteur:
        parts.append(f'"{secteur}"')
    # exclusions ‚Üí NOT group OR
    excl = _split_terms(exclusions)
    if excl:
        parts.append('NOT ' + _or_group(excl))
    if employeur:
        parts.append(f'("{employeur}")')
    return ' AND '.join(filter(None, parts))

def generate_boolean_variants(base_query: str, synonymes: str, comp_opt: str) -> list:
    """G√©n√®re quelques variantes simples:
    - Variante 1: sans comp√©tences optionnelles
    - Variante 2: synonymes en fin
    - Variante 3: suppression des guillemets sur poste/synonymes (si applicable)
    """
    variants = []
    try:
        if not base_query:
            return []
        # Variante 1: retirer groupe optionnel si pr√©sent
        if comp_opt:
            opt_terms = _split_terms(comp_opt)
            if opt_terms:
                opt_group = _or_group(opt_terms)
                v1 = base_query.replace(f" AND {opt_group}", "")
                variants.append(("Sans comp√©tences optionnelles", v1))
        # Variante 2: d√©placer synonymes √† la fin
        if synonymes:
            syn_terms = _split_terms(synonymes)
            syn_group = _or_group(syn_terms)
            if syn_group in base_query:
                parts = base_query.split(" AND ")
                reordered = [p for p in parts if p != syn_group] + [syn_group]
                variants.append(("Synonymes en fin", ' AND '.join(reordered)))
        # Variante 3: retirer guillemets des synonymes individuels si pas d'espaces
        if synonymes:
            syn_terms = _split_terms(synonymes)
            if syn_terms and all(' ' not in t for t in syn_terms):
                syn_group = _or_group(syn_terms)
                no_quotes = '(' + ' OR '.join(syn_terms) + ')'
                variants.append(("Synonymes sans guillemets", base_query.replace(syn_group, no_quotes)))
    except Exception:
        pass
    # d√©duplication titres
    seen = set(); final=[]
    for title, q in variants:
        if q not in seen:
            seen.add(q); final.append((title, q))
    return final[:3]

def generate_xray_query(site_cible: str, poste: str, mots_cles: str, localisation: str) -> str:
    """G√©n√®re une requ√™te X-Ray am√©lior√©e.
    - Support multi mots-cl√©s / localisations
    - Groupes OR pour √©largir la recherche
    """
    site_map = {"LinkedIn": "site:linkedin.com/in", "GitHub": "site:github.com"}
    site = site_map.get(site_cible, "site:linkedin.com/in")
    parts = [site]
    if poste:
        parts.append(f'"{poste}"')
    kws = _split_terms(mots_cles)
    if kws:
        parts.append(_or_group(kws))
    locs = _split_terms(localisation)
    if locs:
        parts.append(_or_group(locs))
    return ' '.join(parts)

def generate_xray_variants(query: str, poste: str, mots_cles: str, localisation: str) -> list:
    variants = []
    try:
        if not query:
            return []
        # intitle sur poste
        if poste:
            v1 = query.replace(f'"{poste}"', f'intitle:"{poste}"') if f'"{poste}"' in query else query + f' intitle:"{poste}"'
            variants.append(("intitle: poste", v1))
        # S√©parer mots-cl√©s en OR explicite si plusieurs
        kws = _split_terms(mots_cles)
        if kws and len(kws) > 1:
            or_block = '(' + ' OR '.join(f'"{k}"' for k in kws) + ')'
            base_no = re.sub(r'\([^)]*\)', '', query)  # tentative retrait ancien groupe
            variants.append(("OR explicite mots-cl√©s", f"{base_no} {or_block}".strip()))
        # Localisations en OR avec pattern "(Casablanca OR Rabat)"
        locs = _split_terms(localisation)
        if locs and len(locs) > 1:
            loc_block = '(' + ' OR '.join(f'"{l}"' for l in locs) + ')'
            if any(l in query for l in locs):
                variants.append(("Localisations OR", query + ' ' + loc_block))
    except Exception:
        pass
    # d√©dup
    seen=set(); final=[]
    for t,q in variants:
        if q not in seen:
            seen.add(q); final.append((t,q))
    return final[:3]

def build_xray_linkedin(poste: str, mots_cles: list[str], localisations: list[str],
                        langues: list[str], entreprises: list[str], ecoles: list[str],
                        seniority: str | None) -> str:
    """Construit une requ√™te X-Ray LinkedIn plus riche.
    seniority peut √™tre: 'junior','senior','manager'
    """
    parts = ["site:linkedin.com/in"]
    if poste:
        parts.append(f'("{poste}" OR intitle:"{poste}")')
    if mots_cles:
        parts.append('(' + ' OR '.join(f'"{m}"' for m in mots_cles) + ')')
    if localisations:
        parts.append('(' + ' OR '.join(f'"{l}"' for l in localisations) + ')')
    if langues:
        # tente de cibler la langue via mots fr√©quents
        for lg in langues:
            if lg.lower().startswith('fr'):
                parts.append('("Fran√ßais" OR "French")')
            elif lg.lower().startswith('en'):
                parts.append('("Anglais" OR "English")')
            elif lg.lower().startswith('ar'):
                parts.append('("Arabe" OR "Arabic")')
    if entreprises:
        parts.append('("' + '" OR "'.join(entreprises) + '")')
    if ecoles:
        parts.append('("' + '" OR "'.join(ecoles) + '")')
    return ' '.join(parts)

def generate_accroche_inmail(url_linkedin, poste_accroche):
    """G√©n√®re un message InMail basique"""
    return f"""Bonjour,

Votre profil sur LinkedIn a retenu mon attention, particuli√®rement votre exp√©rience dans le domaine.

Je me permets de vous contacter concernant une opportunit√© de {poste_accroche} qui correspond parfaitement √† votre profil. Votre expertise serait un atout pr√©cieux pour notre √©quipe.

Seriez-vous ouvert √† un √©change pour discuter de cette opportunit√© ?

Dans l'attente de votre retour,"""

def ask_deepseek(messages, max_tokens=300):
    """Simule l'appel √† l'API DeepSeek avec une logique d'enrichissement."""
    time.sleep(1)  # Simulation de d√©lai
    question = messages[0]["content"].lower()
    
    # 1. Extraction des crit√®res de base
    def extract_field(field_name, content):
        match = re.search(f"{field_name}:\\s*(.*?)(?:\\n|$)", content, re.IGNORECASE)
        return match.group(1).strip() if match else ""
    
    poste = extract_field("poste", messages[0]["content"])
    synonymes = extract_field("synonymes", messages[0]["content"])
    comp_ob = extract_field("comp√©tences obligatoires", messages[0]["content"])
    
    # 2. Logique pour simuler l'enrichissement par l'IA
    
    # Cas 1 : G√©n√©ration de la requ√™te Boolean (enrichissement des synonymes et des comp√©tences)
    if "g√©n√®re une requ√™te boolean" in question:
        
        # Approche MINIMALISTE pour √©viter 0 r√©sultat sur LinkedIn
        # Seulement enrichir les synonymes, pas trop de comp√©tences obligatoires
        
        if "ing√©nieur de travaux" in poste.lower() or "ing√©nieur travaux" in poste.lower():
            # Synonymes conservateurs
            ia_syns = f"{synonymes}, Conducteur de travaux, Chef de chantier" if synonymes else "Conducteur de travaux, Chef de chantier"
            # Comp√©tences tr√®s l√©g√®res ou vides si rien n'est saisi
            ia_comp_ob = comp_ob if comp_ob else ""  # Ne pas ajouter de comp√©tences obligatoires si vide
        elif "ged" in poste.lower() or "gestion √©lectronique" in poste.lower() or "archivage" in poste.lower():
            ia_syns = f"{synonymes}, Gestionnaire documentaire, Archiviste, Document manager" if synonymes else "Gestionnaire documentaire, Archiviste, Document manager"
            ia_comp_ob = comp_ob if comp_ob else ""
        elif "charg√© de recrutement" in poste.lower() or "recruteur" in poste.lower():
            ia_syns = f"{synonymes}, Talent acquisition, Sourcing" if synonymes else "Talent acquisition, Sourcing"
            ia_comp_ob = comp_ob if comp_ob else ""
        elif "d√©veloppeur" in poste.lower() or "developer" in poste.lower():
            ia_syns = f"{synonymes}, Software engineer, Programmeur" if synonymes else "Software engineer, Programmeur"
            ia_comp_ob = comp_ob if comp_ob else ""
        elif "comptable" in poste.lower() or "finance" in poste.lower():
            ia_syns = f"{synonymes}, Expert comptable, Contr√¥leur gestion" if synonymes else "Expert comptable, Contr√¥leur gestion"
            ia_comp_ob = comp_ob if comp_ob else ""
        elif "responsable" in poste.lower() or "manager" in poste.lower() or "directeur" in poste.lower():
            # Synonymes pour postes de management/responsabilit√©
            ia_syns = f"{synonymes}, Manager, Chef de service, Directeur" if synonymes else "Manager, Chef de service, Directeur"
            ia_comp_ob = comp_ob if comp_ob else ""
        else:
            # Pour les postes non reconnus, ajouter seulement des synonymes g√©n√©riques
            ia_syns = f"{synonymes}, Senior, Expert" if synonymes else "Senior, Expert"
            ia_comp_ob = comp_ob  # Garder ce que l'utilisateur a saisi
        
        return {"content": ia_syns, "comp_ob_ia": ia_comp_ob}
    
    # Cas 2 : Analyse de fiche de poste
    elif "analyse cette fiche de poste" in question:
        # Extraire seulement le contenu de la fiche de poste (apr√®s les deux points)
        full_content = messages[0]["content"]
        
        # Trouver le d√©but r√©el de la fiche de poste (apr√®s "cl√©s:" ou similaire)
        if ":" in full_content:
            # Chercher apr√®s le dernier ":" qui marque la fin du prompt
            parts = full_content.split('\n')
            fiche_lines = []
            found_content = False
            
            for line in parts:
                if found_content or (not line.startswith('analyse') and not line.startswith('extrait') and ':' not in line):
                    if line.strip():  # Ignorer les lignes vides au d√©but
                        found_content = True
                        fiche_lines.append(line)
                elif line.strip().endswith(':') or '5. Mots √† exclure' in line:
                    found_content = True  # Commencer √† capturer apr√®s cette ligne
            
            fiche_content = '\n'.join(fiche_lines)
        else:
            fiche_content = full_content
        
        fiche_lower = fiche_content.lower()
        
        # Extraire le titre du poste (premi√®re ligne souvent)
        lines = fiche_content.strip().split('\n')
        titre_candidat = ""
        
        # Chercher le titre dans les premi√®res lignes
        for line in lines[:5]:  # Regarder les 5 premi√®res lignes
            line = line.strip()
            if line and not line.startswith('opportunit√©') and not line.startswith('rejoignez') and not line.startswith('tgcc recrute'):
                if len(line) < 100 and not line.lower().startswith('missions') and not line.lower().startswith('analyse'):  # Probablement un titre
                    titre_candidat = line
                    break
        
        suggestions = []
        
        # Cas sp√©cifiques bas√©s sur le contenu
        if "ged" in fiche_lower or "gestion √©lectronique" in fiche_lower or "archivage" in fiche_lower:
            suggestions.append(f"Titre: {titre_candidat if titre_candidat else 'Responsable GED & Archivage'}")
            suggestions.append("Synonymes: Gestionnaire documentaire, Archiviste, Document manager")
            suggestions.append("Comp√©tences obligatoires: GED, Archivage, D√©mat√©rialisation")
            suggestions.append("Comp√©tences optionnelles: Gouvernance documentaire, Normes ISO, M√©tadonn√©es")
        elif "ing√©nieur" in fiche_lower and "travaux" in fiche_lower:
            suggestions.append(f"Titre: {titre_candidat if titre_candidat else 'Ing√©nieur de travaux'}")
            suggestions.append("Synonymes: Conducteur de travaux, Chef de chantier")
            suggestions.append("Comp√©tences obligatoires: AutoCAD, Gestion projet")
            suggestions.append("Comp√©tences optionnelles: Primavera, Management √©quipe")
        elif "d√©veloppeur" in fiche_lower or "developer" in fiche_lower:
            suggestions.append(f"Titre: {titre_candidat if titre_candidat else 'D√©veloppeur'}")
            suggestions.append("Synonymes: Software engineer, Programmeur")
            suggestions.append("Comp√©tences obligatoires: Programming, Git")
            suggestions.append("Comp√©tences optionnelles: Framework, Base de donn√©es")
        elif "comptable" in fiche_lower or "finance" in fiche_lower:
            suggestions.append(f"Titre: {titre_candidat if titre_candidat else 'Comptable'}")
            suggestions.append("Synonymes: Expert comptable, Contr√¥leur gestion")
            suggestions.append("Comp√©tences obligatoires: SAGE, Fiscalit√©")
            suggestions.append("Comp√©tences optionnelles: Audit, Consolidation")
        elif "responsable" in fiche_lower or "manager" in fiche_lower or "directeur" in fiche_lower:
            suggestions.append(f"Titre: {titre_candidat if titre_candidat else 'Responsable'}")
            suggestions.append("Synonymes: Manager, Chef de service, Directeur")
            
            # Extraire les comp√©tences mentionn√©es dans la fiche
            comp_obligatoires = []
            comp_optionnelles = []
            
            if "leadership" in fiche_lower:
                comp_obligatoires.append("Leadership")
            if "management" in fiche_lower or "encadrer" in fiche_lower:
                comp_obligatoires.append("Management")
            if "gestion" in fiche_lower:
                comp_obligatoires.append("Gestion")
            if "pilotage" in fiche_lower or "piloter" in fiche_lower:
                comp_optionnelles.append("Pilotage projet")
            if "reporting" in fiche_lower:
                comp_optionnelles.append("Reporting")
            if "analyse" in fiche_lower:
                comp_optionnelles.append("Esprit d'analyse")
                
            suggestions.append(f"Comp√©tences obligatoires: {', '.join(comp_obligatoires) if comp_obligatoires else 'Management, Leadership'}")
            suggestions.append(f"Comp√©tences optionnelles: {', '.join(comp_optionnelles) if comp_optionnelles else 'Pilotage projet, Communication'}")
        elif "directeur" in fiche_lower and "capital humain" in fiche_lower:
            suggestions.append(f"Titre: {titre_candidat if titre_candidat else 'Directeur du Capital Humain'}")
            suggestions.append("Synonymes: Chief Human Resources Officer, DRH (Directeur des Ressources Humaines), DRH (Directeur Capital Humain)")
            suggestions.append("Comp√©tences obligatoires: Paie & Administration du personnel, Recrutement, D√©veloppement des talents")
            suggestions.append("Comp√©tences optionnelles: Pilotage projet, Esprit d'analyse")
        else:
            # Essayer d'extraire quand m√™me des informations g√©n√©riques
            if titre_candidat:
                suggestions.append(f"Titre: {titre_candidat}")
                suggestions.append("Synonymes: Expert, Sp√©cialiste, Senior")
            else:
                suggestions.append("Analyse: Titre non d√©tect√© clairement")
            
            # Rechercher des comp√©tences dans le texte
            competences_trouvees = []
            mots_competences = ["gestion", "management", "leadership", "pilotage", "analyse", "organisation", 
                              "supervision", "coordination", "planification", "suivi", "contr√¥le"]
            for mot in mots_competences:
                if mot in fiche_lower:
                    competences_trouvees.append(mot.capitalize())
            
            if competences_trouvees:
                suggestions.append(f"Comp√©tences d√©tect√©es: {', '.join(competences_trouvees[:3])}")
            else:
                suggestions.append("Conseil: Remplissez manuellement les champs ci-dessous")
            
        return {"content": "\n".join(suggestions)}
    
    # Cas 3 : Outils/Logiciels
    elif "outils" in question or "logiciels" in question:
        return {"content": "‚Ä¢ AutoCAD\n‚Ä¢ Revit\n‚Ä¢ Primavera P6\n‚Ä¢ MS Project\n‚Ä¢ Robot Structural Analysis\n‚Ä¢ SketchUp"}
        
    # Cas 4 : Comp√©tences
    elif "comp√©tences" in question:
        return {"content": "‚Ä¢ Gestion de projet\n‚Ä¢ Lecture de plans techniques\n‚Ä¢ Management d'√©quipe\n‚Ä¢ Budget et planning\n‚Ä¢ Conformit√© r√©glementaire\n‚Ä¢ N√©gociation fournisseurs"}
        
    # Cas par d√©faut : G√©n√®re une r√©ponse intelligente selon le contexte
    if "synonymes" in question:
        # Analyser le m√©tier sp√©cifique pour donner des synonymes pertinents
        if "charg√© de recrutement" in question or "charg√©e de recrutement" in question:
            return {"content": "Recruteur, Talent Acquisition Specialist, Responsable Recrutement, Consultant en Recrutement, RH Recrutement, Chasseur de t√™tes, Sourcing Specialist"}
        elif "d√©veloppeur" in question or "programmeur" in question:
            return {"content": "D√©veloppeur, Programmeur, Ing√©nieur logiciel, Software Engineer, Codeur, Dev, Analyste programmeur"}
        elif "commercial" in question or "vente" in question:
            return {"content": "Commercial, Vendeur, Conseiller commercial, Business Developer, Account Manager, Sales Representative, Ing√©nieur commercial"}
        elif "comptable" in question or "finance" in question:
            return {"content": "Comptable, Contr√¥leur de gestion, Analyste financier, Gestionnaire comptable, Assistant comptable, Expert-comptable"}
        elif "ing√©nieur" in question:
            return {"content": "Ing√©nieur, Engineer, Technicien sup√©rieur, Ing√©nieur d'√©tudes, Ing√©nieur de conception, Consultant technique"}
        elif "manager" in question or "responsable" in question:
            return {"content": "Manager, Responsable, Chef d'√©quipe, Superviseur, Directeur, Coordinateur, Team Leader"}
        else:
            # Synonymes g√©n√©riques si m√©tier non reconnu
            return {"content": "Responsable, Sp√©cialiste, Consultant, Expert, Coordinateur, Assistant, Charg√© de mission"}
    elif any(word in question for word in ["junior", "d√©butant", "junior"]):
        return {"content": "‚Ä¢ Junior\n‚Ä¢ D√©butant\n‚Ä¢ Assistant\n‚Ä¢ Stagiaire\n‚Ä¢ Alternant\n‚Ä¢ En formation\n‚Ä¢ Premier emploi\n‚Ä¢ Entry level"}
    elif "outils" in question or "logiciels" in question:
        if "recrutement" in question:
            return {"content": "‚Ä¢ LinkedIn Recruiter\n‚Ä¢ Indeed\n‚Ä¢ Workday\n‚Ä¢ BambooHR\n‚Ä¢ Greenhouse\n‚Ä¢ Lever\n‚Ä¢ SmartRecruiters\n‚Ä¢ Taleo"}
        else:
            return {"content": "‚Ä¢ MS Office\n‚Ä¢ Google Workspace\n‚Ä¢ CRM\n‚Ä¢ ERP\n‚Ä¢ Slack\n‚Ä¢ Teams\n‚Ä¢ Zoom\n‚Ä¢ Project management tools"}
    elif "secteur" in question:
        return {"content": "‚Ä¢ BTP\n‚Ä¢ Construction\n‚Ä¢ Technologie\n‚Ä¢ Finance\n‚Ä¢ Sant√©\n‚Ä¢ Industrie\n‚Ä¢ Services\n‚Ä¢ Consulting\n‚Ä¢ E-commerce"}
    elif "certification" in question:
        if "recrutement" in question or "rh" in question:
            return {"content": "‚Ä¢ CPRP (Certified Professional in Recruitment)\n‚Ä¢ PHR (Professional in Human Resources)\n‚Ä¢ SHRM-CP\n‚Ä¢ CIPD\n‚Ä¢ Certification LinkedIn Recruiter"}
        else:
            return {"content": "‚Ä¢ PMP\n‚Ä¢ ISO 27001\n‚Ä¢ ITIL\n‚Ä¢ Agile/Scrum\n‚Ä¢ Six Sigma\n‚Ä¢ PRINCE2"}
    else:
        # R√©ponse g√©n√©rale pour toute autre question
        return {"content": "Voici quelques suggestions pertinentes pour votre recherche de sourcing :\n‚Ä¢ Variez les mots-cl√©s selon le m√©tier cibl√©\n‚Ä¢ Utilisez des synonymes sp√©cifiques au domaine\n‚Ä¢ Pensez aux comp√©tences transversales\n‚Ä¢ Consid√©rez le niveau d'exp√©rience requis\n‚Ä¢ Adaptez selon le secteur d'activit√©\n‚Ä¢ Incluez les outils m√©tiers sp√©cialis√©s"}

def extract_text_from_pdf(uploaded_file):
    """Extrait le texte d'un fichier PDF upload√©"""
    try:
        if not PDF_AVAILABLE:
            return "Extraction PDF non disponible - Librairies PyPDF2 ou pdfplumber manquantes"

        # Lire les octets et travailler sur un BytesIO (plus fiable avec les librairies)
        try:
            import io
            uploaded_file.seek(0)
            data = uploaded_file.read()
            bio = io.BytesIO(data)
        except Exception:
            # Si on ne peut pas lire le buffer, tenter d'utiliser le fichier tel quel
            bio = uploaded_file

        # 1) pdfplumber
        try:
            import pdfplumber
            bio.seek(0)
            with pdfplumber.open(bio) as pdf:
                text_parts = []
                for page in pdf.pages:
                    page_text = page.extract_text()
                    if page_text:
                        text_parts.append(page_text)
                text = "\n".join(text_parts).strip()
                if text:
                    return text
        except Exception as e:
            print(f"Erreur pdfplumber: {e}")

        # 2) PyPDF2
        try:
            import PyPDF2
            bio.seek(0)
            reader = PyPDF2.PdfReader(bio)
            text_parts = []
            for page in reader.pages:
                try:
                    page_text = page.extract_text()
                except Exception:
                    page_text = None
                if page_text:
                    text_parts.append(page_text)
            text = "\n".join(text_parts).strip()
            if text:
                return text
        except Exception as e:
            print(f"Erreur PyPDF2: {e}")

        # 3) pypdf / pypdf.PdfReader (au cas o√π)
        try:
            from pypdf import PdfReader as PypdfReader
            bio.seek(0)
            reader = PypdfReader(bio)
            text_parts = []
            for page in reader.pages:
                try:
                    page_text = page.extract_text()
                except Exception:
                    page_text = None
                if page_text:
                    text_parts.append(page_text)
            text = "\n".join(text_parts).strip()
            if text:
                return text
        except Exception as e:
            print(f"Erreur pypdf: {e}")

        # 4) Fallback OCR si install√© (pdf2image + pytesseract)
        try:
            # Importer dynamiquement pour √©viter que Pylance signale une erreur
            import importlib
            convert_mod = importlib.import_module('pdf2image')
            pytesseract = importlib.import_module('pytesseract')
            convert_from_bytes = getattr(convert_mod, 'convert_from_bytes')
            bio.seek(0)
            images = convert_from_bytes(bio.read(), dpi=200)
            ocr_text_parts = []
            for img in images:
                try:
                    # petite r√©duction si tr√®s grande
                    ocr_text = pytesseract.image_to_string(img, lang='fra+eng')
                    if ocr_text:
                        ocr_text_parts.append(ocr_text)
                except Exception:
                    continue
            ocr_text = "\n".join(ocr_text_parts).strip()
            if ocr_text:
                return ocr_text
        except Exception as e:
            # Si pdf2image/pytesseract non install√©s ou erreur OCR, ignorer
            print(f"OCR fallback unavailable or failed: {e}")

        # Si on arrive ici, aucune librairie n'a extrait du texte
        return "Aucun texte lisible trouv√©. Le PDF est peut-√™tre un scan (images) ou prot√©g√© par mot de passe."

    except Exception as e:
        return f"Erreur lors de l'extraction PDF: {str(e)}"

def get_email_from_charika(entreprise):
    """Recherche d'email d'entreprise depuis Charika.ma avec am√©lioration"""
    try:
        # Rechercher sur Charika.ma
        search_url = f"https://www.charika.ma/search?q={quote(entreprise)}"
        headers = {
            'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/91.0.4472.124 Safari/537.36'
        }
        
        response = requests.get(search_url, headers=headers, timeout=10)
        soup = BeautifulSoup(response.content, 'html.parser')
        
        # Chercher le lien vers la page de l'entreprise - approche am√©lior√©e
        company_links = soup.find_all('a', href=True)
        company_url = None
        
        # Approche am√©lior√©e pour trouver les liens d'entreprise
        # 1. Chercher d'abord les liens avec "entreprise" ou "fiche" dans l'URL
        for link in company_links:
            href = link.get('href', '')
            text = link.get_text().strip().lower()
            if ('entreprise' in href or 'fiche' in href or 'company' in href) and any(word in text for word in entreprise.lower().split()):
                company_url = "https://www.charika.ma" + href if href.startswith('/') else href
                break
        
        # 2. Si pas trouv√©, chercher plus largement
        if not company_url:
            for link in company_links:
                href = link.get('href', '')
                text = link.get_text().strip().lower()
                # Recherche plus souple dans le texte du lien
                if href and href.startswith('/') and len(href) > 5:  # URLs relatives significatives
                    # V√©rifier si le nom de l'entreprise est dans le texte
                    if (entreprise.lower() in text or 
                        any(word.lower() in text for word in entreprise.split() if len(word) > 2)):
                        company_url = "https://www.charika.ma" + href
                        break
        
        # 3. Essayer des URLs construites directement (fallback)
        if not company_url:
            # Essayer diff√©rents formats d'URL possibles
            possible_urls = [
                f"https://www.charika.ma/entreprise/{entreprise.lower().replace(' ', '-')}",
                f"https://www.charika.ma/fiche/{entreprise.lower().replace(' ', '-')}",
                f"https://www.charika.ma/company/{entreprise.lower().replace(' ', '-')}",
                f"https://www.charika.ma/entreprise/{entreprise.lower().replace(' ', '').replace('-', '')}",
            ]
            
            for test_url in possible_urls:
                try:
                    test_response = requests.get(test_url, headers=headers, timeout=10)
                    # V√©rifier si la page existe et n'est pas une page d'erreur
                    if test_response.status_code == 200 and 'Page manquante' not in test_response.text:
                        company_url = test_url
                        break
                except:
                    continue
        
        if company_url:
            # Acc√©der √† la page de l'entreprise
            company_response = requests.get(company_url, headers=headers, timeout=10)
            company_soup = BeautifulSoup(company_response.content, 'html.parser')
            
            # M√©thode am√©lior√©e bas√©e sur l'inspection de la structure HTML de Charika.ma
            # Structure identifi√©e: <span class="dropdown"> avec <span class="mrg-fiche3"> contenant "E-mail" et lien mailto
            
            # M√©thode 1: Chercher sp√©cifiquement la structure HTML observ√©e
            # Pattern: <span class="dropdown"> contenant "E-mail" et un lien mailto
            dropdown_spans = company_soup.find_all('span', class_='dropdown')
            for dropdown in dropdown_spans:
                dropdown_html = str(dropdown)
                # V√©rifier si ce dropdown contient "E-mail" ET un lien mailto
                if 'E-mail' in dropdown_html and 'mailto:' in dropdown_html:
                    # Extraire directement l'email du HTML avec regex (plus fiable)
                    mailto_pattern = r'mailto:([A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Z|a-z]{2,})'
                    matches = re.findall(mailto_pattern, dropdown_html)
                    for email in matches:
                        if 'charika.ma' not in email.lower():
                            return email
                
                # M√©thode alternative avec BeautifulSoup
                mrg_spans = dropdown.find_all('span', class_='mrg-fiche3')
                for mrg_span in mrg_spans:
                    if 'E-mail' in mrg_span.get_text():
                        # Chercher les liens mailto dans ce dropdown
                        mailto_links = dropdown.find_all('a', href=lambda x: x and x.startswith('mailto:'))
                        for link in mailto_links:
                            email = link.get('href').replace('mailto:', '').strip()
                            if '@' in email and '.' in email.split('@')[1]:
                                # V√©rifier que ce n'est pas l'email de Charika
                                if 'charika.ma' not in email.lower():
                                    return email
            
            # M√©thode 2: Chercher tous les liens mailto dans la page (en excluant Charika)
            mailto_links = company_soup.find_all('a', href=lambda x: x and x.startswith('mailto:'))
            for link in mailto_links:
                email = link.get('href').replace('mailto:', '').strip()
                if '@' in email and '.' in email.split('@')[1]:
                    # Filtrer les emails g√©n√©riques et celui de Charika
                    excluded = ['noreply', 'no-reply', 'donotreply', 'example', 'test', 'charika.ma', 'contact@charika']
                    if not any(generic in email.lower() for generic in excluded):
                        return email
            
            # M√©thode 3: Chercher dans les √©l√©ments contenant "E-mail" avec pattern sp√©cifique
            # Chercher <span class="mrg-fiche3"> contenant "E-mail"
            email_elements = company_soup.find_all('span', class_='mrg-fiche3')
            for element in email_elements:
                if 'E-mail' in element.get_text():
                    # Chercher dans le parent (dropdown) ou les √©l√©ments suivants
                    parent = element.parent
                    if parent:
                        # Chercher les liens mailto dans le parent
                        mailto_links = parent.find_all('a', href=lambda x: x and x.startswith('mailto:'))
                        for link in mailto_links:
                            email = link.get('href').replace('mailto:', '').strip()
                            if '@' in email and '.' in email.split('@')[1] and 'charika.ma' not in email.lower():
                                return email
                    
                    # Chercher aussi dans les √©l√©ments suivants (siblings)
                    for sibling in element.next_siblings:
                        if hasattr(sibling, 'find_all'):
                            mailto_links = sibling.find_all('a', href=lambda x: x and x.startswith('mailto:'))
                            for link in mailto_links:
                                email = link.get('href').replace('mailto:', '').strip()
                                if '@' in email and '.' in email.split('@')[1] and 'charika.ma' not in email.lower():
                                    return email
            
            # M√©thode 4: Fallback - regex dans tout le texte (en excluant Charika)
            email_pattern = r'\b[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Z|a-z]{2,}\b'
            all_text = company_soup.get_text()
            emails = re.findall(email_pattern, all_text)
            if emails:
                # Filtrer les emails g√©n√©riques et celui de Charika
                excluded = ['noreply', 'no-reply', 'donotreply', 'example', 'test', 'charika.ma']
                for email in emails:
                    if not any(generic in email.lower() for generic in excluded):
                        return email
        
        # Retourner None si pas trouv√© pour permettre la gestion d'erreur
        return None
        
    except Exception as e:
        print(f"Erreur lors de la recherche sur Charika: {e}")
        return None

def get_charika_search_url(entreprise):
    """G√©n√®re l'URL de recherche Google pour trouver l'entreprise sur Charika"""
    search_query = f"{entreprise} charika.ma"
    return f"https://www.google.com/search?q={quote(search_query)}"

def search_email_on_website(entreprise):
    """Recherche d'email sur le site officiel de l'entreprise"""
    try:
        # Essayer de trouver le site web de l'entreprise
        search_query = f"{entreprise} site:ma OR site:com contact"
        search_url = f"https://www.google.com/search?q={quote(search_query)}"
        
        headers = {
            'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/91.0.4472.124 Safari/537.36'
        }
        
        response = requests.get(search_url, headers=headers, timeout=10)
        soup = BeautifulSoup(response.content, 'html.parser')
        
        # Extraire les URLs des r√©sultats de recherche
        result_links = soup.find_all('a', href=True)
        potential_websites = []
        
        for link in result_links:
            href = link.get('href', '')
            if href.startswith('/url?q='):
                actual_url = href.split('/url?q=')[1].split('&')[0]
                if any(domain in actual_url for domain in ['.ma', '.com']) and 'google' not in actual_url:
                    potential_websites.append(actual_url)
        
        # Visiter les sites potentiels pour chercher des emails
        email_pattern = r'\b[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Z|a-z]{2,}\b'
        
        for website in potential_websites[:3]:  # Limiter √† 3 sites
            try:
                site_response = requests.get(website, headers=headers, timeout=8)
                site_content = site_response.text
                
                # Chercher sp√©cifiquement dans les pages contact
                if 'contact' in site_content.lower():
                    emails = re.findall(email_pattern, site_content)
                    # Filtrer les emails g√©n√©riques
                    for email in emails:
                        if not any(generic in email.lower() for generic in ['noreply', 'no-reply', 'donotreply']):
                            return email
                            
            except Exception:
                continue
        
        # Format par d√©faut si rien trouv√©
        domain = f"{entreprise.lower().replace(' ', '').replace('-', '')}.ma"
        return f"contact@{domain}"
        
    except Exception as e:
        print(f"Erreur lors de la recherche d'email: {e}")
        # Format par d√©faut
        domain = f"{entreprise.lower().replace(' ', '').replace('-', '')}.ma"
        return f"contact@{domain}"

# -------------------- Initialisation --------------------
init_session_state()
st.set_page_config(
    page_title="TG-Hire IA - Assistant Recrutement",
    page_icon="ü§ñ",
    layout="wide",
    initial_sidebar_state="expanded"
)

# -------------------- Sidebar --------------------
with st.sidebar:
    # Section des statistiques tokens centralis√©e
    st.subheader("üìä Statistiques")
    
    session_tokens = st.session_state.get("api_usage", {}).get("current_session_tokens", 0)
    total_tokens = st.session_state.get("api_usage", {}).get("used_tokens", 0)
    
    # Affichage centr√© des m√©triques
    col1, col2 = st.columns(2)
    with col1:
        st.metric("üîë Session", session_tokens)
    with col2:
        st.metric("üìä Total", total_tokens)
    
    st.markdown("---")
    st.info("üí° Assistant IA pour le sourcing et recrutement")

# -------------------- Onglets --------------------
tab1, tab2, tab3, tab4, tab5, tab6, tab7, tab8, tab9, tab10 = st.tabs([
    "üîç Boolean", "üéØ X-Ray", "üîé CSE LinkedIn", "üê∂ Dogpile", 
    "üï∑Ô∏è Web Scraper", "‚úâÔ∏è InMail", "ü§ñ Magicien", "üìß Permutateur", "üìö Biblioth√®que", "üìÑ Nomination"
])

# -------------------- Tab 1: Boolean Search --------------------
with tab1:
    st.header("üîç Recherche Boolean")
    
    # Option fiche de poste pour l'IA
    with st.expander("üìÑ Fiche de poste (optionnel - pour enrichissement IA)", expanded=False):
        st.markdown("**Choisissez votre m√©thode d'import :**")
        
        # Onglets pour les deux m√©thodes
        tab_text, tab_pdf = st.tabs(["üìù Coller le texte", "üìÑ Uploader PDF"])
        
        fiche_content = ""
        
        with tab_text:
            fiche_poste = st.text_area(
                "Collez ici la fiche de poste compl√®te:",
                height=180,
                key="boolean_fiche_poste",
                placeholder="Mission: ...\nProfil recherch√©: ...\nComp√©tences requises: ...\nExp√©rience: ...\nFormation: ...\nAvantages: ..."
            )
            if fiche_poste:
                fiche_content = fiche_poste
        
        with tab_pdf:
            uploaded_file = st.file_uploader(
                "Choisissez votre fichier PDF:",
                type=['pdf'],
                key="boolean_pdf_uploader",
                help="Formats accept√©s: PDF (max 10MB)"
            )
            
            if uploaded_file is not None:
                try:
                    # V√©rification de la taille du fichier (max 10MB)
                    if uploaded_file.size > 10 * 1024 * 1024:
                        st.error("‚ùå Fichier trop volumineux (max 10MB)")
                    else:
                        st.success(f"‚úÖ Fichier '{uploaded_file.name}' upload√© avec succ√®s!")
                        
                        with st.spinner("üìÑ Extraction du texte en cours..."):
                            # Extraction r√©elle du PDF
                            extracted_text = extract_text_from_pdf(uploaded_file)
                            
                            if extracted_text and "Erreur" not in extracted_text and not extracted_text.strip().startswith("Aucun texte lisible trouv√©"):
                                fiche_content = extracted_text
                                st.success("‚úÖ Texte extrait avec succ√®s!")
                                
                                # Aper√ßu du contenu avec possibilit√© d'√©dition
                                fiche_content = st.text_area(
                                    "Aper√ßu et √©dition du contenu extrait:",
                                    value=extracted_text[:3000] + ("..." if len(extracted_text) > 3000 else ""),
                                    height=200,
                                    help="Vous pouvez modifier le texte extrait si n√©cessaire"
                                )
                            else:
                                # Afficher un message plus clair selon la raison
                                if extracted_text and extracted_text.strip().startswith("Aucun texte lisible trouv√©"):
                                    st.error("‚ùå Aucun texte lisible trouv√© dans le PDF. Il s'agit probablement d'un PDF scann√© (images) ou prot√©g√©.\nüí° Collez manuellement le contenu dans l'onglet 'Coller le texte' ou utilisez un OCR externe.")
                                else:
                                    st.error(f"‚ùå {extracted_text}")
                                # Fallback: permettre √† l'utilisateur de coller le texte manuellement
                                st.warning("üí° Collez manuellement le contenu dans l'onglet 'Coller le texte'")
                                
                except Exception as e:
                    st.error(f"‚ùå Erreur lors du traitement du PDF: {str(e)}")
                    st.warning("üí° Collez manuellement le contenu dans l'onglet 'Coller le texte'")
        
        # Bouton d'analyse commun
        if fiche_content and st.button("ü§ñ Analyser la fiche et pr√©-remplir", key="analyze_fiche", use_container_width=True):
            with st.spinner("üîç Analyse de la fiche en cours..."):
                # Simulation d'analyse de fiche de poste
                analyze_prompt = f"Analyse cette fiche de poste et extrait les √©l√©ments cl√©s:\n{fiche_content}\n\nExtrait:\n1. Titre du poste\n2. 2-3 synonymes du poste\n3. 2-3 comp√©tences obligatoires\n4. 2-3 comp√©tences optionnelles\n5. Mots √† exclure si mentionn√©s"
                result = ask_deepseek([{"role": "user", "content": analyze_prompt}], max_tokens=200)
                
                if result["content"].strip():
                    st.success("‚úÖ Analyse termin√©e ! Utilisez les suggestions ci-dessous pour remplir les champs.")
                    
                    # Affichage des suggestions de mani√®re plus structur√©e
                    with st.container():
                        st.markdown("### üí° Suggestions de l'IA:")
                        suggestions = result["content"].split('\n')
                        for suggestion in suggestions:
                            if suggestion.strip():
                                st.markdown(f"‚Ä¢ {suggestion.strip()}")
                    
                    st.info("ÔøΩ Copiez ces suggestions dans les champs correspondants ci-dessous")
                else:
                    st.warning("‚ö†Ô∏è Impossible d'analyser la fiche. Remplissez manuellement les champs ci-dessous.")
    
    col1, col2 = st.columns(2)
    with col1:
        poste = st.text_input("Poste recherch√©:", key="boolean_poste", placeholder="Ex: Ing√©nieur de travaux")
        synonymes = st.text_input("Synonymes:", key="boolean_synonymes", placeholder="Ex: Conducteur de travaux, Chef de chantier")
        competences_obligatoires = st.text_input("Comp√©tences obligatoires:", key="boolean_comp_oblig", placeholder="Ex: Autocad, Robot Structural Analysis")
        secteur = st.text_input("Secteur d'activit√©:", key="boolean_secteur", placeholder="Ex: BTP, Construction")
    with col2:
        competences_optionnelles = st.text_input("Comp√©tences optionnelles:", key="boolean_comp_opt", placeholder="Ex: Primavera, ArchiCAD")
        exclusions = st.text_input("Mots √† exclure:", key="boolean_exclusions", placeholder="Ex: Stage, Alternance")
        localisation = st.text_input("Localisation:", key="boolean_loc", placeholder="Ex: Casablanca")
        employeur = st.text_input("Employeur:", key="boolean_employeur", placeholder="Ex: TGCC")

    # Mode avanc√© LinkedIn pour Boolean
    with st.expander("‚öôÔ∏è Mode avanc√© LinkedIn", expanded=False):
        col_adv1, col_adv2 = st.columns(2)
        with col_adv1:
            specialite = st.text_input("Sp√©cialit√©", key="boolean_specialite", placeholder="Ex: G√©nie Civil, Informatique")
            entreprises_precedentes = st.text_input("Entreprises pr√©c√©dentes", key="boolean_entreprises_prec", placeholder="Ex: OCP, TGCC")
        with col_adv2:
            ecoles_cibles = st.text_input("√âcoles/Universit√©s", key="boolean_ecoles", placeholder="Ex: EMI, ENSA")
            certifications_bool = st.text_input("Certifications", key="boolean_certifications", placeholder="Ex: PMP, ISO 27001")

    gen_mode = st.selectbox("G√©n√©rer la requ√™te Boolean par :", ["Algorithme", "Intelligence artificielle"], key="boolean_gen_mode")
    if gen_mode == "Intelligence artificielle":
        st.caption("üí° L'IA enrichit les synonymes de fa√ßon conservatrice pour maximiser les r√©sultats LinkedIn")
    
    gen_btn = st.button("üîç G√©n√©rer", type="primary", key="boolean_generate_main", use_container_width=True)
    if gen_btn:
        if gen_mode == "Algorithme":
            with st.spinner("‚è≥ G√©n√©ration en cours..."):
                start_time = time.time()
                base_query = generate_boolean_query(
                    poste, synonymes, competences_obligatoires,
                    competences_optionnelles, exclusions, localisation, secteur
                )
                if employeur:
                    base_query += f' AND ("{employeur}")'
                
                # Add advanced LinkedIn filters
                advanced_parts = []
                if specialite:
                    advanced_parts.append(f'"{specialite}"')
                if entreprises_precedentes:
                    ent_terms = _split_terms(entreprises_precedentes)
                    if ent_terms:
                        advanced_parts.append(_or_group(ent_terms))
                if ecoles_cibles:
                    ecole_terms = _split_terms(ecoles_cibles)
                    if ecole_terms:
                        advanced_parts.append(_or_group(ecole_terms))
                if certifications_bool:
                    cert_terms = _split_terms(certifications_bool)
                    if cert_terms:
                        advanced_parts.append(_or_group(cert_terms))
                
                if advanced_parts:
                    st.session_state["boolean_query"] = base_query + " AND " + " AND ".join(advanced_parts)
                else:
                    st.session_state["boolean_query"] = base_query
                
                st.session_state["boolean_snapshot"] = {
                    "poste": poste,
                    "synonymes": synonymes,
                    "comp_ob": competences_obligatoires,
                    "comp_opt": competences_optionnelles,
                    "exclusions": exclusions,
                    "localisation": localisation,
                    "secteur": secteur,
                    "employeur": employeur or "",
                    "mode": gen_mode,
                    "specialite": specialite,
                    "entreprises_precedentes": entreprises_precedentes,
                    "ecoles_cibles": ecoles_cibles,
                    "certifications_bool": certifications_bool
                }
                total_time = time.time() - start_time
                st.success(f"‚úÖ Requ√™te g√©n√©r√©e en {total_time:.1f}s")
                st.rerun()  # Force la mise √† jour de l'affichage
        else:
            with st.spinner("ü§ñ G√©n√©ration Intelligence artificielle en cours..."):
                start_time = time.time()
                
                # Construct the AI prompt dynamically to avoid redundant fields
                prompt_parts = [
                    f"Poste: {poste}",
                    f"Synonymes: {synonymes}" if synonymes else "",
                    f"Comp√©tences obligatoires: {competences_obligatoires}" if competences_obligatoires else "",
                    f"Comp√©tences optionnelles: {competences_optionnelles}" if competences_optionnelles else "",
                    f"Exclusions: {exclusions}" if exclusions else "",
                    f"Localisation: {localisation}" if localisation else "",
                    f"Secteur: {secteur}" if secteur else "",
                    f"Employeur: {employeur}" if employeur else ""
                ]
                prompt = "G√©n√®re une requ√™te Boolean pour le sourcing avec les crit√®res suivants:\n" + "\n".join(filter(None, prompt_parts))

                ia_result = ask_deepseek([{"role": "user", "content": prompt}], max_tokens=200)

                # Extract enriched synonyms and mandatory skills
                synonymes_ia = ia_result.get("content", synonymes) if ia_result.get("content", "").strip() else synonymes
                comp_ob_ia = ia_result.get("comp_ob_ia", competences_obligatoires)

                # Generate the Boolean query with advanced fields
                base_query = generate_boolean_query(
                    poste, synonymes_ia, comp_ob_ia,
                    competences_optionnelles, exclusions, localisation, secteur, employeur
                )
                
                # Add advanced LinkedIn filters
                advanced_parts = []
                if specialite:
                    advanced_parts.append(f'"{specialite}"')
                if entreprises_precedentes:
                    ent_terms = _split_terms(entreprises_precedentes)
                    if ent_terms:
                        advanced_parts.append(_or_group(ent_terms))
                if ecoles_cibles:
                    ecole_terms = _split_terms(ecoles_cibles)
                    if ecole_terms:
                        advanced_parts.append(_or_group(ecole_terms))
                if certifications_bool:
                    cert_terms = _split_terms(certifications_bool)
                    if cert_terms:
                        advanced_parts.append(_or_group(cert_terms))
                
                if advanced_parts:
                    st.session_state["boolean_query"] = base_query + " AND " + " AND ".join(advanced_parts)
                else:
                    st.session_state["boolean_query"] = base_query

                # Ensure the query ends with NOT if exclusions are specified
                if exclusions:
                    st.session_state["boolean_query"] = st.session_state["boolean_query"].rstrip(" AND") + " NOT"

                # Save the snapshot with the actual values used
                st.session_state["boolean_snapshot"] = {
                    "poste": poste,
                    "synonymes": synonymes_ia,
                    "comp_ob": comp_ob_ia,
                    "comp_opt": competences_optionnelles,
                    "exclusions": exclusions,
                    "localisation": localisation,
                    "secteur": secteur,
                    "employeur": employeur or "",
                    "mode": gen_mode
                }
                total_time = time.time() - start_time
                st.success(f"‚úÖ Requ√™te Boolean g√©n√©r√©e par Intelligence artificielle en {total_time:.1f}s")
                if synonymes_ia != synonymes or comp_ob_ia != competences_obligatoires:
                    st.info("ü§ñ L'IA a enrichi vos crit√®res pour optimiser les r√©sultats LinkedIn")
                st.rerun()  # Force la mise √† jour de l'affichage

    # Affichage unifi√© de la requ√™te Boolean
    snap = st.session_state.get("boolean_snapshot", {})
    query_value = st.session_state.get("boolean_query", "")
    
    # V√©rifier si les param√®tres ont chang√© pour l'indication visuelle
    params_changed = False
    if snap and query_value:
        params_changed = any([
            snap.get("poste") != poste,
            snap.get("comp_ob") != competences_obligatoires,
            snap.get("comp_opt") != competences_optionnelles,
            snap.get("exclusions") != exclusions,
            snap.get("localisation") != localisation,
            snap.get("secteur") != secteur,
            snap.get("employeur") != (employeur or ""),
            snap.get("specialite") != specialite,
            snap.get("entreprises_precedentes") != entreprises_precedentes,
            snap.get("ecoles_cibles") != ecoles_cibles,
            snap.get("certifications_bool") != certifications_bool
        ])
    
    # Label avec indication si obsol√®te
    label = "Requ√™te Boolean:"
    if params_changed:
        label += " ‚ö†Ô∏è (Requ√™te obsol√®te - crit√®res modifi√©s - R√©g√©n√©rez pour mettre √† jour)"
    
    # Widget unifi√© - SANS KEY pour permettre la mise √† jour automatique
    placeholder_text = "Remplissez les crit√®res ci-dessus puis cliquez sur 'G√©n√©rer la requ√™te Boolean'" if not query_value else ""
    st.text_area(label, value=query_value, height=120, placeholder=placeholder_text)
    
    # Boutons et actions (seulement si requ√™te existe)
    if st.session_state.get("boolean_query"):
        # Zone commentaire
        boolean_commentaire = st.text_input("Commentaire (optionnel)", value=st.session_state.get("boolean_commentaire", ""), key="boolean_commentaire")
        # Boutons organis√©s : Copier, Sauvegarder, LinkedIn
        cols_actions = st.columns([0.2,0.4,0.4])
        with cols_actions[0]:
            safe_boolean = st.session_state.get('boolean_query', '').replace('"', '&quot;')
            st.markdown(f'<button data-copy="{safe_boolean}">üìã Copier</button>', unsafe_allow_html=True)
        with cols_actions[1]:
            if st.button("üíæ Sauvegarder", key="boolean_save", use_container_width=True):
                entry = {
                    "timestamp": datetime.now().strftime("%Y-%m-%d %H:%M"),
                    "type": "Boolean",
                    "poste": poste,
                    "requete": st.session_state["boolean_query"],
                    "utilisateur": st.session_state.get("user", ""),
                    "source": "Boolean",
                    "commentaire": st.session_state.get("boolean_commentaire", "")
                }
                st.session_state.library_entries.append(entry)
                save_library_entries()
                save_sourcing_entry_to_gsheet(entry)
                st.success("‚úÖ Sauvegard√©")
        with cols_actions[2]:
            url_linkedin = f"https://www.linkedin.com/search/results/people/?keywords={quote(st.session_state['boolean_query'])}"
            st.link_button("üåê Ouvrir sur LinkedIn", url_linkedin, use_container_width=True)

    # Variantes (seulement si requ√™te existe)
    if st.session_state.get("boolean_query"):
        # G√©n√©rer les variantes avec les valeurs ACTUELLES des champs
        variants = generate_boolean_variants(st.session_state["boolean_query"], synonymes, competences_optionnelles)
        
        st.caption("üîÄ Variantes propos√©es")
        if variants:
            for idx, (title, vq) in enumerate(variants):
                # Supprimer la key pour permettre la mise √† jour automatique
                st.text_area(f"{title}", value=vq, height=80)
                st.text_input(f"Commentaire variante {idx+1}", value=st.session_state.get(f"boolean_commentaire_var_{idx}", ""), key=f"boolean_commentaire_var_{idx}")
                cols_var = st.columns([0.2,0.4,0.4])
                with cols_var[0]:
                    safe_vq = vq.replace('"', '&quot;')
                    st.markdown(f'<button data-copy="{safe_vq}">üìã Copier</button>', unsafe_allow_html=True)
                with cols_var[1]:
                    if st.button(f"üíæ Sauvegarder {idx+1}", key=f"bool_save_{idx}", use_container_width=True):
                        entry = {
                            "timestamp": datetime.now().strftime("%Y-%m-%d %H:%M"),
                            "type": "Boolean",
                            "poste": poste,
                            "requete": vq,
                            "utilisateur": st.session_state.get("user", ""),
                            "source": f"Boolean Variante {idx+1}",
                            "commentaire": st.session_state.get(f"boolean_commentaire_var_{idx}", "")
                        }
                        st.session_state.library_entries.append(entry)
                        save_library_entries()
                        save_sourcing_entry_to_gsheet(entry)
                        st.success(f"‚úÖ Variante {idx+1} sauvegard√©e")
                with cols_var[2]:
                    url_var = f"https://www.linkedin.com/search/results/people/?keywords={quote(vq)}"
                    st.link_button(f"üåê LinkedIn {idx+1}", url_var, use_container_width=True)
        else:
            st.info("Aucune variante g√©n√©r√©e pour la requ√™te actuelle.")



# -------------------- Tab 2: X-Ray --------------------
with tab2:
    st.header("üéØ X-Ray Google")
    col1, col2 = st.columns(2)
    with col1:
        site_cible = st.selectbox("Site cible:", ["LinkedIn", "GitHub"], key="xray_site")
        poste_xray = st.text_input("Poste:", key="xray_poste", placeholder="Ex: D√©veloppeur Python")
        mots_cles = st.text_input("Mots-cl√©s:", key="xray_mots_cles", placeholder="Ex: Django, Flask")
    with col2:
        localisation_xray = st.text_input("Localisation:", key="xray_loc", placeholder="Ex: Casablanca")
        exclusions_xray = st.text_input("Mots √† exclure:", key="xray_exclusions", placeholder="Ex: Stage, Junior")

    if st.button("üîç Construire X-Ray", type="primary", key="xray_build", use_container_width=True):
        with st.spinner("‚è≥ G√©n√©ration en cours..."):
            start_time = time.time()
            # Logic for X-Ray query generation
            xray_query = generate_xray_query(site_cible, poste_xray, mots_cles, localisation_xray)
            
            # Add exclusions if any
            if exclusions_xray:
                exclusion_terms = _split_terms(exclusions_xray)
                if exclusion_terms:
                    xray_query += " -" + " -".join(exclusion_terms)
            
            st.session_state["xray_query"] = xray_query
            st.session_state["xray_snapshot"] = {
                "site": site_cible,
                "poste": poste_xray,
                "mots_cles": mots_cles,
                "localisation": localisation_xray,
                "exclusions": exclusions_xray
            }
            total_time = time.time() - start_time
            st.success(f"‚úÖ Requ√™te X-Ray g√©n√©r√©e en {total_time:.1f}s")

    # Affichage unifi√© de la requ√™te X-Ray
    snapx = st.session_state.get("xray_snapshot", {})
    query_value_xray = st.session_state.get("xray_query", "")
    
    # V√©rifier si les param√®tres ont chang√© pour l'indication visuelle
    params_changed_xray = False
    if snapx and query_value_xray:
        params_changed_xray = any([
            snapx.get("site") != site_cible,
            snapx.get("poste") != poste_xray,
            snapx.get("mots_cles") != mots_cles,
            snapx.get("localisation") != localisation_xray,
            snapx.get("exclusions") != exclusions_xray
        ])
    
    # Label avec indication si obsol√®te
    label_xray = "Requ√™te X-Ray:"
    if params_changed_xray:
        label_xray += " ‚ö†Ô∏è (Requ√™te obsol√®te - param√®tres modifi√©s - Reconstruire pour mettre √† jour)"
    
    # Widget unifi√© - SANS KEY pour permettre la mise √† jour automatique
    placeholder_text_xray = "Remplissez les crit√®res ci-dessus puis cliquez sur 'Construire X-Ray'" if not query_value_xray else ""
    st.text_area(label_xray, value=query_value_xray, height=120, placeholder=placeholder_text_xray)
    
    # Commentaires et actions
    if st.session_state.get("xray_query"):
        # Zone commentaire
        xray_commentaire = st.text_input("Commentaire (optionnel)", value=st.session_state.get("xray_commentaire", ""), key="xray_commentaire")
        # Boutons organis√©s : Copier, Sauvegarder, LinkedIn
        cols_actions = st.columns([0.2, 0.4, 0.4])
        with cols_actions[0]:
            safe_xray = st.session_state.get('xray_query', '').replace('"', '&quot;')
            st.markdown(f'<button data-copy="{safe_xray}">üìã Copier</button>', unsafe_allow_html=True)
        with cols_actions[1]:
            if st.button("üíæ Sauvegarder", key="xray_save", use_container_width=True):
                entry = {
                    "timestamp": datetime.now().strftime("%Y-%m-%d %H:%M"),
                    "type": "X-Ray",
                    "poste": poste_xray,
                    "requete": st.session_state["xray_query"],
                    "utilisateur": st.session_state.get("user", ""),
                    "source": "X-Ray",
                    "commentaire": st.session_state.get("xray_commentaire", "")
                }
                st.session_state.library_entries.append(entry)
                save_library_entries()
                save_sourcing_entry_to_gsheet(entry)
                st.success("‚úÖ Sauvegard√©")
        with cols_actions[2]:
            url_xray = f"https://www.google.com/search?q={quote(st.session_state['xray_query'])}"
            st.link_button("üåê Ouvrir sur Google", url_xray, use_container_width=True)

    # Variantes
    if st.session_state.get("xray_query"):
        x_vars = generate_xray_variants(st.session_state["xray_query"], poste_xray, mots_cles, localisation_xray)
        if x_vars:
            st.caption("üîÄ Variantes propos√©es")
            for i, (title, qv) in enumerate(x_vars):
                st.text_area(title, value=qv, height=80, key=f"xray_var_{i}")
                st.text_input(f"Commentaire variante {i+1}", value=st.session_state.get(f"xray_commentaire_var_{i}", ""), key=f"xray_commentaire_var_{i}")
                cols_var = st.columns([0.33, 0.33, 0.34])
                with cols_var[0]:
                    safe_qv = qv.replace('"', '&quot;')
                    st.markdown(f'<button data-copy="{safe_qv}">üìã Copier</button>', unsafe_allow_html=True)
                with cols_var[1]:
                    if st.button("üíæ Sauvegarder", key=f"xray_save_var_{i}", use_container_width=True):
                        entry = {
                            "timestamp": datetime.now().strftime("%Y-%m-%d %H:%M"),
                            "type": "X-Ray Variante",
                            "poste": poste_xray,
                            "requete": qv,
                            "utilisateur": st.session_state.get("user", ""),
                            "source": "X-Ray",
                            "commentaire": st.session_state.get(f"xray_commentaire_var_{i}", "")
                        }
                        st.session_state.library_entries.append(entry)
                        save_library_entries()
                        save_sourcing_entry_to_gsheet(entry)
                        st.success("‚úÖ Sauvegard√©")
                with cols_var[2]:
                    url_var = f"https://www.google.com/search?q={quote(qv)}"
                    st.link_button("üåê Ouvrir", url_var, use_container_width=True)
        else:
            st.info("Aucune variante g√©n√©r√©e pour la requ√™te actuelle.")

# -------------------- Tab 3: CSE --------------------
with tab3:
    st.header("üîé CSE LinkedIn")
    col1, col2 = st.columns(2)
    with col1:
        poste_cse = st.text_input("Poste recherch√©:", key="cse_poste", placeholder="Ex: D√©veloppeur Python")
        competences_cse = st.text_input("Comp√©tences cl√©s:", key="cse_comp", placeholder="Ex: Django, Flask")
    with col2:
        localisation_cse = st.text_input("Localisation:", key="cse_loc", placeholder="Ex: Casablanca")
        entreprise_cse = st.text_input("Entreprise:", key="cse_ent", placeholder="Ex: TGCC")

    if st.button("üîç Lancer recherche CSE", type="primary", width="stretch", key="cse_search"):
        with st.spinner("‚è≥ Construction de la requ√™te..."):
            start_time = time.time()
            query_parts = []
            if poste_cse: query_parts.append(poste_cse)
            if competences_cse: query_parts.append(competences_cse)
            if localisation_cse: query_parts.append(localisation_cse)
            if entreprise_cse: query_parts.append(entreprise_cse)
            st.session_state["cse_query"] = " ".join(query_parts)
            total_time = time.time() - start_time
            st.success(f"‚úÖ Requ√™te g√©n√©r√©e en {total_time:.1f}s")

    if st.session_state.get("cse_query"):
        st.text_area("Requ√™te CSE:", value=st.session_state["cse_query"], height=100, key="cse_area")
        
        # Boutons align√©s : Copier, Sauvegarder, Ouvrir
        cols_actions = st.columns([0.33, 0.33, 0.34])
        with cols_actions[0]:
            safe_cse = st.session_state.get('cse_query', '').replace('"', '&quot;')
            st.markdown(f'<button data-copy="{safe_cse}">üìã Copier</button>', unsafe_allow_html=True)
        with cols_actions[1]:
            if st.button("üíæ Sauvegarder", key="cse_save", use_container_width=True):
                entry = {
                    "timestamp": datetime.now().strftime("%Y-%m-%d %H:%M"),
                    "type": "CSE",
                    "poste": poste_cse,
                    "requete": st.session_state["cse_query"],
                    "utilisateur": st.session_state.get("user", ""),
                    "source": "CSE",
                    "commentaire": ""
                }
                st.session_state.library_entries.append(entry)
                save_library_entries()
                save_sourcing_entry_to_gsheet(entry)
                st.success("‚úÖ Sauvegard√©")
        with cols_actions[2]:
            cse_url = f"https://cse.google.fr/cse?cx=004681564711251150295:d-_vw4klvjg&q={quote(st.session_state['cse_query'])}"
            st.link_button("üåê Ouvrir sur CSE", cse_url, use_container_width=True)

# -------------------- Tab 4: Dogpile --------------------
with tab4:
    st.header("üê∂ Dogpile Search")
    query = st.text_input("Requ√™te Dogpile:", key="dogpile_query_input", placeholder="Ex: Python developer Casablanca")
    if st.button("üîç Rechercher", key="dogpile_search_btn", type="primary", use_container_width=True):
        if query:
            st.session_state["dogpile_query"] = query
            st.session_state["dogpile_snapshot"] = query
            st.success("‚úÖ Requ√™te enregistr√©e")
            st.rerun()

    # Affichage uniquement si une requ√™te a √©t√© g√©n√©r√©e
    if st.session_state.get("dogpile_query"):
        # Affichage unifi√© avec d√©tection de changement
        snap_dogpile = st.session_state.get("dogpile_snapshot", "")
        query_value_dogpile = st.session_state.get("dogpile_query", "")
        
        # V√©rifier si les param√®tres ont chang√©
        params_changed_dogpile = False
        if snap_dogpile and query_value_dogpile:
            params_changed_dogpile = snap_dogpile != query
        
        # Label avec indication si obsol√®te
        label_dogpile = "Requ√™te Dogpile:"
        if params_changed_dogpile:
            label_dogpile += " ‚ö†Ô∏è (Requ√™te obsol√®te - param√®tres modifi√©s - Rechercher pour mettre √† jour)"
        
        # Widget unifi√©
        st.text_area(label_dogpile, value=query_value_dogpile, height=80)
    
    # Boutons et commentaires (seulement si requ√™te existe)
    if st.session_state.get("dogpile_query"):
        # Zone commentaire
        dogpile_commentaire = st.text_input("Commentaire (optionnel)", value=st.session_state.get("dogpile_commentaire", ""), key="dogpile_commentaire")
        
        # Boutons organis√©s : Copier, Sauvegarder, Ouvrir
        cols_actions = st.columns([0.33, 0.33, 0.34])
        with cols_actions[0]:
            safe_dogpile = st.session_state.get('dogpile_query', '').replace('"', '&quot;')
            st.markdown(f'<button data-copy="{safe_dogpile}">üìã Copier</button>', unsafe_allow_html=True)
        with cols_actions[1]:
            if st.button("üíæ Sauvegarder", key="dogpile_save_btn", use_container_width=True):
                entry = {
                    "timestamp": datetime.now().strftime("%Y-%m-%d %H:%M"),
                    "type": "Dogpile",
                    "poste": "Recherche Dogpile",
                    "requete": st.session_state["dogpile_query"],
                    "utilisateur": st.session_state.get("user", ""),
                    "source": "Dogpile",
                    "commentaire": st.session_state.get("dogpile_commentaire", "")
                }
                st.session_state.library_entries.append(entry)
                save_library_entries()
                save_sourcing_entry_to_gsheet(entry)
                st.success("‚úÖ Sauvegard√©")
        with cols_actions[2]:
            dogpile_url = f"http://www.dogpile.com/serp?q={quote(st.session_state['dogpile_query'])}"
            st.link_button("üåê Ouvrir sur Dogpile", dogpile_url, use_container_width=True)

# -------------------- Tab 5: Web Scraper - Analyse Concurrentielle --------------------
# -------------------- Tab 5: Web Scraper - Analyse Concurrentielle --------------------
with tab5:
    st.header("üîç Analyse Concurrentielle - Offres d'Emploi")
    
    # Configuration du scraping
    with st.expander("‚öôÔ∏è Configuration", expanded=True):
        col1, col2 = st.columns(2)
        with col1:
            concurrents = st.text_area(
                "Sites des concurrents √† analyser (1 par ligne):", 
                placeholder="https://jobs.vinci.com/fr/recherche-d'offres/Maroc\nhttps://www.rekrute.com/sogea-maroc-emploi.html",
                height=100
            )
            max_pages = st.slider("Nombre maximum de pages √† analyser par site:", 1, 20, 5)
        
        with col2:
            mots_cles = st.text_input(
                "Mots-cl√©s √† rechercher (s√©par√©s par des virgules):",
                placeholder="ing√©nieur, coordinateur, m√©canicien, acheteur"
            )
            delay = st.slider("D√©lai entre les requ√™tes (secondes):", 1, 10, 3)
    
    # Options d'analyse
    with st.expander("üìä Options d'analyse", expanded=False):
        analyse_options = st.multiselect(
            "√âl√©ments √† analyser:",
            ["Comp√©tences recherch√©es", "Niveaux d'exp√©rience", "Avantages propos√©s", 
             "Types de contrats", "Localisations", "Salaires mentionn√©s", "Processus de recrutement"],
            default=["Comp√©tences recherch√©es", "Niveaux d'exp√©rience", "Avantages propos√©s"]
        )
    
    if st.button("üöÄ Lancer l'analyse concurrentielle", width="stretch", key="scraper_btn"):
        if concurrents:
            concurrents_list = [url.strip() for url in concurrents.split('\n') if url.strip()]
            mots_cles_list = [mot.strip().lower() for mot in mots_cles.split(',')] if mots_cles else []
            
            # Initialiser les r√©sultats
            results = {
                "concurrent": [],
                "url": [],
                "titre_poste": [],
                "competences": [],
                "experience": [],
                "avantages": [],
                "mots_cles_trouves": []
            }
            
            progress_bar = st.progress(0)
            status_text = st.empty()
            
            for i, url in enumerate(concurrents_list):
                status_text.text(f"Analyse de {url}...")
                
                try:
                    # Simulation de scraping - √Ä remplacer par votre logique r√©elle
                    time.sleep(delay)  # Respect du d√©lai
                    
                    # V√©rifier si c'est le site Vinci
                    if "vinci.com" in url:
                        try:
                            # Tentative de scraping r√©el du site Vinci
                            headers = {
                                'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/91.0.4472.124 Safari/537.36'
                            }
                            response = requests.get(url, headers=headers, timeout=10)
                            soup = BeautifulSoup(response.content, 'html.parser')
                            
                            # Chercher les √©l√©ments qui contiennent les offres d'emploi
                            # (Cette s√©lecteur est un exemple et doit √™tre adapt√© au site r√©el)
                            offres = soup.select('.job-listing, .offer-item, .job-title, [class*="job"]')
                            
                            if offres:
                                for offre in offres[:20]:  # Limiter √† 20 offres
                                    try:
                                        titre = offre.get_text(strip=True)
                                        if titre and len(titre) > 10:  # Filtrer les textes trop courts
                                            results["concurrent"].append("Vinci")
                                            results["url"].append(url)
                                            results["titre_poste"].append(titre)
                                            results["competences"].append("√Ä analyser")
                                            results["experience"].append("Non sp√©cifi√©")
                                            results["avantages"].append("√Ä analyser")
                                            
                                            # V√©rifier quels mots-cl√©s correspondent
                                            mots_trouves = []
                                            for mot in mots_cles_list:
                                                if mot in titre.lower():
                                                    mots_trouves.append(mot)
                                            results["mots_cles_trouves"].append(", ".join(mots_trouves))
                                    except:
                                        continue
                            else:
                                st.warning(f"Aucune offre trouv√©e sur {url}. Utilisation des donn√©es simul√©es.")
                                # Fallback aux donn√©es simul√©es si le scraping √©choue
                                postes_vinci = [
                                    {"titre": "Coordinateur HSE", "competences": "HSE, Normes de s√©curit√©, Gestion des risques", "experience": "5+ ans", "avantages": "Assurance, Formation, Transport"},
                                    {"titre": "Ing√©nieur √©lectrom√©canicien - Traitement des Eaux", "competences": "√âlectrom√©canique, Traitement des eaux, Maintenance", "experience": "3+ ans", "avantages": "Logement, Transport, Mutuelle"},
                                    # ... (ajouter d'autres postes simul√©s)
                                ]
                                
                                for poste in postes_vinci:
                                    results["concurrent"].append("Vinci")
                                    results["url"].append(url)
                                    results["titre_poste"].append(poste["titre"])
                                    results["competences"].append(poste["competences"])
                                    results["experience"].append(poste["experience"])
                                    results["avantages"].append(poste["avantages"])
                                    mots_trouves = []
                                    for mot in mots_cles_list:
                                        if mot in poste["titre"].lower() or mot in poste["competences"].lower():
                                            mots_trouves.append(mot)
                                    results["mots_cles_trouves"].append(", ".join(mots_trouves))
                                
                        except Exception as e:
                            st.error(f"Erreur lors du scraping de {url}: {str(e)}")
                            # Fallback aux donn√©es simul√©es en cas d'erreur
                            # ... (code de fallback similaire √† ci-dessus)
                    
                    # V√©rifier si c'est le site Rekrute (Sogea Maroc)
                    elif "rekrute.com" in url and "sogea" in url:
                        try:
                            # Tentative de scraping r√©el du site Rekrute
                            headers = {
                                'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/91.0.4472.124 Safari/537.36'
                            }
                            response = requests.get(url, headers=headers, timeout=10)
                            soup = BeautifulSoup(response.content, 'html.parser')
                            
                            # Chercher les √©l√©ments qui contiennent les offres d'emploi
                            # (Cette s√©lecteur est un exemple et doit √™tre adapt√© au site r√©el)
                            offres = soup.select('.job-item, .offer-title, [class*="job"]')
                            
                            if offres:
                                for offre in offres[:10]:  # Limiter √† 10 offres
                                    try:
                                        titre = offre.get_text(strip=True)
                                        if titre and len(titre) > 10:  # Filtrer les textes trop courts
                                            results["concurrent"].append("Sogea Maroc (Vinci)")
                                            results["url"].append(url)
                                            results["titre_poste"].append(titre)
                                            results["competences"].append("√Ä analyser")
                                            results["experience"].append("Non sp√©cifi√©")
                                            results["avantages"].append("√Ä analyser")
                                            
                                            # V√©rifier quels mots-cl√©s correspondent
                                            mots_trouves = []
                                            for mot in mots_cles_list:
                                                if mot in titre.lower():
                                                    mots_trouves.append(mot)
                                            results["mots_cles_trouves"].append(", ".join(mots_trouves))
                                    except:
                                        continue
                            else:
                                st.warning(f"Aucune offre trouv√©e sur {url}. Utilisation des donn√©es simul√©es.")
                                # Fallback aux donn√©es simul√©es si le scraping √©choue
                                postes_sogea = [
                                    {"titre": "Directeur de Travaux Hydraulique (H/F)", "competences": "Hydraulique, Gestion de projet, Management", "experience": "10+ ans", "avantages": "Voiture de fonction, Logement, Assurance"},
                                    {"titre": "M√©canicien Atelier", "competences": "M√©canique, R√©paration, Maintenance", "experience": "3+ ans", "avantages": "Transport, Formation, Prime de performance"},
                                    # ... (ajouter d'autres postes simul√©s)
                                ]
                                
                                for poste in postes_sogea:
                                    results["concurrent"].append("Sogea Maroc (Vinci)")
                                    results["url"].append(url)
                                    results["titre_poste"].append(poste["titre"])
                                    results["competences"].append(poste["competences"])
                                    results["experience"].append(poste["experience"])
                                    results["avantages"].append(poste["avantages"])
                                    mots_trouves = []
                                    for mot in mots_cles_list:
                                        if mot in poste["titre"].lower() or mot in poste["competences"].lower():
                                            mots_trouves.append(mot)
                                    results["mots_cles_trouves"].append(", ".join(mots_trouves))
                                
                        except Exception as e:
                            st.error(f"Erreur lors du scraping de {url}: {str(e)}")
                            # Fallback aux donn√©es simul√©es en cas d'erreur
                            # ... (code de fallback similaire √† ci-dessus)
                    
                    # Pour les autres sites
                    else:
                        try:
                            # Tentative de scraping g√©n√©rique pour les autres sites
                            headers = {
                                'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/91.0.4472.124 Safari/537.36'
                            }
                            response = requests.get(url, headers=headers, timeout=10)
                            soup = BeautifulSoup(response.content, 'html.parser')
                            
                            # Chercher les √©l√©ments qui pourraient contenir des offres d'emploi
                            # (Cette approche est tr√®s g√©n√©rale et peut ne pas fonctionner)
                            potential_selectors = [
                                '.job', '.offer', '.employment', '.career', 
                                '[class*="job"]', '[class*="offer"]', '[class*="employment"]',
                                'h1', 'h2', 'h3', 'h4', 'h5', 'h6'  # Les titres peuvent contenir des offres
                            ]
                            
                            offres_trouvees = False
                            for selector in potential_selectors:
                                offres = soup.select(selector)
                                for offre in offres[:5]:  # Limiter √† 5 offres par s√©lecteur
                                    try:
                                        texte = offre.get_text(strip=True)
                                        if texte and len(texte) > 20 and len(texte) < 200:  # Filtrer les textes
                                            # V√©rifier si le texte ressemble √† un titre d'offre d'emploi
                                            mots_emploi = ["emploi", "job", "offre", "recrutement", "poste", "h/f", "f/h"]
                                            if any(mot in texte.lower() for mot in mots_emploi):
                                                results["concurrent"].append("Autre entreprise")
                                                results["url"].append(url)
                                                results["titre_poste"].append(texte)
                                                results["competences"].append("√Ä analyser")
                                                results["experience"].append("Non sp√©cifi√©")
                                                results["avantages"].append("√Ä analyser")
                                                
                                                # V√©rifier quels mots-cl√©s correspondent
                                                mots_trouves = []
                                                for mot in mots_cles_list:
                                                    if mot in texte.lower():
                                                        mots_trouves.append(mot)
                                                results["mots_cles_trouves"].append(", ".join(mots_trouves))
                                                offres_trouvees = True
                                    except:
                                        continue
                            
                            if not offres_trouvees:
                                st.warning(f"Aucune offre d√©tect√©e sur {url}. Le site peut n√©cessiter une configuration sp√©cifique.")
                                # Ajouter une entr√©e g√©n√©rique
                                results["concurrent"].append("Autre entreprise")
                                results["url"].append(url)
                                results["titre_poste"].append("Poste vari√© - Analyse manuelle requise")
                                results["competences"].append("Comp√©tences diverses")
                                results["experience"].append("Non sp√©cifi√©")
                                results["avantages"].append("Avantages standards")
                                results["mots_cles_trouves"].append("")
                                
                        except Exception as e:
                            st.error(f"Erreur lors du scraping de {url}: {str(e)}")
                            # Ajouter une entr√©e d'erreur
                            results["concurrent"].append("Erreur de scraping")
                            results["url"].append(url)
                            results["titre_poste"].append(f"Erreur: {str(e)}")
                            results["competences"].append("N/A")
                            results["experience"].append("N/A")
                            results["avantages"].append("N/A")
                            results["mots_cles_trouves"].append("")
                
                except Exception as e:
                    st.error(f"Erreur avec {url}: {str(e)}")
                    # Ajouter une entr√©e d'erreur
                    results["concurrent"].append("Erreur")
                    results["url"].append(url)
                    results["titre_poste"].append(f"Erreur: {str(e)}")
                    results["competences"].append("N/A")
                    results["experience"].append("N/A")
                    results["avantages"].append("N/A")
                    results["mots_cles_trouves"].append("")
                
                progress_bar.progress((i + 1) / len(concurrents_list))
            
            status_text.text("Analyse termin√©e!")
            
            # Affichage des r√©sultats
            if results["concurrent"]:
                total_postes = len(results["concurrent"])
                st.success(f"‚úÖ {total_postes} postes trouv√©s sur {len(concurrents_list)} sites")
                
                # Cr√©ation d'un DataFrame pour une meilleure visualisation
                try:
                    df_results = pd.DataFrame(results)
                    st.dataframe(df_results, width="stretch")
                    
                    # Afficher un r√©sum√© par entreprise
                    st.subheader("üìä R√©sum√© par entreprise")
                    entreprises = {}
                    for i, entreprise in enumerate(results["concurrent"]):
                        if entreprise not in entreprises:
                            entreprises[entreprise] = 0
                        entreprises[entreprise] += 1
                    
                    for entreprise, count in entreprises.items():
                        st.write(f"- **{entreprise}**: {count} poste(s)")
                        
                except NameError:
                    st.error("Erreur: pandas n'est pas install√©. Impossible de cr√©er le DataFrame.")
                    # On continue sans DataFrame
                    for i, concurrent in enumerate(results["concurrent"]):
                        st.write(f"**{concurrent}** - {results['titre_poste'][i]}")
                        st.write(f"Comp√©tences: {results['competences'][i]}")
                        st.write(f"Exp√©rience: {results['experience'][i]}")
                        st.write(f"Avantages: {results['avantages'][i]}")
                        st.write("---")
                
                # Analyses avanc√©es
                st.subheader("üìà Analyses")
                
                # Nuage de mots des comp√©tences recherch√©es
                if "Comp√©tences recherch√©es" in analyse_options:
                    st.write("**Comp√©tences les plus recherch√©es:**")
                    all_skills = ", ".join(results["competences"]).lower()
                    skills_counter = Counter([skill.strip() for skill in all_skills.split(',')])
                    
                    if skills_counter:
                        # Affichage simplifi√© des comp√©tences (sans nuage de mots)
                        st.write("R√©partition des comp√©tences:")
                        for skill, count in skills_counter.most_common(10):
                            st.write(f"- {skill}: {count} occurrence(s)")
                
                # Analyse des niveaux d'exp√©rience
                if "Niveaux d'exp√©rience" in analyse_options:
                    st.write("**Niveaux d'exp√©rience requis:**")
                    exp_counter = Counter(results["experience"])
                    for exp, count in exp_counter.items():
                        st.write(f"- {exp}: {count} offre(s)")
                
                # Export des r√©sultats (uniquement si pandas est disponible)
                try:
                    csv_data = df_results.to_csv(index=False).encode('utf-8')
                    st.download_button(
                        label="üì• T√©l√©charger les r√©sultats (CSV)",
                        data=csv_data,
                        file_name="analyse_concurrentielle_emplois.csv",
                        mime="text/csv",
                        width="stretch"
                    )
                except NameError:
                    st.warning("Impossible de g√©n√©rer le fichier CSV car pandas n'est pas disponible.")
            else:
                st.warning("Aucun r√©sultat √† afficher.")
        else:
            st.error("Veuillez entrer au moins une URL de concurrent √† analyser.")
    
    # Section d'aide
    with st.expander("‚ùì Comment utiliser cet outil", expanded=False):
        st.markdown("""
        ### Guide d'utilisation de l'analyse concurrentielle
        
        1. **Listez les sites de vos concurrents** - Entrez les URLs des pages carri√®res ou offres d'emploi
        2. **D√©finissez les mots-cl√©s** - Sp√©cifiez les comp√©tences ou postes qui vous int√©ressent
        3. **Configurez l'analyse** - Choisissez ce que vous voulez analyser pr√©cis√©ment
        4. **Lancez l'extraction** - L'outil parcourt les sites et extrait les informations
        5. **Consultez les r√©sultats** - Visualisez les tendances et t√©l√©chargez les donn√©es
        
        ### Conseils pour de meilleurs r√©sultats:
        - Ciblez des pages listant plusieurs offres d'emploi
        - Utilisez des mots-cl√©s pr√©cis li√©s √† vos besoins
        - Augmentez le d√©lai entre les requ√™tes pour √©viter le blocage
        - Testez d'abord avec 2-3 sites pour valider la configuration
        
        ### Limitations:
        - Le scraping web peut √™tre bloqu√© par certains sites
        - La structure des pages peut changer, n√©cessitant une mise √† jour des s√©lecteurs
        - Certains sites utilisent JavaScript pour charger le contenu, ce qui peut ne pas √™tre compatible avec cette approche
        """)

# -------------------- Tab 6: InMail --------------------
with tab6:
    st.header("‚úâÔ∏è G√©n√©rateur d'InMail Personnalis√©")

    # --------- FONCTIONS UTILES ---------
    def generate_cta(cta_type, prenom, genre):
        suffix = "e" if genre == "F√©minin" else ""
        if cta_type == "Proposer un appel":
            return f"Je serai ravi{suffix} d'√©changer avec vous par t√©l√©phone cette semaine afin d‚Äôen discuter davantage."
        elif cta_type == "Partager le CV":
            return f"Seriez-vous int√©ress√©{suffix} √† partager votre CV afin que je puisse examiner cette opportunit√© avec vous ?"
        elif cta_type == "D√©couvrir l'opportunit√© sur notre site":
            return f"Souhaiteriez-vous consulter plus de d√©tails sur cette opportunit√© via notre site carri√®re ?"
        elif cta_type == "Accepter un rendez-vous":
            return f"Je serai ravi{suffix} de convenir d‚Äôun rendez-vous afin d‚Äô√©changer sur cette opportunit√©."
        return ""

    # --------- PARAM√àTRES G√âN√âRAUX ---------
    col1, col2, col3, col4 = st.columns(4)
    with col1:
        url_linkedin = st.text_input("Profil LinkedIn", key="inmail_url", placeholder="linkedin.com/in/nom-prenom")
    with col2:
        entreprise = st.selectbox("Entreprise", ["TGCC", "TG ALU", "TG COVER", "TG WOOD", "TG STEEL", "TG STONE", "TGEM", "TGCC Immobilier"], key="inmail_entreprise")
    with col3:
        ton_message = st.selectbox("Ton du message", ["Persuasif", "Professionnel", "Convivial", "Direct"], key="inmail_ton")
    with col4:
        genre_profil = st.selectbox("Genre du profil", ["Masculin", "F√©minin"], key="inmail_genre")

    col5, col6, col7, col8 = st.columns(4)
    with col5:
        poste_accroche = st.text_input("Poste √† pourvoir", key="inmail_poste", placeholder="Ex: Directeur Financier")
    with col6:
        longueur_message = st.slider("Longueur (mots)", 10, 200, 50, key="inmail_longueur")
    with col7:
        analyse_profil = st.selectbox("M√©thode d'analyse du profil LinkedIn", ["Manuel", "Intelligence artificielle"], index=0, key="inmail_analyse")
        if analyse_profil == "Intelligence artificielle":
            # V√©rifier si l'API est disponible
            api_key = st.secrets.get("DEEPSEEK_API_KEY")
            if not api_key:
                st.warning("‚ö†Ô∏è API non configur√©e")
    with col8:
        cta_option = st.selectbox("Call to action (Conclusion)", ["Proposer un appel", "Partager le CV", "D√©couvrir l'opportunit√© sur notre site", "Accepter un rendez-vous"], key="inmail_cta")

    # --------- INFORMATIONS CANDIDAT ---------
    with st.expander("üìä Informations candidat", expanded=False):
        default_profil = {
            "prenom": "",
            "nom": "",
            "poste_actuel": "",
            "entreprise_actuelle": "",
            "competences_cles": ["", "", ""],
            "experience_annees": "",
            "formation": "",
            "localisation": ""
        }
        profil_data = {**default_profil, **st.session_state.get("inmail_profil_data", {})}

        cols = st.columns(5)
        profil_data["prenom"] = cols[0].text_input("Pr√©nom", profil_data.get("prenom", ""), key="inmail_prenom")
        profil_data["nom"] = cols[1].text_input("Nom", profil_data.get("nom", ""), key="inmail_nom")
        profil_data["poste_actuel"] = cols[2].text_input("Poste actuel", profil_data.get("poste_actuel", ""), key="inmail_poste_actuel")
        profil_data["entreprise_actuelle"] = cols[3].text_input("Entreprise actuelle", profil_data.get("entreprise_actuelle", ""), key="inmail_entreprise_actuelle")
        profil_data["experience_annees"] = cols[4].text_input("Ann√©es d'exp√©rience", profil_data.get("experience_annees", ""), key="inmail_exp")

        cols2 = st.columns(5)
        profil_data["formation"] = cols2[0].text_input("Domaine de formation", profil_data.get("formation", ""), key="inmail_formation")
        profil_data["competences_cles"][0] = cols2[1].text_input("Comp√©tence 1", profil_data["competences_cles"][0], key="inmail_comp1")
        profil_data["competences_cles"][1] = cols2[2].text_input("Comp√©tence 2", profil_data["competences_cles"][1], key="inmail_comp2")
        profil_data["competences_cles"][2] = cols2[3].text_input("Comp√©tence 3", profil_data["competences_cles"][2], key="inmail_comp3")
        profil_data["localisation"] = cols2[4].text_input("Localisation", profil_data.get("localisation", ""), key="inmail_loc")

    # Garder les donn√©es du profil √† jour
    if not st.session_state.get("inmail_profil_data"):
        st.session_state["inmail_profil_data"] = profil_data
    else:
        # Mettre √† jour avec les nouvelles valeurs
        st.session_state["inmail_profil_data"].update(profil_data)

    # --------- G√âN√âRATION ---------
    if st.button("‚ú® G√©n√©rer", type="primary", use_container_width=True, key="btn_generate_inmail"):
        donnees_profil = st.session_state.get("inmail_profil_data", profil_data)
        
        # Si Intelligence artificielle est s√©lectionn√©e, avertir de la limitation LinkedIn
        if analyse_profil == "Intelligence artificielle" and url_linkedin.strip():
            st.warning("‚ö†Ô∏è **Limitation** : L'IA ne peut pas acc√©der directement aux profils LinkedIn pour des raisons de s√©curit√© et de confidentialit√©.")
            st.info("üí° **Recommandation** : Veuillez remplir manuellement les informations du candidat ci-dessus pour une g√©n√©ration d'InMail pr√©cise et personnalis√©e.")
            
            # D√©sactiver l'analyse automatique LinkedIn et utiliser les donn√©es manuelles
            # with st.spinner("ü§ñ Analyse IA du profil LinkedIn..."):
            #     analyse_prompt = f"""
            #     IMPORTANT: Tu ne peux pas acc√©der aux profils LinkedIn r√©els. 
            #     Au lieu d'inventer des informations, r√©ponds avec ce JSON d'erreur :
            #     {{
            #         "erreur": "Impossible d'acc√©der au profil LinkedIn",
            #         "message": "Veuillez remplir les informations manuellement"
            #     }}
            #     """

        
        # Utiliser l'IA pour g√©n√©rer le message
        ia_prompt = f"""
        G√©n√®re UNIQUEMENT le contenu du message InMail personnalis√© (sans objet, sans titre) avec les informations suivantes:
        - Candidat: {donnees_profil.get('prenom', '')} {donnees_profil.get('nom', '')}
        - Poste actuel: {donnees_profil.get('poste_actuel', '')}
        - Entreprise actuelle: {donnees_profil.get('entreprise_actuelle', '')}
        - Comp√©tences: {', '.join(filter(None, donnees_profil.get('competences_cles', [])))}
        - Formation: {donnees_profil.get('formation', '')}
        - Exp√©rience: {donnees_profil.get('experience_annees', '')} ans
        - Localisation: {donnees_profil.get('localisation', '')}
        
        Poste √† pourvoir: {poste_accroche}
        Entreprise: {entreprise}
        Ton: {ton_message}
        Genre: {genre_profil}
        Call-to-action: {cta_option}
        
        Le message doit faire environ {longueur_message} mots, √™tre {ton_message.lower()}, et commencer directement par la salutation (ex: Bonjour {donnees_profil.get('prenom', 'Candidat')},).
        Ne pas inclure d'objet, de titre ou de formatage en gras pour l'objet.
        """
        
        with st.spinner("ü§ñ G√©n√©ration IA en cours..."):
            ia_result = get_deepseek_response(ia_prompt, [], "normale", "InMail Generation")
            if ia_result.get("content") and "Erreur: Cl√© API DeepSeek manquante" not in ia_result["content"]:
                msg = ia_result["content"]
            elif "Erreur: Cl√© API DeepSeek manquante" in str(ia_result.get("content", "")):
                st.error("üîë **Configuration manquante** : La cl√© API DeepSeek n'est pas configur√©e.")
                st.info("üí° **Solution** : Contactez l'administrateur pour configurer la cl√© API DeepSeek dans les secrets Streamlit.")
                st.warning("‚ö†Ô∏è **Mode de secours** : G√©n√©ration d'un message de base sans IA.")
                # Fallback avec un message g√©n√©rique
                msg = f"""Bonjour {donnees_profil.get('prenom', 'Candidat')},

Votre profil a retenu notre attention chez {entreprise}. Nous recherchons actuellement un {poste_accroche} et pensons que votre exp√©rience pourrait correspondre √† nos besoins.

Seriez-vous disponible pour √©changer sur cette opportunit√© ?

Cordialement."""
            else:
                # Fallback si l'IA ne r√©pond pas
                msg = generate_inmail(donnees_profil, poste_accroche, entreprise, ton_message, longueur_message, cta_option, genre_profil, "entreprise")
        
        st.session_state["inmail_message"] = msg
        st.session_state["inmail_objet"] = "Nouvelle opportunit√©: " + poste_accroche
        st.session_state["inmail_generated"] = True
        st.session_state["inmail_snapshot"] = {
            "poste_accroche": poste_accroche,
            "entreprise": entreprise,
            "ton_message": ton_message,
            "longueur_message": longueur_message,
            "cta_option": cta_option,
            "genre_profil": genre_profil,
            "profil_data": donnees_profil.copy()
        }

    # --------- R√âSULTAT ---------
    if st.session_state.get("inmail_generated"):
        # V√©rifier si les param√®tres ont chang√©
        snap_inmail = st.session_state.get("inmail_snapshot", {})
        current_profil = st.session_state.get("inmail_profil_data", profil_data)
        
        params_changed_inmail = False
        if snap_inmail:
            params_changed_inmail = any([
                snap_inmail.get("poste_accroche") != poste_accroche,
                snap_inmail.get("entreprise") != entreprise,
                snap_inmail.get("ton_message") != ton_message,
                snap_inmail.get("longueur_message") != longueur_message,
                snap_inmail.get("cta_option") != cta_option,
                snap_inmail.get("genre_profil") != genre_profil,
                snap_inmail.get("profil_data") != current_profil
            ])
        
        # Titre avec indication si obsol√®te
        titre_inmail = "üìù Message InMail g√©n√©r√©"
        if params_changed_inmail:
            titre_inmail += " ‚ö†Ô∏è (Param√®tres modifi√©s - R√©g√©n√©rer pour mettre √† jour)"
        
        st.subheader(titre_inmail)
        st.text_input("üìß Objet", st.session_state.get("inmail_objet", ""), key="inmail_objet_display")
        
        # Debug: Afficher les informations du message
        msg = st.session_state.get("inmail_message", "")
        if not msg:
            msg = "Aucun message g√©n√©r√©"
        
        # Utiliser une cl√© dynamique pour forcer la mise √† jour du contenu
        import time
        dynamic_key = f"inmail_msg_display_{int(time.time() * 1000) % 10000}"
        
        try:
            st.text_area("Message", value=msg, height=250, key=dynamic_key)
        except Exception as e:
            st.error(f"Erreur affichage message: {e}")
            st.text_area("Message", msg, height=250, key=f"fallback_{dynamic_key}")
        st.caption(f"üìè {len(msg.split())} mots | {len(msg)} caract√®res")

        col1, col2 = st.columns(2)
        with col1:
            if st.button("üîÑ R√©g√©n√©rer (nouvelle version)", key="btn_regen_inmail"):
                donnees_profil = st.session_state.get("inmail_profil_data", profil_data)
                
                # G√©n√©rer une nouvelle version avec l'IA
                ia_prompt = f"""
                G√©n√®re une NOUVELLE version d'un message InMail UNIQUEMENT LE CONTENU (sans objet, sans titre) diff√©rente de la pr√©c√©dente avec:
                - Candidat: {donnees_profil.get('prenom', '')} {donnees_profil.get('nom', '')}
                - Poste actuel: {donnees_profil.get('poste_actuel', '')}
                - Entreprise actuelle: {donnees_profil.get('entreprise_actuelle', '')}
                - Comp√©tences: {', '.join(filter(None, donnees_profil.get('competences_cles', [])))}
                - Formation: {donnees_profil.get('formation', '')}
                - Exp√©rience: {donnees_profil.get('experience_annees', '')} ans
                
                Poste √† pourvoir: {poste_accroche}
                Entreprise: {entreprise}
                Ton: {ton_message}
                
                G√©n√®re une approche diff√©rente, avec un angle nouveau mais professionnel.
                Commence directement par la salutation (ex: Bonjour {donnees_profil.get('prenom', 'Candidat')},).
                Ne pas inclure d'objet, de titre ou de formatage en gras.
                """
                
                with st.spinner("üîÑ R√©g√©n√©ration IA en cours..."):
                    ia_result = get_deepseek_response(ia_prompt, [], "normale", "InMail Regeneration")
                    if ia_result.get("content"):
                        new_msg = ia_result["content"]
                    else:
                        new_msg = generate_inmail(donnees_profil, poste_accroche, entreprise, ton_message, longueur_message, cta_option, genre_profil, "entreprise")
                
                st.session_state["inmail_message"] = new_msg
                st.session_state["inmail_objet"] = "Nouvelle opportunit√©: " + poste_accroche
                # Forcer le rechargement de la page pour mettre √† jour l'affichage
                st.rerun()
        with col2:
            if st.button("üíæ Sauvegarder comme mod√®le", key="btn_save_inmail"):
                entry = {
                    "timestamp": datetime.now().strftime("%Y-%m-%d %H:%M"),
                    "type": "InMail",
                    "poste": poste_accroche,
                    "requete": st.session_state["inmail_message"],
                    "utilisateur": st.session_state.get("user", ""),
                    "source": "InMail",
                    "commentaire": f"Ton: {ton_message}, Longueur: {longueur_message} mots"
                }
                st.session_state.library_entries.append(entry)
                save_library_entries()
                save_sourcing_entry_to_gsheet(entry)
                st.success(f"‚úÖ Mod√®le '{poste_accroche} - {entry['timestamp']}' sauvegard√©")


# -------------------- Tab 7: Magicien --------------------
with tab7:
    st.header("ü§ñ Magicien de sourcing")

    questions_pretes = [
        "Quels sont les synonymes possibles pour le m√©tier de",
        "Quels outils ou logiciels sont li√©s au m√©tier de", 
        "Quels mots-cl√©s pour cibler les juniors pour le poste de",
        "Quels intitul√©s similaires au poste de",
        "Quels crit√®res √©liminatoires fr√©quents pour le poste de",
        "Quels secteurs d'activit√© embauchent souvent pour le poste de",
        "Quelles certifications utiles pour le m√©tier de",
        "Quels r√¥les proches √† consid√©rer lors du sourcing pour",
        "Quelles tendances de recrutement r√©centes pour le m√©tier de"
    ]

    # Zone unique fusionn√©e : selectbox avec option "Autre" pour saisie libre
    option_choisie = st.selectbox(
        "üìå Choisissez une question ou tapez la v√¥tre :",
        questions_pretes + ["Autre (tapez votre question)"],
        help="S√©lectionnez une question pr√™te ou choisissez 'Autre' pour taper votre propre question."
    )
    
    # Zone de saisie unique selon le choix
    if option_choisie == "Autre (tapez votre question)":
        question_complete = st.text_input(
            "Votre question :",
            placeholder="Ex: Quelles sont les comp√©tences cl√©s pour un chef de projet BTP ?"
        )
    else:
        question_complete = st.text_input(
            "Compl√©tez la question :",
            value=option_choisie + " ",
            placeholder="Ex: " + option_choisie + " d√©veloppeur web"
        )

    mode_rapide_magicien = st.checkbox("‚ö° R√©ponse concise", key="magicien_fast")
    
    if st.button("‚ú® Poser la question √† l'IA", type="primary", key="ask_magicien", use_container_width=True):
        if question_complete and question_complete.strip():
            with st.spinner("‚è≥ G√©n√©ration en cours..."):
                start_time = time.time()
                prompt = question_complete.strip()
                # Ajout d'une instruction pour forcer une liste d'intitul√©s/synonymes si la question le demande
                if "synonymes" in prompt.lower() or "intitul√©s similaires" in prompt.lower():
                    prompt += ". R√©ponds uniquement par une liste de synonymes ou intitul√©s similaires, s√©par√©s par des virgules, sans introduction."
                elif "outils" in prompt.lower() or "logiciels" in prompt.lower():
                    prompt += ". R√©ponds avec une liste √† puces des outils, sans introduction."
                elif "comp√©tences" in prompt.lower() or "skills" in prompt.lower():
                    prompt += ". R√©ponds avec une liste √† puces, sans introduction."
                if mode_rapide_magicien:
                    prompt += " R√©ponse concise et directe."
                result = get_deepseek_response(prompt, [], "normale" if not mode_rapide_magicien else "courte", "Magicien Sourcing")
                total_time = int(time.time() - start_time)
                st.success(f"‚úÖ R√©ponse g√©n√©r√©e en {total_time}s")
                if result.get("content"):
                    st.subheader("üí° R√©ponse :")
                    st.write(result["content"])
                    if not hasattr(st.session_state, 'magicien_history'):
                        st.session_state.magicien_history = []
                    st.session_state.magicien_history.append({
                        "q": prompt,
                        "r": result["content"],
                        "time": total_time
                    })
                else:
                    st.error("‚ùå Aucune r√©ponse g√©n√©r√©e. Veuillez reformuler votre question.")
        else:
            st.warning("‚ö†Ô∏è Veuillez saisir une question")

    if st.session_state.get("magicien_history"):
        st.subheader("üìù Historique des r√©ponses")
        for i, item in enumerate(st.session_state.magicien_history[::-1]):
            with st.expander(f"‚ùì {item['q']} ({item['time']}s)"):
                st.write(item["r"])
                if st.button("üóëÔ∏è Supprimer", key=f"del_magicien_{i}"):
                    st.session_state.magicien_history.remove(item)
                    st.rerun()
        
        if st.button("üßπ Supprimer tout", key="clear_magicien_all", width="stretch"):
            st.session_state.magicien_history.clear()
            st.success("‚úÖ Historique vid√©")
            st.rerun()
            
# -------------------- Tab 8: Permutateur --------------------
with tab8:
    st.header("üìß Permutateur Email")

    # G√©n√©ration de noms marocains al√©atoires
    if "random_names" not in st.session_state:
        import random
        noms_masculins = ["Ahmed", "Mohamed", "Youssef", "Omar", "Khalid", "Rachid", "Hassan", "Abdelkader", "Mustapha", "Sa√Ød"]
        noms_feminins = ["Fatima", "Aicha", "Khadija", "Zineb", "Salma", "Nadia", "Houda", "Laila", "Amina", "Sanaa"]
        noms_famille = ["Alami", "Bennani", "Cherkaoui", "Filali", "Idrissi", "Jamal", "Kettani", "Lahlou", "Mahfoudi", "Naciri", "Ouazzani", "Qadiri"]
        
        # S√©lectionner al√©atoirement
        random_prenom_m = random.choice(noms_masculins)
        random_prenom_f = random.choice(noms_feminins)
        random_nom = random.choice(noms_famille)
        
        st.session_state["random_names"] = {
            "masculin": f"{random_prenom_m}",
            "feminin": f"{random_prenom_f}",
            "nom": random_nom
        }

    col1, col2 = st.columns(2)
    with col1:
        prenom = st.text_input("Pr√©nom:", key="perm_prenom", placeholder=st.session_state['random_names']['masculin'])
        nom = st.text_input("Nom:", key="perm_nom", placeholder=st.session_state['random_names']['nom'])
    with col2:
        entreprise = st.text_input("Entreprise:", key="perm_entreprise", placeholder="TGCC")
        source = st.radio("Source de d√©tection :", ["Site officiel", "Charika.ma"], key="perm_source", horizontal=True)
    



    if st.button("üîÆ G√©n√©rer permutations", use_container_width=True):
        if prenom and nom and entreprise:
            with st.spinner("‚è≥ G√©n√©ration des permutations..."):
                start_time = time.time()
                permutations = []
                detected = get_email_from_charika(entreprise) if source == "Charika.ma" else None
                
                if detected and source == "Charika.ma":
                    st.info(f"üìß Format d√©tect√© sur Charika.ma : {detected}")
                    domain = detected.split("@")[1]
                elif source == "Charika.ma":
                    # Email non d√©tect√© sur Charika
                    st.error(f"‚ùå Format d'email non d√©tect√© sur Charika.ma pour '{entreprise}'")
                    domain = f"{entreprise.lower().replace(' ', '').replace('-', '')}.ma"
                    # Ajout du bouton/lien Google
                    google_url = get_charika_search_url(entreprise)
                    st.markdown(f"<a href='{google_url}' target='_blank' style='font-size:16px;'>üîé Rechercher sur Google</a>", unsafe_allow_html=True)
                else:
                    domain = f"{entreprise.lower().replace(' ', '').replace('-', '')}.ma"
                
                # G√©n√©ration des permutations
                patterns = [
                    f"{prenom.lower()}.{nom.lower()}@{domain}",
                    f"{prenom[0].lower()}{nom.lower()}@{domain}",
                    f"{nom.lower()}.{prenom.lower()}@{domain}",
                    f"{prenom.lower()}{nom.lower()}@{domain}",
                    f"{prenom.lower()}-{nom.lower()}@{domain}",
                    f"{nom.lower()}{prenom[0].lower()}@{domain}",
                    f"{prenom[0].lower()}.{nom.lower()}@{domain}",
                    f"{nom.lower()}.{prenom[0].lower()}@{domain}"
                ]
                
                total_time = time.time() - start_time
                st.session_state["perm_result"] = list(set(patterns))
                st.success(f"‚úÖ {len(patterns)} permutations g√©n√©r√©es en {total_time:.1f}s")
        else:
            st.warning("‚ö†Ô∏è Veuillez remplir tous les champs")

    if st.session_state.get("perm_result"):
        st.text_area("R√©sultats:", value="\n".join(st.session_state["perm_result"]), height=150)
        st.caption("üîç Tester sur : [Hunter.io](https://hunter.io/) ou [NeverBounce](https://neverbounce.com/)")

# -------------------- Tab 9: Biblioth√®que --------------------
with tab9:
    st.header("üìö Biblioth√®que des recherches")
    # Actualisation auto depuis Google Sheets
    entries_local = st.session_state.library_entries if st.session_state.library_entries else []
    entries_gsheet = load_sourcing_entries_from_gsheet()
    # Fusion et d√©duplication (par requ√™te + type + poste)
    all_entries = entries_local.copy()
    for e in entries_gsheet:
        if not any((e.get("requete") == x.get("requete") and e.get("type") == x.get("type") and e.get("poste") == x.get("poste")) for x in all_entries):
            all_entries.append(e)
    col1, col2 = st.columns(2)
    with col1:
        search_term = st.text_input("üîé Rechercher:", placeholder="Rechercher par poste ou requ√™te")
    with col2:
        sort_by = st.selectbox("üìå Trier par:", ["Date r√©cente", "Date ancienne", "Type", "Poste"], key="sort_by")

    entries = all_entries
    if search_term:
        entries = [e for e in entries if search_term.lower() in str(e.get("requete","")) .lower() or 
                 search_term.lower() in str(e.get("poste","")) .lower() or search_term.lower() in str(e.get("type","")) .lower()]

    # Utilise timestamp si pr√©sent, sinon date
    def get_date(e):
        return e.get("timestamp") or e.get("date") or ""

    if sort_by == "Type":
        entries = sorted(entries, key=lambda x: x.get("type",""))
    elif sort_by == "Poste":
        entries = sorted(entries, key=lambda x: x.get("poste",""))
    elif sort_by == "Date ancienne":
        entries = sorted(entries, key=get_date)
    else:
        entries = sorted(entries, key=get_date, reverse=True)

    st.info(f"üìä {len(entries)} recherche(s) trouv√©e(s)")
    for i, entry in enumerate(entries):
        with st.expander(f"{get_date(entry)} - {entry.get('type','')} - {entry.get('poste','')}"):
            st.text_area("Requ√™te:", value=entry.get('requete',''), height=100, key=f"req_{i}")
            col1, col2 = st.columns(2)
            with col1:
                if st.button("üóëÔ∏è Supprimer", key=f"del_{i}"):
                    if entry in st.session_state.library_entries:
                        st.session_state.library_entries.remove(entry)
                        save_library_entries()
                        st.success("‚úÖ Recherche supprim√©e")
                        st.rerun()
            with col2:
                if entry.get('type') == 'Boolean':
                    url = f"https://www.linkedin.com/search/results/people/?keywords={quote(entry.get('requete',''))}"
                    st.link_button("üåê Ouvrir", url)
                elif entry.get('type') == 'X-Ray':
                    url = f"https://www.google.com/search?q={quote(entry.get('requete',''))}"
                    st.link_button("üåê Ouvrir", url)
    if not entries:
        st.info("üìù Aucune recherche sauvegard√©e pour le moment")

# -------------------- Tab 10: Documents de Nomination --------------------
with tab10:
    st.header("üìÑ G√©n√©rateur de Documents de Nomination")
    
    # Fonction pour g√©n√©rer le document Word de nomination
    def generate_nomination_document(nom, prenom, poste, departement_rattachement, date_effet):
        try:
            from docx import Document
            from docx.shared import Inches, Pt
            from docx.enum.text import WD_ALIGN_PARAGRAPH
            from io import BytesIO
            
            doc = Document()
            
            # En-t√™te TGCC
            header_p = doc.add_paragraph()
            header_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = header_p.add_run("TGCC\nCONSTRUISONS ENSEMBLE")
            run.font.size = Pt(16)
            run.bold = True
            
            # Espacement
            doc.add_paragraph()
            doc.add_paragraph()
            
            # Objet
            objet_p = doc.add_paragraph()
            objet_run = objet_p.add_run(f"Objet : Annonce de nomination ‚Äî {prenom} {nom}")
            objet_run.bold = True
            
            doc.add_paragraph()
            
            # Corps du message
            doc.add_paragraph("Chers Collaborateurs,")
            doc.add_paragraph()
            
            message_p = doc.add_paragraph()
            message_p.add_run(f"Nous avons le plaisir de vous annoncer la nomination de {prenom} {nom} au poste de ")
            run_poste = message_p.add_run(poste)
            run_poste.bold = True
            message_p.add_run(f", effective √† compter du {date_effet}.")
            
            doc.add_paragraph()
            
            rattachement_p = doc.add_paragraph()
            rattachement_p.add_run(f"{prenom} {nom} sera rattach√© au ")
            run_dept = rattachement_p.add_run(departement_rattachement)
            run_dept.bold = True
            rattachement_p.add_run(" et aura pour mission de piloter la strat√©gie des ressources humaines de TGCC, en accompagnant notre croissance et en renfor√ßant notre culture d'entreprise.")
            
            doc.add_paragraph()
            
            doc.add_paragraph("Nous lui souhaitons la bienvenue au sein de notre groupe et lui exprimons notre enti√®re confiance pour relever les d√©fis √† venir.")
            doc.add_paragraph()
            doc.add_paragraph("Restons mobilis√©s pour construire ensemble !")
            doc.add_paragraph()
            
            # Ligne de s√©paration
            doc.add_paragraph("---")
            doc.add_paragraph()
            
            # Signature
            signature_p = doc.add_paragraph()
            signature_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            signature_run = signature_p.add_run(f"{departement_rattachement}\nTGCC ‚Äî CONSTRUISONS ENSEMBLE")
            signature_run.bold = True
            
            # Sauvegarder dans un buffer
            buffer = BytesIO()
            doc.save(buffer)
            buffer.seek(0)
            return buffer.getvalue()
            
        except ImportError:
            st.error("Erreur: La biblioth√®que python-docx n'est pas install√©e")
            return None
        except Exception as e:
            st.error(f"Erreur lors de la g√©n√©ration du document: {str(e)}")
            return None
    
    # Interface utilisateur
    col1, col2 = st.columns(2)
    with col1:
        nom = st.text_input("Nom du collaborateur:", key="nom_nomination", placeholder="DUPONT")
        poste = st.text_input("Poste:", key="poste_nomination", placeholder="Directeur du Capital Humain")
        date_effet = st.date_input("Date d'effet:", key="date_nomination")
    
    with col2:
        prenom = st.text_input("Pr√©nom du collaborateur:", key="prenom_nomination", placeholder="Jean")
        departement_rattachement = st.text_input("D√©partement de rattachement:", key="dept_nomination", placeholder="P√¥le des Affaires G√©n√©rales")
    
    if st.button("üìÑ G√©n√©rer Document de Nomination", type="primary", use_container_width=True):
        if nom and prenom and poste and departement_rattachement:
            with st.spinner("üìÑ G√©n√©ration du document en cours..."):
                doc_data = generate_nomination_document(
                    nom, prenom, poste, departement_rattachement, 
                    date_effet.strftime("%d %B %Y") if date_effet else "Date √† d√©finir"
                )
                
                if doc_data:
                    filename = f"Nomination_{prenom}_{nom}_{date_effet.strftime('%Y%m%d') if date_effet else 'DateADefinir'}.docx"
                    st.download_button(
                        label="üì• T√©l√©charger le Document Word",
                        data=doc_data,
                        file_name=filename,
                        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                        use_container_width=True
                    )
                    st.success("‚úÖ Document g√©n√©r√© avec succ√®s !")
                else:
                    st.error("‚ùå Erreur lors de la g√©n√©ration du document")
        else:
            st.warning("‚ö†Ô∏è Veuillez remplir tous les champs obligatoires")
    
    # Pr√©visualisation du contenu
    if nom and prenom and poste and departement_rattachement:
        st.subheader("üëÅÔ∏è Pr√©visualisation du contenu")
        with st.expander("Voir le contenu du document", expanded=False):
            preview_content = f"""
**TGCC**
**CONSTRUISONS ENSEMBLE**

---

**Objet : Annonce de nomination ‚Äî {prenom} {nom}**

Chers Collaborateurs,

Nous avons le plaisir de vous annoncer la nomination de **{prenom} {nom}** au poste de **{poste}**, effective √† compter du **{date_effet.strftime('%d %B %Y') if date_effet else 'Date √† d√©finir'}**.

{prenom} {nom} sera rattach√© au **{departement_rattachement}** et aura pour mission de piloter la strat√©gie des ressources humaines de TGCC, en accompagnant notre croissance et en renfor√ßant notre culture d'entreprise.

Nous lui souhaitons la bienvenue au sein de notre groupe et lui exprimons notre enti√®re confiance pour relever les d√©fis √† venir.

Restons mobilis√©s pour construire ensemble !

---

**{departement_rattachement}**
**TGCC ‚Äî CONSTRUISONS ENSEMBLE**
            """
            st.markdown(preview_content)

# -------------------- CSS pour masquer le prompt en bas --------------------