import streamlit as st
# -------------------- CONFIGURATION GOOGLE SHEETS pour la Biblioth√®que Sourcing --------------------
SOURCING_SHEET_URL = "https://docs.google.com/spreadsheets/d/1Yw99SS4vU5v0DuD7S1AwaEispJCo-cwioxSsAYnzRkE/edit"
SOURCING_WORKSHEET_NAME = "Sourcing DB"
SOURCING_HEADERS = [
    "timestamp", "type", "poste", "requete", "utilisateur", "source", "commentaire"
]

@st.cache_resource
def get_sourcing_gsheet_client():
    if not GSPREAD_AVAILABLE:
        return None
    try:
        service_account_info = _build_service_account_info_from_st_secrets()
        gc = gspread.service_account_from_dict(service_account_info)
        spreadsheet = gc.open_by_url(SOURCING_SHEET_URL)
        worksheet = spreadsheet.worksheet(SOURCING_WORKSHEET_NAME)
        return worksheet
    except Exception as e:
        st.error(f"‚ùå Erreur Google Sheets Sourcing DB: {e}")
        return None

def save_sourcing_entry_to_gsheet(entry: dict):
    worksheet = get_sourcing_gsheet_client()
    if worksheet is None:
        return False
    row = [entry.get(h, "") for h in SOURCING_HEADERS]
    try:
        worksheet.append_row(row)
        st.toast("‚úÖ Requ√™te sauvegard√©e dans Sourcing DB", icon="‚òÅÔ∏è")
        return True
    except Exception as e:
        st.error(f"‚ùå Erreur sauvegarde Sourcing DB: {e}")
        return False

def load_sourcing_entries_from_gsheet():
    worksheet = get_sourcing_gsheet_client()
    if worksheet is None:
        return []
    try:
        records = worksheet.get_all_records()
        return records
    except Exception as e:
        st.warning(f"Erreur chargement Sourcing DB: {e}")
        return []
import streamlit as st
# utils.py
# -*- coding: utf-8 -*-
import os
import json
import pandas as pd
from datetime import datetime
import io
import random
from io import BytesIO

# --- IMPORTS CONDITIONNELS (CORRIG√âS) ---
# Importations standard
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

# PDF (ReportLab)
try:
    from reportlab.lib.pagesizes import A4
    from reportlab.pdfgen import canvas
    from reportlab.lib.units import mm
    PDF_PRETTY_AVAILABLE = True
except ImportError:
    PDF_PRETTY_AVAILABLE = False
    class Dummy:
        def __init__(self, *args, **kwargs): pass
    SimpleDocTemplate, Table, TableStyle, Paragraph, Spacer, getSampleStyleSheet, colors = [Dummy] * 7
    A4 = None

# === AJOUT COMPATIBILIT√â ANCIEN NOM ===
PDF_AVAILABLE = PDF_PRETTY_AVAILABLE   # pour les fonctions existantes (export_brief_pdf)

# Word (Python-docx)
try:
    from docx import Document
    WORD_AVAILABLE = True
except ImportError:
    WORD_AVAILABLE = False

# Google Sheets & IA
try:
    import gspread
    from google.oauth2 import service_account
    GSPREAD_AVAILABLE = True
except ImportError:
    GSPREAD_AVAILABLE = False

def _build_service_account_info_from_st_secrets():
    """Build a service account dict from st.secrets supporting a couple of key-name variants.

    This helps when secrets were added under slightly different names (with or without _X509).
    Returns a dict suitable for gspread.service_account_from_dict or raises KeyError.
    """
    # Required keys
    required = [
        "GCP_TYPE", "GCP_PROJECT_ID", "GCP_PRIVATE_KEY_ID",
        "GCP_PRIVATE_KEY", "GCP_CLIENT_EMAIL", "GCP_CLIENT_ID"
    ]
    missing = [k for k in required if not st.secrets.get(k)]
    if missing:
        # Raise KeyError so callers can show a single informative message
        raise KeyError(
            ", ".join(missing)
        )

    # Build info using .get to avoid KeyError for optional fields
    private_key_raw = st.secrets.get("GCP_PRIVATE_KEY", "") or ""
    info = {
        "type": st.secrets.get("GCP_TYPE"),
        "project_id": st.secrets.get("GCP_PROJECT_ID"),
        "private_key_id": st.secrets.get("GCP_PRIVATE_KEY_ID"),
        "private_key": private_key_raw.replace("\\n", "\n") if private_key_raw else "",
        "client_email": st.secrets.get("GCP_CLIENT_EMAIL"),
        "client_id": st.secrets.get("GCP_CLIENT_ID"),
        "auth_uri": st.secrets.get("GCP_AUTH_URI", "https://accounts.google.com/o/oauth2/auth"),
        "token_uri": st.secrets.get("GCP_TOKEN_URI", "https://oauth2.googleapis.com/token"),
        # optional cert URLs with sensible defaults
        "auth_provider_x509_cert_url": st.secrets.get("GCP_AUTH_PROVIDER_X509_CERT_URL") or st.secrets.get("GCP_AUTH_PROVIDER_CERT_URL") or "https://www.googleapis.com/oauth2/v1/certs",
        "client_x509_cert_url": st.secrets.get("GCP_CLIENT_X509_CERT_URL") or st.secrets.get("GCP_CLIENT_CERT_URL") or "",
    }

    return info


def get_api_secret(name: str, alt_names=None, env_fallback=True):
    """R√©cup√®re un secret en essayant plusieurs variantes et en retombant sur les variables d'environnement.

    - name: nom principal du secret attendu dans st.secrets
    - alt_names: liste de noms alternatifs √† tester
    - env_fallback: si True, essaie os.environ[name] si st.secrets ne contient rien

    Retourne la valeur (str) ou None.
    """
    alt_names = alt_names or []
    # 1) st.secrets direct
    try:
        val = st.secrets.get(name)
        if val:
            return val
    except Exception:
        # st.secrets peut ne pas √™tre disponible en environnement de test
        pass

    # 2) noms alternatifs
    for alt in alt_names:
        try:
            val = st.secrets.get(alt)
            if val:
                return val
        except Exception:
            continue

    # 3) fallback vers variables d'environnement
    if env_fallback:
        val = os.environ.get(name) or next((os.environ.get(a) for a in alt_names if os.environ.get(a)), None)
        if val:
            return val

    return None

try:
    from openai import OpenAI
    OPENAI_CLIENT_AVAILABLE = True
except ImportError:
    OPENAI_CLIENT_AVAILABLE = False

# -------------------- CONFIGURATION GOOGLE SHEETS pour les Briefs --------------------
BRIEFS_SHEET_URL = "https://docs.google.com/spreadsheets/d/1DDw-aypZ9zDwAUuL6-p4TsKBIwV3pua8O3ACw_K-bHs/edit"
BRIEFS_WORKSHEET_NAME = "Briefs"

# Ent√™tes de colonnes (doivent correspondre EXACTEMENT √† la Ligne 1 de votre Google Sheet)
BRIEFS_HEADERS = [
    "BRIEF_NAME", "POSTE_INTITULE", "MANAGER_NOM", "RECRUTEUR",
    "AFFECTATION_TYPE", "AFFECTATION_NOM", "DATE_BRIEF",
    "RAISON_OUVERTURE", "IMPACT_STRATEGIQUE", "TACHES_PRINCIPALES",
    "MUST_HAVE_EXPERIENCE", "MUST_HAVE_DIPLOMES", "MUST_HAVE_COMPETENCES",
    "MUST_HAVE_SOFTSKILLS", "NICE_TO_HAVE_EXPERIENCE", "NICE_TO_HAVE_DIPLOMES",
    "NICE_TO_HAVE_COMPETENCES", "RATTACHEMENT", "BUDGET",
    "ENTREPRISES_PROFIL", "SYNONYMES_POSTE", "CANAUX_PROFIL",
    "LIEN_PROFIL_1", "LIEN_PROFIL_2", "LIEN_PROFIL_3",
    "COMMENTAIRES", "NOTES_LIBRES",
    "CRITERES_EXCLUSION", "PROCESSUS_EVALUATION", "MANAGER_NOTES",
    "MANAGER_COMMENTS_JSON", "KSA_MATRIX_JSON", "DATE_MAJ"
]

# -------------------- CONFIGURATION GOOGLE SHEETS pour les Annonces --------------------
ANNONCES_SHEET_URL = "https://docs.google.com/spreadsheets/d/1yibT8jcQ35a9VWAKpcdczKAHdnIOLi2UHnQvsfulH_s/edit"
ANNONCES_WORKSHEET_NAME = "Annonces"

# Ent√™tes de colonnes pour les annonces (doivent correspondre EXACTEMENT √† la Ligne 1 de votre Google Sheet)
ANNONCES_HEADERS = [
    "timestamp", "poste", "entreprise", "localisation", "plateforme", "format_type",
    "contenu", "fiche_text", "type_contrat", "niveau_experience",
    "formation_requise", "affectation", "competences_techniques",
    "missions_principales", "soft_skills", "date_creation", "fields_present"
]

# -------------------- FONCTIONS DE GESTION GOOGLE SHEETS (CORRIG√âES POUR SECRETS GCP_) --------------------

@st.cache_resource
def get_briefs_gsheet_client():
    """Initialise et retourne le client gspread en utilisant les secrets GCP_..."""
    if not GSPREAD_AVAILABLE:
        return None
    try:
        service_account_info = _build_service_account_info_from_st_secrets()
        gc = gspread.service_account_from_dict(service_account_info)
        spreadsheet = gc.open_by_url(BRIEFS_SHEET_URL)
        worksheet = spreadsheet.worksheet(BRIEFS_WORKSHEET_NAME)
        return worksheet
    except KeyError as e:
        st.error(f"‚ùå Cl√© de secret manquante pour Google Sheets: {e}. V√©rifiez la configuration des secrets GCP_...")
        return None
    except Exception as e:
        st.error(f"‚ùå Erreur de connexion/ouverture de Google Sheets: {e}")
        return None

def save_brief_to_gsheet(brief_name, brief_data):
    """Sauvegarde (cr√©ation / update) du brief dans Google Sheets avec normalisation champs."""
    worksheet = get_briefs_gsheet_client()
    if worksheet is None:
        return False

    # Normalisation cl√©s courtes -> longues (compat anciennes donn√©es)
    upgrade_map = {
        "MUST_HAVE_EXP": "MUST_HAVE_EXPERIENCE",
        "MUST_HAVE_DIP": "MUST_HAVE_DIPLOMES",
        "NICE_TO_HAVE_EXP": "NICE_TO_HAVE_EXPERIENCE",
        "NICE_TO_HAVE_DIP": "NICE_TO_HAVE_DIPLOMES"
    }
    for old, new in upgrade_map.items():
        if old in brief_data and new not in brief_data:
            brief_data[new] = brief_data[old]

    # Normalisation KSA si DataFrame pr√©sente
    ksa_df = brief_data.get("ksa_matrix")
    if isinstance(ksa_df, pd.DataFrame) and not ksa_df.empty:
        ksa_df = ksa_df.copy()
        # Harmonise colonnes possibles
        if "Cible / Standard attendu" in ksa_df.columns and "Question pour l'entretien" not in ksa_df.columns:
            ksa_df["Question pour l'entretien"] = ksa_df["Cible / Standard attendu"]
        if "√âchelle d'√©valuation (1-5)" in ksa_df.columns and "√âvaluation (1-5)" not in ksa_df.columns:
            ksa_df["√âvaluation (1-5)"] = ksa_df["√âchelle d'√©valuation (1-5)"]
        need = ["Rubrique","Crit√®re","Type de question","Question pour l'entretien","√âvaluation (1-5)","√âvaluateur"]
        for c in need:
            if c not in ksa_df.columns:
                ksa_df[c] = ""
        brief_data["KSA_MATRIX_JSON"] = ksa_df[need].to_json(orient="records", force_ascii=False)

    # Construit la ligne dans l‚Äôordre des headers
    row = []
    for header in BRIEFS_HEADERS:
        if header == "KSA_MATRIX_JSON":
            val = brief_data.get("KSA_MATRIX_JSON", "")
        elif header == "DATE_MAJ":
            val = datetime.now().strftime("%Y-%m-%d %H:%M:%S")
        else:
            val = brief_data.get(header, "")
        row.append("" if val is None else str(val))

    # Force le nom
    row[0] = brief_name
    try:
        cell = worksheet.find(brief_name, in_column=1, case_sensitive=True)
        last_col = col_to_letter(len(BRIEFS_HEADERS))
        if cell:
            worksheet.update(f"A{cell.row}:{last_col}{cell.row}", [row])
            st.toast(f"‚úÖ Brief mis √† jour (Sheets)", icon="‚òÅÔ∏è")
        else:
            worksheet.append_row(row)
            st.toast(f"‚úÖ Brief ajout√© (Sheets)", icon="‚òÅÔ∏è")
        return True
    except Exception as e:
        st.error(f"‚ùå Erreur sauvegarde Google Sheets: {e}")
        return False


# -------------------- Directory for Briefs --------------------
BRIEFS_DIR = "briefs"

def ensure_briefs_directory():
    """Ensure the briefs directory exists."""
    if not os.path.exists(BRIEFS_DIR):
        os.makedirs(BRIEFS_DIR)

# -------------------- Persistance (JSON locale - la version active) --------------------
def save_briefs():
    """
    Sauvegarde locale des briefs en JSON.
    - Convertit dates -> YYYY-MM-DD
    - Convertit DataFrame -> JSON (records)
    - Assure pr√©sence de KSA_MATRIX_JSON si ksa_matrix pr√©sent
    """
    import json, os
    from datetime import date, datetime
    import pandas as pd

    briefs = st.session_state.get("saved_briefs", {}) or {}

    def convert(obj):
        if isinstance(obj, (date, datetime)):
            return obj.strftime("%Y-%m-%d")
        if isinstance(obj, pd.DataFrame):
            # on ne garde pas la DataFrame brute
            return obj.to_json(orient="records", force_ascii=False)
        if isinstance(obj, dict):
            return {k: convert(v) for k, v in obj.items()}
        if isinstance(obj, list):
            return [convert(x) for x in obj]
        return obj

    safe = {}
    for name, data in briefs.items():
        # si une cl√© ksa_matrix (DataFrame) existe, la transformer en KSA_MATRIX_JSON
        if "ksa_matrix" in data and isinstance(data["ksa_matrix"], pd.DataFrame):
            try:
                data["KSA_MATRIX_JSON"] = data["ksa_matrix"].to_json(orient="records", force_ascii=False)
            except Exception:
                pass
            # on ne sauvegarde pas la DataFrame brute
            data.pop("ksa_matrix", None)
        safe[name] = convert(data)

    os.makedirs("briefs", exist_ok=True)
    try:
        with open("briefs/briefs.json", "w", encoding="utf-8") as f:
            json.dump(safe, f, ensure_ascii=False, indent=2)
    except Exception as e:
        st.error(f"Erreur lors de la sauvegarde locale des briefs: {e}")

def load_briefs():
    """Charge tous les briefs depuis Google Sheets."""
    try:
        briefs = {}
        worksheet = get_briefs_gsheet_client()
        if worksheet:
            records = worksheet.get_all_records()
            for record in records:
                brief_name = record.get("BRIEF_NAME")
                if brief_name:
                    briefs[brief_name] = record
                    if record.get("KSA_MATRIX_JSON"):
                        try:
                            record["ksa_matrix"] = pd.DataFrame.from_records(json.loads(record["KSA_MATRIX_JSON"]))
                        except Exception:
                            record["ksa_matrix"] = pd.DataFrame()
        return briefs
    except Exception as e:
        st.warning(f"Erreur lors du chargement des briefs depuis Google Sheets : {e}")
        return {}

def refresh_saved_briefs():
    """
    Recharge depuis Google Sheets et stocke dans session_state.saved_briefs.
    """
    briefs = load_briefs()
    st.session_state.saved_briefs = briefs
    return briefs

# === AJOUT : chargement local (manquait) ===
@st.cache_data(ttl=120)
def load_all_local_briefs():
    """
    Charge les briefs locaux (briefs/briefs.json + fichiers individuels).
    N'interroge PAS Google Sheets.
    """
    folder = "briefs"
    collected = {}
    global_path = os.path.join(folder, "briefs.json")
    try:
        if os.path.exists(global_path):
            with open(global_path, "r", encoding="utf-8") as f:
                collected = json.load(f)
        else:
            if os.path.isdir(folder):
                for fn in os.listdir(folder):
                    if fn.endswith(".json") and fn != "briefs.json":
                        try:
                            with open(os.path.join(folder, fn), "r", encoding="utf-8") as f:
                                collected[fn[:-5]] = json.load(f)
                        except:
                            continue
    except Exception:
        pass
    return collected

# (Optionnel : fusion locale + m√©moire)
def merge_local_with_session():
    local = load_all_local_briefs()
    mem = st.session_state.get("saved_briefs", {}) or {}
    merged = {**local, **mem}
    st.session_state.saved_briefs = merged
    return merged

# -------------------- Initialisation Session --------------------
def init_session_state():
    """Initialise l'√©tat de la session Streamlit avec des valeurs par d√©faut."""
    defaults = {
        "poste_intitule": "", "service": "", "niveau_hierarchique": "", "type_contrat": "",
        "localisation": "", "budget_salaire": "", "date_prise_poste": "", "recruteur": "",
        "manager_nom": "", "affectation_type": "", "affectation_nom": "", "raison_ouverture": "",
        "impact_strategique": "", "rattachement": "", "taches_principales": "", "must_have_experience": "",
        "must_have_diplomes": "", "must_have_competences": "", "must_have_softskills": "", 
        "nice_to_have_experience": "", "nice_to_have_diplomes": "", "nice_to_have_competences": "",
        "entreprises_profil": "", "synonymes_poste": "", "canaux_profil": "", "commentaires": "", 
        "notes_libres": "", "profil_links": ["", "", ""], "ksa_data": {}, "ksa_matrix": pd.DataFrame(),
        "saved_briefs": load_briefs(), "current_brief_name": None, "filtered_briefs": {},
        "show_filtered_results": False, "brief_data": {}, "comment_libre": "", "brief_phase": "Gestion",
        "saved_job_descriptions": [], "temp_extracted_data": None, "temp_job_title": "",
        "canaux_prioritaires": [], "criteres_exclusion": "", "processus_evaluation": "",
        "manager_comments": {}, "manager_notes": "", "job_library": load_library(),
    }
    for k, v in defaults.items():
        if k not in st.session_state:
            st.session_state[k] = v

# -------------------- Conseils IA --------------------
def generate_checklist_advice(section_title, field_title):
    """G√©n√®re un conseil IA pour les champs du brief."""
    advice_db = {
        "Contexte du poste": {
            "Raison de l'ouverture": [
                "- Clarifier si remplacement pour un chantier en cours, cr√©ation pour un nouveau projet BTP ou √©volution interne due √† une promotion.",
                "- Identifier le niveau d'urgence li√© aux d√©lais de construction et √† la priorisation des sites.",
                "- Expliquer le contexte strat√©gique dans le secteur BTP, comme un nouveau contrat de travaux publics.",
                "- Pr√©ciser si le poste est une cr√©ation pour renforcer l'√©quipe sur un grand chantier ou une r√©affectation pour optimiser les ressources.",
                "- Relier le poste √† la strat√©gie globale de l'entreprise en mati√®re de d√©veloppement durable dans le BTP."
            ],
            "Mission globale": [
                "- La mission globale consiste √† diriger l'√©quipe de gestion de chantier pour optimiser les processus de construction et respecter les normes de s√©curit√©.",
                "- Le r√¥le consiste √† am√©liorer l'efficacit√© sur les sites en supervisant les projets BTP complexes et en maintenant les d√©lais et le budget.",
                "- Superviser la transformation des chantiers et les projets d'innovation pour soutenir la comp√©titivit√© dans le BTP.",
                "- Assurer le bon d√©roulement des travaux en maintenant un √©quilibre entre qualit√©, co√ªts et respect des normes environnementales.",
                "- Garantir une communication fluide entre les √©quipes de terrain et les parties prenantes pour maximiser l'impact du projet BTP."
            ],
            "T√¢ches principales": [
                "- Piloter le process de recrutement pour des profils BTP, d√©finir la strat√©gie de sourcing sur les chantiers, interviewer les candidats, g√©rer les entretiens, effectuer le reporting.",
                "- Superviser la gestion des √©quipes sur site, organiser des formations s√©curit√©, garantir la conformit√© aux normes BTP, coordonner les projets transversaux.",
                "- Planifier les objectifs trimestriels pour les chantiers, analyser les donn√©es de co√ªts, optimiser les performances des √©quipes, pr√©parer les rapports de performance.",
                "- Coordonner les efforts entre les d√©partements BTP, g√©rer les budgets de mat√©riaux, superviser les ressources humaines sur site, suivre les projets et t√¢ches assign√©es.",
                "- Assurer la planification des travaux, organiser des sessions de formation aux normes de s√©curit√©, coordonner les activit√©s de d√©veloppement des talents en BTP."
            ]
        },
        "Must-have (Indispensables)": {
            "Exp√©rience": [
                "- Sp√©cifier le nombre d'ann√©es d'exp√©rience requis dans le BTP et le secteur des chantiers.",
                "- Mentionner les types de projets similaires, comme la gestion de grands travaux publics ou de construction r√©sidentielle."
            ],
            "Connaissances / Dipl√¥mes / Certifications": [
                "- Indiquer les dipl√¥mes exig√©s en g√©nie civil ou BTP, certifications comme Habilitation √âlectrique ou CACES.",
                "- Pr√©ciser les connaissances en normes de s√©curit√© (ISO 45001) ou r√©glementaires (RT 2012 pour le BTP)."
            ],
            "Comp√©tences / Outils": [
                "- Sugg√©rer des comp√©tences techniques comme la ma√Ætrise d'AutoCAD, Revit ou logiciels de gestion de chantier.",
                "- Exemple : 'Expertise en gestion de projet BTP avec m√©thodes Agile adapt√©es aux chantiers'."
            ],
            "Soft skills / aptitudes comportementales": [
                "- Sugg√©rer des aptitudes comme 'Leadership sur terrain', 'Rigueur en s√©curit√©', 'Communication avec sous-traitants', 'Autonomie en gestion de crises'."
            ]
        },
        "Nice-to-have (Atouts)": {
            "Exp√©rience additionnelle": [
                "- Ex. projets internationaux de BTP ou multi-sites avec coordination de grands chantiers."
            ],
            "Dipl√¥mes / Certifications valorisantes": [
                "- Certifications suppl√©mentaires comme LEED pour le d√©veloppement durable en BTP."
            ],
            "Comp√©tences compl√©mentaires": [
                "- Comp√©tences en BIM (Building Information Modeling) ou en gestion environnementale des chantiers."
            ]
        },
        "Sourcing et march√©": {
            "Entreprises o√π trouver ce profil": [
                "- Sugg√©rer des entreprises concurrentes dans le BTP comme Bouygues, Vinci ou Eiffage."
            ],
            "Synonymes / intitul√©s proches": [
                "- Titres alternatifs comme 'Chef de Chantier', 'Ing√©nieur Travaux', 'Conducteur de Travaux BTP'."
            ],
            "Canaux √† utiliser": [
                "- Proposer LinkedIn pour profils BTP, jobboards sp√©cialis√©s (Batiweb), cooptation sur chantiers, r√©seaux professionnels du BTP."
            ]
        }
    }

    section_advice = advice_db.get(section_title, {})
    field_advice = section_advice.get(field_title, [])
    if field_advice:
        return random.choice(field_advice)
    else:
        return "Pas de conseil disponible."

# -------------------- Filtre --------------------
def filter_briefs(briefs, month, recruteur, brief_type, manager, affectation, nom_affectation):
    """Filtre les briefs selon les crit√®res donn√©s."""
    filtered = {}
    for name, data in briefs.items():
        match = True
        try:
            if month and month != "" and datetime.strptime(data.get("date_brief", datetime.today().strftime("%Y-%m-%d")), "%Y-%m-%d").strftime("%m") != month:
                match = False
            if recruteur and recruteur != "" and recruteur.lower() not in data.get("recruteur", "").lower():
                match = False
            if brief_type and brief_type != "" and data.get("brief_type", "") != brief_type:
                match = False
            if manager and manager != "" and manager.lower() not in data.get("manager_nom", "").lower():
                match = False
            if affectation and affectation != "" and data.get("affectation_type", "") != affectation:
                match = False
            if nom_affectation and nom_affectation != "" and nom_affectation.lower() not in data.get("affectation_nom", "").lower():
                match = False
            if match:
                filtered[name] = data
        except ValueError:  # G√®re les dates non valides
            continue
    return filtered

# -------------------- Exportation PDF --------------------
def export_brief_pdf():
    """Exporte un brief au format PDF."""
    if not PDF_AVAILABLE or not st.session_state.current_brief_name:
        return None
    
    brief_data = st.session_state.saved_briefs.get(st.session_state.current_brief_name, {})
    buffer = io.BytesIO()
    doc = SimpleDocTemplate(buffer, pagesize=A4)
    styles = getSampleStyleSheet()
    story = []

    story.append(Paragraph("üìã Brief Recrutement", styles['Heading1']))
    story.append(Spacer(1, 12))

    # Section 1: Identit√© du poste
    story.append(Paragraph("1. Identit√© du poste", styles['Heading2']))
    story.append(Paragraph(f"Intitul√©: {brief_data.get('poste_intitule', '')}", styles['Normal']))
    story.append(Paragraph(f"Service: {brief_data.get('affectation_nom', '')}", styles['Normal']))
    story.append(Paragraph(f"Niveau Hi√©rarchique: {brief_data.get('niveau_hierarchique', 'N/A')}", styles['Normal']))  # Placeholder
    story.append(Paragraph(f"Type de Contrat: {brief_data.get('affectation_type', 'N/A')}", styles['Normal']))  # Adjust if needed
    story.append(Paragraph(f"Localisation: {brief_data.get('rattachement', '')}", styles['Normal']))
    story.append(Paragraph(f"Budget Salaire: {brief_data.get('budget', '')}", styles['Normal']))
    story.append(Paragraph(f"Date Prise de Poste: {brief_data.get('date_brief', '')}", styles['Normal']))
    story.append(Spacer(1, 12))

    # Section 2: Contexte & Enjeux
    story.append(Paragraph("2. Contexte & Enjeux", styles['Heading2']))
    context_text = f"{brief_data.get('raison_ouverture', '')} {brief_data.get('impact_strategique', '')}"
    story.append(Paragraph(context_text, styles['Normal']))
    story.append(Spacer(1, 12))

    # Section 3: Exigences
    story.append(Paragraph("3. Exigences", styles['Heading2']))
    story.append(Paragraph(f"Exp√©rience: {brief_data.get('must_have_experience', '')}", styles['Normal']))
    story.append(Paragraph(f"Dipl√¥mes: {brief_data.get('must_have_diplomes', '')}", styles['Normal']))
    story.append(Paragraph(f"Comp√©tences: {brief_data.get('must_have_competences', '')}", styles['Normal']))
    story.append(Paragraph(f"Soft Skills: {brief_data.get('must_have_softskills', '')}", styles['Normal']))
    story.append(Spacer(1, 12))

    # Section 4: Matrice KSA
    story.append(Paragraph("4. Matrice KSA", styles['Heading2']))
    ksa_matrix = brief_data.get("ksa_matrix")
    if isinstance(ksa_matrix, pd.DataFrame) and not ksa_matrix.empty:
        for _, row in ksa_matrix.iterrows():
            question = row.get('Question pour l\'entretien', '')
            ksa_text = f"- {row.get('Rubrique', '')}: {row.get('Crit√®re', '')} (Question: {question}, √âval: {row.get('√âvaluation (1-5)', '')})"
            story.append(Paragraph(ksa_text, styles['Normal']))
    else:
        story.append(Paragraph("Aucune donn√©e KSA disponible.", styles['Normal']))
    story.append(Spacer(1, 12))

    # Section 5: Strat√©gie Recrutement
    story.append(Paragraph("5. Strat√©gie Recrutement", styles['Heading2']))
    story.append(Paragraph(f"Canaux: {brief_data.get('canaux_profil', '')}", styles['Normal']))
    story.append(Paragraph(f"Crit√®res d'exclusion: {brief_data.get('criteres_exclusion', '')}", styles['Normal']))
    story.append(Spacer(1, 12))

    # Section 6: Notes du Manager
    story.append(Paragraph("6. Notes du Manager", styles['Heading2']))
    story.append(Paragraph(brief_data.get("manager_notes", ""), styles['Normal']))

    doc.build(story)
    buffer.seek(0)
    return buffer.getvalue()

# -------------------- Exportation Word --------------------
def export_brief_word():
    """Exporte un brief au format Word."""
    if not WORD_AVAILABLE or not st.session_state.current_brief_name:
        return None
    
    brief_data = st.session_state.saved_briefs.get(st.session_state.current_brief_name, {})
    doc = Document()
    doc.add_heading("üìã Brief Recrutement", 0)

    # Section 1: Identit√© du poste
    doc.add_heading("1. Identit√© du poste", level=1)
    doc.add_paragraph(f"Intitul√©: {brief_data.get('poste_intitule', '')}")
    doc.add_paragraph(f"Service: {brief_data.get('affectation_nom', '')}")
    doc.add_paragraph(f"Niveau Hi√©rarchique: {brief_data.get('niveau_hierarchique', 'N/A')}")  # Placeholder
    doc.add_paragraph(f"Type de Contrat: {brief_data.get('affectation_type', 'N/A')}")  # Adjust if needed
    doc.add_paragraph(f"Localisation: {brief_data.get('rattachement', '')}")
    doc.add_paragraph(f"Budget Salaire: {brief_data.get('budget', '')}")
    doc.add_paragraph(f"Date Prise de Poste: {brief_data.get('date_brief', '')}")

    # Section 2: Contexte & Enjeux
    doc.add_heading("2. Contexte & Enjeux", level=1)
    doc.add_paragraph(f"{brief_data.get('raison_ouverture', '')} {brief_data.get('impact_strategique', '')}")

    # Section 3: Exigences
    doc.add_heading("3. Exigences", level=1)
    doc.add_paragraph(f"Exp√©rience: {brief_data.get('must_have_experience', '')}")
    doc.add_paragraph(f"Dipl√¥mes: {brief_data.get('must_have_diplomes', '')}")
    doc.add_paragraph(f"Comp√©tences: {brief_data.get('must_have_competences', '')}")
    doc.add_paragraph(f"Soft Skills: {brief_data.get('must_have_softskills', '')}")

    # Section 4: Matrice KSA
    doc.add_heading("4. Matrice KSA", level=1)
    ksa_matrix = brief_data.get("ksa_matrix")
    if isinstance(ksa_matrix, pd.DataFrame) and not ksa_matrix.empty:
        for _, row in ksa_matrix.iterrows():
            question_text = row.get('Question pour l\'entretien', '')
            doc.add_paragraph(f"- {row.get('Rubrique', '')}: {row.get('Crit√®re', '')} (Question: {question_text}, √âval: {row.get('√âvaluation (1-5)', '')})")
    else:
        doc.add_paragraph("Aucune donn√©e KSA disponible.")

    # Section 5: Strat√©gie Recrutement
    doc.add_heading("5. Strat√©gie Recrutement", level=1)
    doc.add_paragraph(f"Canaux: {brief_data.get('canaux_profil', '')}")
    doc.add_paragraph(f"Crit√®res d'exclusion: {brief_data.get('criteres_exclusion', '')}")

    # Section 6: Notes du Manager
    doc.add_heading("6. Notes du Manager", level=1)
    doc.add_paragraph(brief_data.get("manager_notes", ""))

    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer.getvalue()

# -------------------- Gestion de la Biblioth√®que de fiches de poste --------------------
LIBRARY_FILE = "job_library.json"

def load_library():
    """Charge les fiches de poste depuis le fichier de la biblioth√®que."""
    if os.path.exists(LIBRARY_FILE):
        with open(LIBRARY_FILE, "r", encoding="utf-8") as f:
            return json.load(f)
    return []

def save_library(library_data):
    """Sauvegarde les fiches de poste dans le fichier de la biblioth√®que."""
    with open(LIBRARY_FILE, "w", encoding="utf-8") as f:
        json.dump(library_data, f, indent=4, ensure_ascii=False)

# -------------------- Pr√©-r√©daction IA avec DeepSeek --------------------
def get_ai_pre_redaction(fiche_data):
    """G√©n√®re une pr√©-r√©daction synth√©tique avec DeepSeek API via OpenAI client."""
    try:
        from openai import OpenAI
    except ImportError:
        raise ImportError("Le module 'openai' n'est pas install√©. Veuillez l'installer avec 'pip install openai'.")

    api_key = st.secrets.get("DEEPSEEK_API_KEY")
    if not api_key:
        raise ValueError("Cl√© API DeepSeek non trouv√©e dans st.secrets")

    client = OpenAI(
        api_key=api_key,
        base_url="https://api.deepseek.com/v1"
    )

    prompt = (
        f"Synth√©tise les informations de cette fiche de poste en une version courte et concise. "
        f"Modifie uniquement les sections suivantes :\n"
        f"- Mission globale : une phrase courte.\n"
        f"- T√¢ches principales : 5-6 missions courtes en bullet points.\n"
        f"- Must have : liste des exigences essentielles compl√®tes en bullet points.\n"
        f"- Nice to have : liste des exigences optionnelles compl√®tes en bullet points.\n"
        f"Ne touche pas aux autres sections. Utilise un format markdown clair pour chaque section.\n"
        f"Fiche de poste :\n{fiche_data}"
    )

    response = client.chat.completions.create(
        model="deepseek-chat",
        messages=[{"role": "user", "content": prompt}],
        temperature=0.5,
        max_tokens=500
    )

    return response.choices[0].message.content

# -------------------- G√©n√©ration de question IA avec DeepSeek --------------------
def generate_ai_question(prompt, concise=False):
    """G√©n√®re une question d'entretien et une r√©ponse exemple via l'API DeepSeek."""
    try:
        from openai import OpenAI
    except ImportError:
        raise ImportError("Le module 'openai' n'est pas install√©. Veuillez l'installer avec 'pip install openai'.")

    api_key = st.secrets.get("DEEPSEEK_API_KEY")
    if not api_key:
        raise ValueError("Cl√© API DeepSeek non trouv√©e dans st.secrets")

    client = OpenAI(
        api_key=api_key,
        base_url="https://api.deepseek.com/v1"
    )

    max_tokens = 700

    context = "recruitment for a KSA (Knowledge, Skills, Abilities) matrix"
    question_type = "technical"
    skill = prompt
    role = "candidate"
    
    if "une question" in prompt and "pour √©valuer" in prompt and "par" in prompt:
        parts = prompt.split("une question")
        if len(parts) > 1:
            question_type_part = parts[1].split("pour")[0].strip().lower()
            if "g√©n√©rale" in question_type_part:
                question_type = "general"
            elif "comportementale" in question_type_part:
                question_type = "behavioral"
            elif "situationnelle" in question_type_part:
                question_type = "situational"
            elif "technique" in question_type_part:
                question_type = "technical"
            
            skill_part = prompt.split("√©valuer")
            if len(skill_part) > 1:
                skill = skill_part[1].split("par")[0].strip()
            
            role_part = prompt.split("par")
            if len(role_part) > 1:
                role = role_part[1].strip()

    if question_type == "behavioral":
        context += ", using the STAR method (Situation, Task, Action, Result)"
    elif question_type == "situational":
        context += ", presenting a hypothetical scenario"
    elif question_type == "general":
        context += ", focusing on overall experience"

    full_prompt = (
        f"Dans le contexte de {context}, g√©n√®re une {question_type} question d'entretien pour √©valuer : {skill} by a {role}. "
        f"Adapte la question au domaine sp√©cifi√© (ex. recrutement ou BTP) si applicable. "
        f"Retourne une r√©ponse exemple pertinente. "
        f"Assure-toi que la r√©ponse corresponde au type de question (ex. STAR pour comportementale, sc√©nario pour situationnelle). "
        f"Utilise uniquement ce format :\n"
        f"Question: [votre question]\n"
        f"R√©ponse: [exemple de r√©ponse]"
    )
    
    if concise:
        full_prompt = (
            f"G√©n√®re une question d'entretien et une r√©ponse tr√®s concise et directe pour √©valuer le crit√®re : '{skill}'. "
            f"La question doit √™tre {question_type} et la r√©ponse ne doit pas d√©passer 50 mots. "
            f"Format : 'Question: [votre question]\nR√©ponse: [r√©ponse concise]'"
        )
        max_tokens = 150

    response = client.chat.completions.create(
        model="deepseek-chat",
        messages=[{"role": "user", "content": full_prompt}],
        temperature=0.7,
        max_tokens=max_tokens
    )

    result = response.choices[0].message.content.strip()
    if result.startswith("- Question:"):
        result = result.replace("- Question:", "Question:", 1)
    if "R√©ponse: R√©ponse:" in result:
        result = result.replace("R√©ponse: R√©ponse:", "R√©ponse:")
    
    return result

# -------------------- Test de connexion DeepSeek --------------------
def test_deepseek_connection():
    """Teste la connexion √† l'API DeepSeek."""
    try:
        from openai import OpenAI
        api_key = st.secrets.get("DEEPSEEK_API_KEY")
        if not api_key:
            st.error("Cl√© API DeepSeek non trouv√©e dans st.secrets")
            return False
        client = OpenAI(
            api_key=api_key,
            base_url="https://api.deepseek.com/v1"
        )
        client.chat.completions.create(
            model="deepseek-chat",
            messages=[{"role": "user", "content": "Test de connexion"}],
            max_tokens=1
        )
        st.success("‚úÖ Connexion √† DeepSeek r√©ussie !")
        return True
    except Exception as e:
        st.error(f"‚ùå Erreur de connexion √† DeepSeek : {e}")
        return False

# -------------------- G√©n√©ration de contenu avec DeepSeek --------------------
def deepseek_generate(prompt, max_tokens=2000, temperature=0.7):
    """G√©n√®re du contenu via l'API DeepSeek."""
    try:
        from openai import OpenAI
    except ImportError:
        raise ImportError("Le module 'openai' n'est pas install√©. Veuillez l'installer avec 'pip install openai'.")

    api_key = st.secrets.get("DEEPSEEK_API_KEY")
    if not api_key:
        raise ValueError("Cl√© API DeepSeek non trouv√©e dans st.secrets")

    client = OpenAI(
        api_key=api_key,
        base_url="https://api.deepseek.com/v1"
    )

    response = client.chat.completions.create(
        model="deepseek-chat",
        messages=[{"role": "user", "content": prompt}],
        temperature=temperature,
        max_tokens=max_tokens
    )

    return response.choices[0].message.content

# -------------------- Exemples contextuels --------------------
def get_example_for_field(section_title, field_title):
    """Retourne un exemple contextuel pour un champ donn√©, adapt√© au BTP."""
    # Pour chaque champ nous construisons une petite banque d'exemples
    # et retournons jusqu'√† 10 exemples al√©atoires (s√©par√©s par une ligne).
    import random

    base_examples = {
        "Contexte du poste": {
            "Raison de l'ouverture": [
                "Remplacement d‚Äôun d√©part en retraite sur un chantier majeur",
                "Ouverture d‚Äôun nouveau site suite √† un gain d‚Äôappel d‚Äôoffres",
                "Renforcement temporaire pour pic d‚Äôactivit√© saisonnier",
                "Cr√©ation d‚Äôun poste pour piloter un projet de r√©novation urbaine",
                "Remplacement d‚Äôun cong√© maternit√© pour 12 mois",
                "Augmentation de l‚Äô√©quipe suite extension du p√©rim√®tre travaux",
                "Remplacement d‚Äôun collaborateur promu sur un autre site",
                "Besoin de comp√©tences sp√©cifiques pour un chantier industriel",
                "Remplacement suite √† un d√©part impr√©vu (d√©mission)",
                "Poste cr√©√© pour am√©liorer le pilotage s√©curit√© chantier"
            ],
            "Impact strat√©gique": [
                "Assurer la gestion des projets BTP strat√©giques de l‚Äôentreprise",
                "R√©duire les retards de livraison sur les chantiers cl√©s",
                "Am√©liorer la conformit√© s√©curit√© et r√©duire les incidents",
                "Optimiser les co√ªts de sous-traitance et ressources internes",
                "Renforcer la capacit√© de pilotage de projets multi-sites",
                "Acc√©l√©rer la mise en service d‚Äôun projet d‚Äôinfrastructure critique",
                "Permettre la mont√©e en comp√©tence d‚Äôune √©quipe locale",
                "Soutenir la strat√©gie de croissance sur le segment infrastructures",
                "Garantir la qualit√© technique sur des lots sensibles",
                "Contribuer √† l‚Äôobtention d‚Äôune certification qualit√© environnementale"
            ],
            "T√¢ches principales": [
                "Gestion de chantier complexe, coordination d‚Äô√©quipe sur site, suivi des normes de s√©curit√© BTP",
                "Planification des phases travaux et coordination des sous-traitants",
                "Suivi budg√©taire et reporting hebdomadaire au manager de projet",
                "Contr√¥le qualit√© des ouvrages et respect des cahiers des charges",
                "Animation des r√©unions de chantier et gestion des al√©as",
                "Pilotage des livrables techniques et validation des plans",
                "Supervision s√©curit√© et application des proc√©dures HSE",
                "Gestion des commandes mat√©rielles et approvisionnement",
                "R√©daction des comptes-rendus et reporting client",
                "Coordination interface entre bureaux d‚Äô√©tudes et √©quipes op√©rationnelles"
            ]
        },
        "Must-have (Indispensables)": {
            "Exp√©rience": [
                "5 ans d‚Äôexp√©rience dans le secteur BTP sur des chantiers similaires",
                "3 √† 7 ans dans la gestion de chantiers industriels",
                "Exp√©rience significative en pilotage de projets en extension",
                "Ant√©c√©dents sur chantiers d‚Äôinfrastructures et grands projets",
                "Exp√©rience en coordination multi-stakeholders",
                "Preuve de livraison de projets respectant d√©lais et budget",
                "Exp√©rience en management d‚Äô√©quipes de 10+ personnes",
                "Historique de management de sous-traitants",
                "Exp√©rience en conduite de travaux et lecture de plans techniques",
                "Exp√©rience en suivi de travaux HSE et conformit√©"
            ],
            "Connaissances / Dipl√¥mes / Certifications": [
                "Dipl√¥me en g√©nie civil, certification CACES pour engins de chantier",
                "BTS/DUT ou licence en BTP ou g√©nie civil",
                "Certifications s√©curit√© chantiers (HSE) requises",
                "Formation en conduite de travaux confirm√©e",
                "Certification en management de projets (Prince2/PMI) appr√©ci√©e",
                "Attestation CACES pour engins de terrassement",
                "Formation en r√©glementation du b√¢timent (Code du travail)",
                "Dipl√¥me d‚Äôing√©nieur sp√©cialis√© BTP souhaitable",
                "Certificat en gestion financi√®re de projet",
                "Qualification professionnelle reconnue dans le secteur"
            ],
            "Comp√©tences / Outils": [
                "Ma√Ætrise d'AutoCAD et de la gestion de projet BTP",
                "Bon niveau sur MS Project / Gantt / planning chantier",
                "Capacit√© √† lire et interpr√©ter plans d‚Äôex√©cution",
                "Connaissance des process achats et commandes chantier",
                "Exp√©rience avec outils de suivi HSE et s√©curit√©",
                "Comp√©tences en estimation et chiffrage des lots",
                "Ma√Ætrise d‚Äôoutils collaboratifs (Teams, Sharepoint)",
                "Capacit√© √† piloter sous-traitance et relation fournisseurs",
                "Ma√Ætrise de la gestion documentaire chantier",
                "Savoir-faire en coordination technique multi-lots"
            ],
            "Soft skills / aptitudes comportementales": [
                "Leadership sur terrain et rigueur en s√©curit√©",
                "Capacit√© √† g√©rer le stress et les priorit√©s",
                "Communication claire avec √©quipes et clients",
                "Orientation r√©sultats et sens du service",
                "Esprit d‚Äô√©quipe et mentorat des juniors",
                "Autonomie et prise d‚Äôinitiative sur site",
                "Rigueur administrative et reporting r√©gulier",
                "Adaptabilit√© face aux impr√©vus chantier",
                "Sens de l‚Äôorganisation et gestion du temps",
                "Capacit√© d√©cisionnelle lors d‚Äôal√©as techniques"
            ]
        },
        "Nice-to-have (Atouts)": {
            "Exp√©rience additionnelle": [
                "Projets internationaux de BTP ou multi-sites",
                "Exp√©rience en chantiers en zone urbaine dense",
                "Participation √† projets de renovation patrimoniale",
                "Exp√©rience sur projets d‚Äô√©nergie ou infrastructures",
                "Gestion de projet √©cologique / chantier bas-carbone",
                "Projets en BIM / mod√©lisation num√©rique",
                "Exp√©rience en march√©s publics et appels d‚Äôoffres",
                "Connaissance projets PPP ou financement priv√©-public",
                "Travail sur sites industriels sensibles",
                "Coordination multi-pays / √©quipes internationales"
            ],
            "Dipl√¥mes / Certifications valorisantes": [
                "Certification LEED pour le BTP durable",
                "Formation compl√©mentaire en BIM ou num√©rique",
                "Certificat en management QSE",
                "Titre d‚Äôing√©nieur sp√©cialis√© en structures",
                "Certification management de projet avanc√©",
                "Formation en coordination s√©curit√© avanc√©e",
                "Attestation en gestion des risques environnementaux",
                "Dipl√¥me compl√©mentaire en √©conomique de la construction",
                "Certification en gestion contractuelle",
                "Qualification en conduite d‚Äô√©quipes pluridisciplinaires"
            ],
            "Comp√©tences compl√©mentaires": [
                "Connaissance en BIM pour mod√©lisation de chantiers",
                "Comp√©tences en estimation co√ªts et optimisation",
                "Ma√Ætrise des normes environnementales BTP",
                "Exp√©rience en digitalisation des processus chantier",
                "Savoir utiliser les logiciels de simulation chantier",
                "Capacit√© √† conduire audits qualit√©",
                "Connaissance en supervision d‚Äôessais mat√©riaux",
                "Comp√©tences en relation client et contrats",
                "Id√©es innovantes pour r√©duction des co√ªts",
                "Exp√©rience en lean construction"
            ]
        },
        "Conditions et contraintes": {
            "Localisation": [
                "Chantier principal √† Paris avec 20% de t√©l√©travail",
                "Poste bas√© sur site (d√©placements fr√©quents inter-chantiers)",
                "Localisation r√©gionale avec h√©bergement temporaire possible",
                "Travail principalement sur site, quelques r√©unions au si√®ge",
                "D√©placements nationaux ponctuels pr√©vus",
                "Mission sur zone industrielle avec contraintes horaires",
                "Pr√©sence requise sur site d√®s 7h30 certains jours",
                "Site √©loign√© avec voiture de service fournie",
                "Mission avec rotation entre 2 chantiers principaux",
                "Localisation multi-sites selon planning trimestriel"
            ],
            "Budget recrutement": [
                "Salaire entre 40k et 50k‚Ç¨ + primes",
                "R√©mun√©ration selon profil + 13√®me mois possible",
                "Package mobilit√© et prime d‚Äôexpatriation selon cas",
                "Budget fixe avec fourchette 38k-45k‚Ç¨",
                "Salaire comp√©titif sur march√© local",
                "Contrat CDI avec prime d‚Äôobjectif annuelle",
                "R√©mun√©ration n√©gociable selon exp√©rience",
                "Forfait jour ou heures selon convention collective",
                "Avantages en nature (v√©hicule, t√©l√©phone) inclus",
                "Indemnit√©s de chantier et panier repas pr√©vues"
            ]
        },
        "Sourcing et march√©": {
            "Entreprises o√π trouver ce profil": [
                "Vinci, Bouygues, Eiffage",
                "Entreprises locales de construction sp√©cialis√©es",
                "Sous-traitants de r√©seaux et TP",
                "Soci√©t√©s d‚Äôing√©nierie structurelle",
                "PME r√©gionales actives sur grands chantiers",
                "Groupes internationaux de BTP pr√©sents en France",
                "Prestataires de maintenance industrielle",
                "Ateliers de pr√©fabrication et mod√©lisation BIM",
                "Int√©grateurs de solutions chantiers connect√©s",
                "Entreprises sp√©cialis√©es en infrastructures routi√®res"
            ],
            "Synonymes / intitul√©s proches": [
                "Conducteur de Travaux, Ing√©nieur Chantier",
                "Chef de chantier, Responsable travaux",
                "Coordinateur de r√©alisation, Pilotage de travaux",
                "Chef de projet chantier junior/senior",
                "Responsable de lot, Charg√© d‚Äôaffaires technique",
                "Ing√©nieur d‚Äôex√©cution, Superviseur terrain",
                "Coordinateur HSE chantier",
                "Responsable √©tudes d‚Äôex√©cution",
                "Chef de groupe travaux",
                "Gestionnaire op√©rations chantier"
            ],
            "Canaux √† utiliser": [
                "LinkedIn pour profils BTP, jobboards comme Batiweb",
                "Candidatures via r√©seaux professionnels et salons",
                "Approche directe (chasse) sur LinkedIn",
                "Partenariat √©coles d‚Äôing√©nieurs locales",
                "Diffusion sur jobboards sp√©cialis√©s BTP",
                "Utiliser CVth√®ques et viviers internes",
                "Annonces sur journaux sectoriels",
                "Groupes professionnels et forums m√©tiers",
                "Recommandations internes (employee referral)",
                "Chasse via cabinets sp√©cialis√©s" 
            ]
        },
        "Profils pertinents": {
            "Lien profil 1": [
                "https://linkedin.com/in/exemple-btp",
                "https://linkedin.com/in/exemple-btp-2",
                "https://linkedin.com/in/exemple-btp-3"
            ],
            "Lien profil 2": [
                "https://linkedin.com/in/exemple-btp-4",
                "https://linkedin.com/in/exemple-btp-5",
                "https://linkedin.com/in/exemple-btp-6"
            ],
            "Lien profil 3": [
                "https://linkedin.com/in/exemple-btp-7",
                "https://linkedin.com/in/exemple-btp-8",
                "https://linkedin.com/in/exemple-btp-9"
            ]
        },
        "Notes libres": {
            "Points √† discuter ou √† clarifier avec le manager": [
                "Clarifier les d√©lais du projet",
                "Valider la disponibilit√© des ressources internes",
                "Pr√©ciser le p√©rim√®tre exact des responsabilit√©s",
                "Discuter des contraintes HSE sp√©cifiques",
                "Valider la plage budg√©taire disponible",
                "S‚Äôassurer des conditions d‚Äôacc√®s site et s√©curit√©",
                "Pr√©ciser la dur√©e estim√©e du chantier",
                "V√©rifier les interfaces avec autres prestataires",
                "Confirmer les priorit√©s clients sur les lots",
                "Demander les √©l√©ments manquants du dossier technique"
            ],
            "Case libre": [
                "Note sur les priorit√©s du chantier",
                "Remarques g√©n√©rales √† partager avec l‚Äô√©quipe",
                "Points de vigilance pour la phase d‚Äô√©tudes",
                "Id√©es d‚Äôam√©lioration du processus recrutement",
                "Observations terrain report√©es par le manager",
                "Notes administratives et contacts cl√©s",
                "R√©sum√©s d‚Äô√©changes avec client/ma√Ætre d‚Äôouvrage",
                "Consignes particuli√®res pour les habilitations",
                "Informations utiles pour la prise de poste",
                "Bref historique du projet et enjeux"
            ]
        }
    }

    pool = base_examples.get(section_title, {}).get(field_title)
    if not pool:
        return "Exemple non disponible"

    # Toujours retourner exactement 2 exemples (ou moins si indisponible)
    n = 2
    if isinstance(pool, list):
        if len(pool) <= n:
            chosen = pool
        else:
            chosen = random.sample(pool, n)
        return "\n".join([f"- {s}" for s in chosen])

    return str(pool)


def delete_brief_from_gsheet(brief_name: str) -> bool:
    """Supprime une ligne correspondant √† brief_name dans la feuille Google Sheets "Briefs".

    Retourne True si la suppression r√©ussit ou si la feuille n'est pas disponible.
    """
    try:
        worksheet = get_briefs_gsheet_client()
        if not worksheet:
            # Pas de client Sheets : consid√©rer comme succ√®s silencieux
            return True

        # R√©cup√©rer toutes les valeurs de la colonne A (BRIEF_NAME)
        try:
            col_values = worksheet.col_values(1)
        except Exception:
            # Si col_values √©choue, retomber sur find (ancienne m√©thode)
            try:
                cell = worksheet.find(brief_name, in_column=1, case_sensitive=True)
                if cell:
                    worksheet.delete_row(cell.row)
                    st.info(f"‚úÖ Ligne supprim√©e dans Google Sheets (ligne {cell.row}).")
                    return True
                return False
            except Exception:
                return False

        # Normalisation simple: unicode normalize, strip, lower
        import unicodedata
        def _norm(s):
            if s is None:
                return ""
            s2 = str(s)
            try:
                s2 = unicodedata.normalize('NFKC', s2)
            except Exception:
                pass
            return s2.strip().lower()

        target = _norm(brief_name)

        # 1) recherche exacte (normalis√©e)
        rows_to_delete = [i+1 for i, v in enumerate(col_values) if _norm(v) == target]

        # 2) si aucun match exact, essayer une correspondance "contains" (moins stricte)
        if not rows_to_delete:
            rows_to_delete = [i+1 for i, v in enumerate(col_values) if target and target in _norm(v)]

        if not rows_to_delete:
            # rien √† supprimer: informer et laisser caller d√©cider
            st.warning(f"Aucune ligne trouv√©e dans Google Sheets pour: '{brief_name}' (col A).")
            return False

        deleted = 0
        # Supprimer de la plus grande ligne vers la plus petite pour garder les indices valides
        for row in sorted(rows_to_delete, reverse=True):
            try:
                worksheet.delete_row(row)
                deleted += 1
            except Exception:
                # si une suppression particuli√®re √©choue, continuer
                continue

        if deleted:
            st.info(f"‚úÖ {deleted} ligne(s) supprim√©e(s) dans Google Sheets.")
            return True
        else:
            st.warning("Aucune suppression effective r√©alis√©e (erreur lors des op√©rations delete_row).")
            return False
    except Exception:
        return False

# -------------------- G√©n√©ration de nom de brief automatique --------------------
def generate_automatic_brief_name(poste: str = None, manager: str = None, date_obj=None):
    """
    G√©n√®re un nom unique de brief: POSTE_MANAGER_YYYYMMDD[_n].
    Param√®tres optionnels (fallback sur st.session_state pour compatibilit√©).
    """
    from datetime import datetime, date
    # Fallback session_state
    if poste is None:
        poste = st.session_state.get("poste_intitule", "Poste")
    if manager is None:
        manager = st.session_state.get("manager_nom", "Manager")
    if date_obj is None:
        date_obj = st.session_state.get("date_brief", datetime.today())
    # Normalisation date
    if isinstance(date_obj, str):
        for fmt in ("%Y-%m-%d", "%d/%m/%Y"):
            try:
                date_obj = datetime.strptime(date_obj, fmt).date()
                break
            except Exception:
                continue
    if isinstance(date_obj, datetime):
        date_obj = date_obj.date()
    if not isinstance(date_obj, date):
        date_obj = datetime.today().date()
    date_str = date_obj.strftime("%Y%m%d")

    # Nettoyage simple
    def slugify(val):
        return "".join(c for c in val.strip().replace(" ", "_") if c.isalnum() or c in ("_", "-"))[:40] or "X"
    poste_slug = slugify(poste)
    manager_slug = slugify(manager)

    base_name = f"{poste_slug}_{manager_slug}_{date_str}"
    saved_briefs = st.session_state.get("saved_briefs", {})
    brief_name = base_name
    i = 1
    while brief_name in saved_briefs:
        brief_name = f"{base_name}_{i}"
        i += 1
    return brief_name

# -------------------- Conversion de colonne --------------------
def col_to_letter(col_index):
    """Convertit l'index de colonne (base 1) en lettre Excel (e.g., 1 -> A, 27 -> AA)"""
    letter = ''
    while col_index > 0:
        col_index, remainder = divmod(col_index - 1, 26)
        letter = chr(65 + remainder) + letter
    return letter

def export_brief_pdf_pretty(brief_name: str, brief_data: dict, ksa_df):
    """
    G√©n√®re un PDF structur√© et lisible (brief + strat√©gie + matrice KSA).
    Retourne un buffer BytesIO ou None si lib absente.
    """
    if not PDF_PRETTY_AVAILABLE:
        return None
    buf = BytesIO()
    c = canvas.Canvas(buf, pagesize=A4)
    W, H = A4
    margin_x = 20 * mm
    y = H - 25 * mm

    def line(txt, size=9, leading=12, bold=False, max_len=150):
        nonlocal y
        if y < 30 * mm:
            c.showPage()
            y = H - 25 * mm
        font = "Helvetica-Bold" if bold else "Helvetica"
        c.setFont(font, size)
        for seg in str(txt).split("\n"):
            c.drawString(margin_x, y, seg[:max_len])
            y -= leading

    # Dessiner le logo TGCC centr√© en haut si disponible
    possible_logo_paths = [
        os.path.join(os.path.dirname(__file__), '..', 'tgcc.png'),
        os.path.join(os.getcwd(), 'tgcc.png'),
        os.path.join(os.path.dirname(__file__), 'tgcc.png')
    ]
    logo_path = next((p for p in possible_logo_paths if os.path.exists(p)), None)
    if logo_path:
        try:
            logo_w = 40 * mm
            x = (W - logo_w) / 2
            c.drawImage(logo_path, x, H - 20 * mm, width=logo_w, preserveAspectRatio=True, mask='auto')
            y = H - 35 * mm
        except Exception:
            y = H - 25 * mm
    else:
        y = H - 25 * mm

    line(f"Brief: {brief_name}", size=14, leading=18, bold=True)
    line(f"Poste: {brief_data.get('POSTE_INTITULE', brief_data.get('poste_intitule',''))}", bold=True)
    line(f"Manager: {brief_data.get('MANAGER_NOM', brief_data.get('manager_nom',''))}")
    line(f"Affectation: {brief_data.get('AFFECTATION_NOM','')} ({brief_data.get('AFFECTATION_TYPE','')})")
    line(f"Date: {brief_data.get('DATE_BRIEF', brief_data.get('date_brief',''))}")
    line("")

    line("1. Contexte & Exigences", bold=True)
    ordered_keys = [
        ("Raison de l'ouverture","RAISON_OUVERTURE"),
        ("Impact strat√©gique","IMPACT_STRATEGIQUE"),
        ("T√¢ches principales","TACHES_PRINCIPALES"),
        ("Must Have - Exp√©rience","MUST_HAVE_EXPERIENCE"),
        ("Must Have - Dipl√¥mes","MUST_HAVE_DIPLOMES"),
        ("Must Have - Comp√©tences","MUST_HAVE_COMPETENCES"),
        ("Must Have - Soft skills","MUST_HAVE_SOFTSKILLS"),
        ("Nice to Have - Exp√©rience","NICE_TO_HAVE_EXPERIENCE"),
        ("Nice to Have - Dipl√¥mes","NICE_TO_HAVE_DIPLOMES"),
        ("Nice to Have - Comp√©tences","NICE_TO_HAVE_COMPETENCES"),
        ("Localisation","RATTACHEMENT"),
        ("Budget","BUDGET"),
        ("Entreprises cibles","ENTREPRISES_PROFIL"),
        ("Synonymes","SYNONYMES_POSTE"),
        ("Canaux sugg√©r√©s","CANAUX_PROFIL"),
        ("Lien profil 1","LIEN_PROFIL_1"),
        ("Lien profil 2","LIEN_PROFIL_2"),
        ("Lien profil 3","LIEN_PROFIL_3"),
        ("Notes libres","NOTES_LIBRES")
    ]
    for label, key in ordered_keys:
        val = brief_data.get(key, "")
        if val:
            line(f"- {label}: {val}")

    line("")
    line("2. Strat√©gie & Processus", bold=True)
    for label, key in [
        ("Strat√©gie de sourcing","STRATEGIE_SOURCING"),
        ("Crit√®res d'exclusion","CRITERES_EXCLUSION"),
        ("Processus d'√©valuation","PROCESSUS_EVALUATION"),
        ("Notes manager","MANAGER_NOTES")
    ]:
        v = brief_data.get(key, "")
        if v:
            line(f"- {label}: {v}")

    if ksa_df is not None and hasattr(ksa_df, "empty") and not ksa_df.empty:
        line("")
        line("3. Matrice KSA (R√©sum√©)", bold=True)
        for _, row in ksa_df.iterrows():
            rub = row.get("Rubrique","")
            crit = row.get("Crit√®re","")
            tq = row.get("Type de question","")
            q = row.get("Question pour l'entretien", row.get("Cible / Standard attendu",""))
            ev = row.get("√âvaluation (1-5)", row.get("√âvaluation (1-5)",""))
            line(f"[{rub}/{tq}] {crit} (Cible: {ev}/5)")
            if q:
                line(f"    Q: {str(q)[:110]}")
        if "√âvaluation (1-5)" in ksa_df.columns:
            try:
                vals = ksa_df["√âvaluation (1-5)"].dropna().astype(float)
                if len(vals) > 0:
                    avg = round(vals.mean(), 2)
                    line(f"Score cible moyen: {avg}/5", bold=True)
            except:
                pass

    c.showPage()
    c.save()
    buf.seek(0)
    return buf

import pandas as pd
import streamlit as st
import json

def get_brief_value(brief_dict: dict, key: str, default: str = ""):
    if not isinstance(brief_dict, dict):
        return default
    if key.startswith("profil_link_"):
        suf = key.split("_")[-1]
        candidates = [f"LIEN_PROFIL_{suf}", key.upper(), key]
    else:
        candidates = [key, key.upper()]
    for c in candidates:
        v = brief_dict.get(c)
        if v not in ("", None):
            return v
    return default

def save_ksa_matrix_to_current_brief():
    """Sauvegarde robuste de la KSA en JSON (accepte anciens noms de colonnes)."""
    bname = st.session_state.get("current_brief_name")
    if not bname:
        return
    df = st.session_state.get("ksa_matrix")
    if not isinstance(df, pd.DataFrame) or df.empty:
        return
    df = df.copy()
    # Harmonisation colonnes
    if "Cible / Standard attendu" in df.columns and "Question pour l'entretien" not in df.columns:
        df["Question pour l'entretien"] = df["Cible / Standard attendu"]
    if "√âchelle d'√©valuation (1-5)" in df.columns and "√âvaluation (1-5)" not in df.columns:
        df["√âvaluation (1-5)"] = df["√âchelle d'√©valuation (1-5)"]
    needed = ["Rubrique","Crit√®re","Type de question","Question pour l'entretien","√âvaluation (1-5)","√âvaluateur"]
    for c in needed:
        if c not in df.columns:
            df[c] = ""
    df = df[needed]
    brief = st.session_state.saved_briefs.get(bname, {})
    brief["KSA_MATRIX_JSON"] = df.to_json(orient="records", force_ascii=False)
    brief.pop("ksa_matrix", None)
    st.session_state.saved_briefs[bname] = brief
    save_briefs()
    save_brief_to_gsheet(bname, brief)


# ======================== ANALYSE CV AVANC√âE ========================
# Fonctions pour l'analyse combin√©e, le traitement par lots et le feedback utilisateur
# Ces fonctions √©taient pr√©c√©demment dans utils_cv_analyzer.py

# -------------------- Configuration --------------------

# -------------------- Fonctions pour l'analyse combin√©e et le feedback --------------------
def rank_resumes_with_ensemble(job_description, resumes, file_names, 
                              cosinus_weight=0.2, semantic_weight=0.4, rules_weight=0.4,
                              cosine_func=None, semantic_func=None, rules_func=None):
    """
    Combine les r√©sultats de plusieurs m√©thodes d'analyse pour obtenir un score global pond√©r√©.
    
    Args:
        job_description: Description du poste √† pourvoir
        resumes: Liste des textes de CV √† analyser
        file_names: Liste des noms de fichiers correspondants
        cosinus_weight: Poids de la m√©thode cosinus (0 √† 1)
        semantic_weight: Poids de la m√©thode s√©mantique (0 √† 1)
        rules_weight: Poids de la m√©thode par r√®gles (0 √† 1)
        cosine_func: Fonction d'analyse cosinus (obligatoire)
        semantic_func: Fonction d'analyse s√©mantique (obligatoire)
        rules_func: Fonction d'analyse par r√®gles (obligatoire)
        
    Returns:
        Dict contenant les scores combin√©s et les explications d√©taill√©es
    """
    # V√©rification des fonctions n√©cessaires
    if cosine_func is None or semantic_func is None or rules_func is None:
        raise ValueError("Les fonctions d'analyse doivent √™tre fournies pour √©viter les r√©f√©rences circulaires")
    
    # Calcul des poids normalis√©s
    total_weight = cosinus_weight + semantic_weight + rules_weight
    if total_weight == 0:
        cosinus_weight, semantic_weight, rules_weight = 0.33, 0.33, 0.34
    else:
        cosinus_weight /= total_weight
        semantic_weight /= total_weight
        rules_weight /= total_weight
    
    # Analyse avec chaque m√©thode
    with st.spinner("Analyse par m√©thode cosinus..."):
        cosine_results = cosine_func(job_description, resumes, file_names)
        cosine_scores = cosine_results.get("scores", [0] * len(file_names))
        cosine_logic = cosine_results.get("logic", {})
    
    with st.spinner("Analyse par m√©thode s√©mantique..."):
        semantic_results = semantic_func(job_description, resumes, file_names)
        semantic_scores = semantic_results.get("scores", [0] * len(file_names))
        semantic_logic = semantic_results.get("logic", {})
    
    with st.spinner("Analyse par r√®gles..."):
        rule_results = rules_func(job_description, resumes, file_names)
        rule_scores = [r["score"] for r in rule_results]
        rule_logic = {r["file_name"]: r["logic"] for r in rule_results}
    
    # Combinaison des scores
    combined_scores = []
    for i in range(len(file_names)):
        cosine_score = cosine_scores[i] if i < len(cosine_scores) else 0
        semantic_score = semantic_scores[i] if i < len(semantic_scores) else 0
        rule_score = rule_scores[i] if i < len(rule_scores) else 0
        
        weighted_score = (
            cosine_score * cosinus_weight +
            semantic_score * semantic_weight +
            rule_score * rules_weight
        )
        combined_scores.append(weighted_score)
    
    # Construction d'une logique combin√©e
    combined_logic = {}
    for i, file_name in enumerate(file_names):
        cosine_score = cosine_scores[i] if i < len(cosine_scores) else 0
        semantic_score = semantic_scores[i] if i < len(semantic_scores) else 0
        rule_score = rule_scores[i] if i < len(rule_scores) else 0
        
        combined_logic[file_name] = {
            "M√©thode Cosinus": {
                "Score": f"{cosine_score*100:.1f}%",
                "Poids": f"{cosinus_weight:.2f}",
                "Contribution": f"{cosine_score*cosinus_weight*100:.1f}%",
                "D√©tails": cosine_logic.get(file_name, {})
            },
            "M√©thode S√©mantique": {
                "Score": f"{semantic_score*100:.1f}%",
                "Poids": f"{semantic_weight:.2f}",
                "Contribution": f"{semantic_score*semantic_weight*100:.1f}%",
                "D√©tails": semantic_logic.get(file_name, {})
            },
            "M√©thode R√®gles": {
                "Score": f"{rule_score*100:.1f}%",
                "Poids": f"{rules_weight:.2f}",
                "Contribution": f"{rule_score*rules_weight*100:.1f}%",
                "D√©tails": rule_logic.get(file_name, {})
            },
            "Score Final": f"{combined_scores[i]*100:.1f}%"
        }
    
    return {"scores": combined_scores, "logic": combined_logic}

def batch_process_resumes(job_description, file_list, analysis_method, 
                          batch_size=10, progress_callback=None, 
                          extract_text_from_pdf_func=None, rank_resumes_funcs=None):
    """
    Traite un grand nombre de CVs par lots pour √©viter les probl√®mes de m√©moire.
    
    Args:
        job_description: Description du poste
        file_list: Liste de fichiers CV (objets UploadedFile de Streamlit)
        analysis_method: M√©thode d'analyse √† utiliser
        batch_size: Taille des lots pour le traitement
        progress_callback: Fonction appel√©e pour mettre √† jour la progression
        extract_text_from_pdf_func: Fonction pour extraire le texte des PDFs (obligatoire)
        rank_resumes_funcs: Dictionnaire des fonctions d'analyse (obligatoire)
            {'cosine': func, 'embeddings': func, 'rules': func, 'ai': func, 'ensemble': func}
        
    Returns:
        Dict contenant les r√©sultats combin√©s de tous les lots
    """
    # V√©rification des fonctions n√©cessaires
    if extract_text_from_pdf_func is None or rank_resumes_funcs is None:
        raise ValueError("Les fonctions d'analyse et d'extraction doivent √™tre fournies pour √©viter les r√©f√©rences circulaires")
    
    # Initialisation des r√©sultats
    all_scores = []
    all_file_names = []
    all_logic = {}
    all_explanations = {}
    
    # D√©terminer la fonction d'analyse appropri√©e
    if analysis_method == "Analyse par IA (DeepSeek)":
        analysis_func = rank_resumes_funcs.get('ai')
    elif analysis_method == "Scoring par R√®gles (Regex)":
        analysis_func = lambda jd, res, fnames: rank_resumes_funcs.get('rules')(jd, res, fnames)
    elif analysis_method == "M√©thode S√©mantique (Embeddings)":
        analysis_func = rank_resumes_funcs.get('embeddings')
    elif analysis_method == "Analyse combin√©e (Ensemble)":
        analysis_func = rank_resumes_funcs.get('ensemble')
    else:  # Par d√©faut, m√©thode cosinus
        analysis_func = rank_resumes_funcs.get('cosine')
    
    # Traitement par lots
    num_batches = (len(file_list) + batch_size - 1) // batch_size
    
    for batch_idx in range(num_batches):
        start_idx = batch_idx * batch_size
        end_idx = min((batch_idx + 1) * batch_size, len(file_list))
        batch_files = file_list[start_idx:end_idx]
        
        # Extraction du texte des CVs
        batch_resumes = []
        batch_file_names = []
        
        for file in batch_files:
            text = extract_text_from_pdf_func(file)
            if "Erreur" not in text:
                batch_resumes.append(text)
                batch_file_names.append(file.name)
        
        # Analyse du lot
        if batch_resumes:
            if analysis_method == "Analyse par IA (DeepSeek)":
                results = analysis_func(job_description, batch_resumes, batch_file_names)
                batch_scores = results.get("scores", [])
                batch_explanations = results.get("explanations", {})
                all_explanations.update(batch_explanations)
            elif analysis_method == "Scoring par R√®gles (Regex)":
                rule_results = analysis_func(job_description, batch_resumes, batch_file_names)
                batch_scores = [r["score"] for r in rule_results]
                batch_logic = {r["file_name"]: r["logic"] for r in rule_results}
                all_logic.update(batch_logic)
            elif analysis_method == "Analyse combin√©e (Ensemble)":
                results = analysis_func(job_description, batch_resumes, batch_file_names)
                batch_scores = results.get("scores", [])
                batch_logic = results.get("logic", {})
                all_logic.update(batch_logic)
            else:
                results = analysis_func(job_description, batch_resumes, batch_file_names)
                batch_scores = results.get("scores", [])
                batch_logic = results.get("logic", {})
                all_logic.update(batch_logic)
            
            all_scores.extend(batch_scores)
            all_file_names.extend(batch_file_names)
        
        # Mise √† jour de la progression
        if progress_callback:
            progress_callback((batch_idx + 1) / num_batches)
    
    # Construction des r√©sultats finaux
    final_results = {
        "scores": all_scores,
        "logic": all_logic
    }
    
    if all_explanations:
        final_results["explanations"] = all_explanations
    
    return final_results, all_file_names

# -------------------- FONCTIONS DE GESTION DU FEEDBACK --------------------

# Configuration pour Google Sheets Feedback
FEEDBACK_DATA_PATH = "feedback_data.json"
FEEDBACK_GSHEET_NAME = "TG_Hire_Feedback_Analytics"  # Nom correct de l'onglet
FEEDBACK_GSHEET_URL = "https://docs.google.com/spreadsheets/d/1FBeN0s7ESjZ6BPoG4iB4VQ6w-MfRL1GWBGEkbqPR0gI/edit"

def get_feedback_google_credentials():
    """Cr√©e les identifiants √† partir des secrets Streamlit."""
    try:
        service_account_info = {
            "type": st.secrets["GCP_TYPE"],
            "project_id": st.secrets["GCP_PROJECT_ID"],
            "private_key_id": st.secrets.get("GCP_PRIVATE_KEY_ID", ""),
            "private_key": st.secrets["GCP_PRIVATE_KEY"].replace('\\n', '\n'),
            "client_email": st.secrets["GCP_CLIENT_EMAIL"],
            "client_id": st.secrets.get("GCP_CLIENT_ID", ""),
            "auth_uri": st.secrets.get("GCP_AUTH_URI", "https://accounts.google.com/o/oauth2/auth"),
            "token_uri": st.secrets["GCP_TOKEN_URI"],
            "auth_provider_x509_cert_url": st.secrets.get("GCP_AUTH_PROVIDER_CERT_URL", "https://www.googleapis.com/oauth2/v1/certs"),
            "client_x509_cert_url": st.secrets.get("GCP_CLIENT_CERT_URL", "")
        }
        return service_account.Credentials.from_service_account_info(service_account_info)
    except Exception as e:
        st.error(f"‚ùå Erreur de format des secrets Google: {e}")
        return None

def get_feedback_gsheet_client():
    """Authentification pour Google Sheets."""
    try:
        creds = get_feedback_google_credentials()
        if creds and GSPREAD_AVAILABLE:
            scoped_creds = creds.with_scopes([
                "https://www.googleapis.com/auth/spreadsheets",
                "https://www.googleapis.com/auth/drive",
            ])
            gc = gspread.authorize(scoped_creds)
            return gc
    except Exception as e:
        st.error(f"‚ùå Erreur d'authentification Google Sheets: {str(e)}")
    return None

def save_feedback(analysis_method, job_title, job_description_snippet, cv_count, feedback_score, feedback_text="",
                 user_criteria=None, improvement_suggestions=None):
    """
    Sauvegarde un feedback utilisateur localement et dans Google Sheets si disponible.

    Args:
        analysis_method: M√©thode d'analyse utilis√©e (Cosinus, S√©mantique, R√®gles, IA, Ensemble)
        job_title: Intitul√© du poste analys√©
        job_description_snippet: Extrait de la description du poste (200 premiers caract√®res)
        cv_count: Nombre de CV analys√©s dans cette session
        feedback_score: Note de satisfaction (1-5)
        feedback_text: Commentaires textuels de l'utilisateur (optionnel)
        user_criteria: Crit√®res d'√©valuation sp√©cifiques mentionn√©s par l'utilisateur
        improvement_suggestions: Suggestions d'am√©lioration

    Returns:
        Boolean indiquant si le feedback a √©t√© sauvegard√© avec succ√®s
    """
    # Pr√©paration des donn√©es de feedback
    timestamp = datetime.now().strftime("%Y-%m-%d %H:%M:%S")
    app_version = "1.0.0"

    feedback_entry = {
        "timestamp": timestamp,
        "analysis_method": analysis_method,
        "job_title": job_title,
        "job_description_snippet": job_description_snippet[:500] if job_description_snippet else "",  # √âtendu √† 500 caract√®res
        "cv_count": cv_count,
        "feedback_score": feedback_score,
        "feedback_text": feedback_text,
        "user_criteria": user_criteria or "",
        "improvement_suggestions": improvement_suggestions or "",
        "version_app": app_version
    }
    
    # Sauvegarde locale
    feedback_data = []
    if os.path.exists(FEEDBACK_DATA_PATH):
        try:
            with open(FEEDBACK_DATA_PATH, 'r', encoding='utf-8') as f:
                feedback_data = json.load(f)
        except:
            feedback_data = []
    
    feedback_data.append(feedback_entry)
    
    try:
        with open(FEEDBACK_DATA_PATH, 'w', encoding='utf-8') as f:
            json.dump(feedback_data, f, ensure_ascii=False, indent=2)
    except Exception as e:
        print(f"Erreur lors de la sauvegarde locale du feedback: {e}")
        return False
    
    # Sauvegarde dans Google Sheets si disponible
    try:
        if GSPREAD_AVAILABLE:
            gc = get_feedback_gsheet_client()
            if gc:
                # Ouvrir la feuille Google Sheets par URL
                sh = gc.open_by_url(FEEDBACK_GSHEET_URL)
                
                # Acc√©der √† la feuille par nom
                try:
                    worksheet = sh.worksheet(FEEDBACK_GSHEET_NAME)
                except gspread.exceptions.WorksheetNotFound:
                # Cr√©er l'onglet s'il n'existe pas
                    worksheet = sh.add_worksheet(title=FEEDBACK_GSHEET_NAME, rows=1000, cols=10)
                    # Ajouter les en-t√™tes
                    headers = [
                        "timestamp", "analysis_method", "job_title", "job_description_snippet",
                        "cv_count", "feedback_score", "feedback_text", "Crit√®res √©valu√©s :",
                        "version_app"
                    ]
                    worksheet.update('A1:I1', [headers])
                
                # Pr√©parer les donn√©es √† ajouter
                # Construction explicite pour garantir version_app en colonne J (index 9)
                row_data = [
                    str(feedback_entry.get("timestamp", "")),
                    str(feedback_entry.get("analysis_method", "")),
                    str(feedback_entry.get("job_title", "")),
                    str(feedback_entry.get("job_description_snippet", "")),
                    str(feedback_entry.get("cv_count", "")),
                    str(feedback_entry.get("feedback_score", "")),
                    str(feedback_entry.get("feedback_text", "")),
                    str(", ".join(feedback_entry.get("user_criteria", [])) if isinstance(feedback_entry.get("user_criteria", []), list) else feedback_entry.get("user_criteria", "")),
                    str(feedback_entry.get("version_app", ""))
                ]
                # Tronquer √† 9 colonnes si jamais il y a plus
                row_data = row_data[:9]
                # Compl√©ter si moins de 9 colonnes
                while len(row_data) < 9:
                    row_data.append("")
                
                # Ajouter la ligne en for√ßant l'insertion √† la premi√®re colonne d'une nouvelle ligne
                last_row = len(worksheet.get_all_values()) + 1
                worksheet.insert_row(row_data, index=last_row)
                st.success("‚úÖ Feedback envoy√© avec succ√®s √† Google Sheets")
                return True
                
    except Exception as e:
        st.error(f"‚ùå Erreur lors de la sauvegarde du feedback dans Google Sheets : {e}")
        # Ne pas faire √©chouer si Google Sheets n'est pas disponible
        return True
    
    return True

def get_average_feedback_score(analysis_method=None):
    """
    R√©cup√®re le score moyen de feedback pour une m√©thode donn√©e ou pour toutes les m√©thodes.
    
    Args:
        analysis_method: M√©thode d'analyse sp√©cifique (ou None pour toutes les m√©thodes)
        
    Returns:
        Score moyen et nombre d'√©valuations
    """
    if not os.path.exists(FEEDBACK_DATA_PATH):
        return 0, 0
    
    try:
        with open(FEEDBACK_DATA_PATH, 'r', encoding='utf-8') as f:
            feedback_data = json.load(f)
        
        if analysis_method:
            relevant_feedback = [f for f in feedback_data if f["analysis_method"] == analysis_method]
        else:
            relevant_feedback = feedback_data
        
        if not relevant_feedback:
            return 0, 0
        
        total_score = sum(f["feedback_score"] for f in relevant_feedback)
        return total_score / len(relevant_feedback), len(relevant_feedback)
    
    except Exception as e:
        print(f"Erreur lors de la r√©cup√©ration des scores de feedback: {e}")
        return 0, 0

def get_feedback_summary():
    """
    G√©n√®re un r√©sum√© des feedbacks re√ßus pour chaque m√©thode d'analyse.
    
    Returns:
        DataFrame contenant les statistiques de feedback
    """
    methods = [
        "M√©thode Cosinus (Mots-cl√©s)",
        "M√©thode S√©mantique (Embeddings)",
        "Scoring par R√®gles (Regex)",
        "Analyse combin√©e (Ensemble)",
        "Analyse par IA (DeepSeek)"
    ]
    
    summary_data = []
    
    for method in methods:
        avg_score, count = get_average_feedback_score(method)
        summary_data.append({
            "M√©thode": method,
            "Score moyen": round(avg_score, 2),
            "Nombre d'√©valuations": count
        })
    
    # Ajouter le score global
    overall_avg, overall_count = get_average_feedback_score()
    summary_data.append({
        "M√©thode": "Toutes m√©thodes confondues",
        "Score moyen": round(overall_avg, 2),
        "Nombre d'√©valuations": overall_count
    })
    
    return pd.DataFrame(summary_data)

# -------------------- FONCTIONS GOOGLE SHEETS pour les Annonces --------------------

@st.cache_resource
def get_annonces_gsheet_client():
    """Initialise et retourne le client gspread pour les annonces."""
    if not GSPREAD_AVAILABLE:
        return None
    try:
        service_account_info = _build_service_account_info_from_st_secrets()
        gc = gspread.service_account_from_dict(service_account_info)
        return gc
    except KeyError as e:
        st.error(f"‚ùå Cl√© de secret manquante pour Google Sheets: {e}. V√©rifiez la configuration des secrets GCP_...")
        return None
    except Exception as e:
        st.error(f"‚ùå Erreur de connexion/ouverture de Google Sheets: {e}")
        return None

def save_annonce_to_gsheet(annonce_data):
    """Sauvegarde une annonce dans Google Sheets."""
    try:
        gc = get_annonces_gsheet_client()
        if not gc:
            return False
        
        sheet = gc.open_by_url(ANNONCES_SHEET_URL)
        worksheet = sheet.worksheet(ANNONCES_WORKSHEET_NAME)
        
        # Pr√©parer les donn√©es dans l'ordre des colonnes
        # Calculer fields_present (True/False pour chaque champ utile)
        keys_to_check = [
            "poste", "entreprise", "localisation", "plateforme", "contenu", "fiche_text",
            "type_contrat", "niveau_experience", "formation_requise", "affectation",
            "competences_techniques", "missions_principales", "soft_skills"
        ]
        fields_present = {k: bool(annonce_data.get(k)) for k in keys_to_check}

        row_data = [
            annonce_data.get("timestamp", ""),
            annonce_data.get("poste", ""),
            annonce_data.get("entreprise", "TGCC"),
            annonce_data.get("localisation", ""),
            annonce_data.get("plateforme", ""),
            annonce_data.get("format_type", annonce_data.get("plateforme", "")),
            annonce_data.get("contenu", ""),
            annonce_data.get("fiche_text", ""),
            annonce_data.get("type_contrat", ""),
            annonce_data.get("niveau_experience", ""),
            annonce_data.get("formation_requise", ""),
            annonce_data.get("affectation", ""),
            annonce_data.get("competences_techniques", ""),
            annonce_data.get("missions_principales", ""),
            annonce_data.get("soft_skills", ""),
            json.dumps(fields_present, ensure_ascii=False)
        ]
        
        # Ajouter la ligne
        worksheet.append_row(row_data)
        return True
        
    except Exception as e:
        st.error(f"‚ùå Erreur sauvegarde annonce Google Sheets: {e}")
        return False

def load_annonces_from_gsheet():
    """Charge toutes les annonces depuis Google Sheets."""
    try:
        gc = get_annonces_gsheet_client()
        if not gc:
            return []
        
        sheet = gc.open_by_url(ANNONCES_SHEET_URL)
        worksheet = sheet.worksheet(ANNONCES_WORKSHEET_NAME)
        
        # R√©cup√©rer toutes les donn√©es
        records = worksheet.get_all_records()
        return records
        
    except Exception as e:
        st.warning(f"Erreur lors du chargement des annonces depuis Google Sheets : {e}")
        return []

def delete_annonce_from_gsheet(annonce):
    """Supprime une annonce dans Google Sheets en fonction du timestamp et du poste."""
    try:
        gc = get_annonces_gsheet_client()
        if not gc:
            return False
        sheet = gc.open_by_url(ANNONCES_SHEET_URL)
        worksheet = sheet.worksheet(ANNONCES_WORKSHEET_NAME)
        # On identifie la ligne √† supprimer par le timestamp et le poste
        records = worksheet.get_all_records()
        for idx, r in enumerate(records):
            if r.get("timestamp") == annonce.get("date") and r.get("poste") == annonce.get("poste"):
                worksheet.delete_rows(idx + 2)  # +2 car get_all_records saute l'en-t√™te et est 0-index√©
                return True
        return False
    except Exception as e:
        st.error(f"‚ùå Erreur lors de la suppression dans Google Sheets : {e}")
        return False